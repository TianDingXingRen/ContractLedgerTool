"""Rule-based procurement Word documents for the first release."""

from __future__ import annotations

from decimal import Decimal
import json
import re
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

import procurement_store
from services import procurement_file_service
from utils.constants import PROCUREMENT_METHOD_LABELS


def _money(value):
    return f'{Decimal(int(value or 0)) / 100:,.2f}'


def _strip_text(value):
    return str(value or '').strip()


def _safe_docx_filename(value, fallback):
    raw = _strip_text(value) or fallback
    stem = raw[:-5] if raw.lower().endswith('.docx') else raw
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', stem).strip(' ._')
    if not stem:
        stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', fallback).strip(' ._')
    stem = stem[:80].strip(' ._') or '谈判预案'
    return f'{stem}.docx'


_DOCX_CJK_FONT = '仿宋'


def _set_element_fonts(element, font_name=_DOCX_CJK_FONT):
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        r_fonts.set(qn(attr), font_name)


def _set_run_font(run, font_name=_DOCX_CJK_FONT):
    run.font.name = font_name
    _set_element_fonts(run._element, font_name)


def _set_paragraph_font(paragraph, font_name=_DOCX_CJK_FONT):
    for run in paragraph.runs:
        _set_run_font(run, font_name)


def _apply_document_font(document, font_name=_DOCX_CJK_FONT):
    for paragraph in document.paragraphs:
        _set_paragraph_font(paragraph, font_name)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_paragraph_font(paragraph, font_name)


def _set_default_font(document):
    style = document.styles['Normal']
    style.font.name = _DOCX_CJK_FONT
    style.font.size = Pt(10.5)
    _set_element_fonts(style._element)


def _title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _save_and_register(document, project, file_type, filename):
    path = procurement_file_service.target_path(project, file_type, filename)
    _apply_document_font(document)
    document.save(path)
    procurement_store.register_project_file(
        project['id'], file_type, procurement_file_service.relative_path(path), filename,
        procurement_file_service.sha256_file(path), path.stat().st_size,
    )
    return str(path)


def _append_negotiation_signature_box(document):
    document.add_paragraph()
    table = document.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    title_cell = table.rows[0].cells[0].merge(table.rows[0].cells[-1])
    title_cell.text = '谈判人员签字'
    for cell, label in zip(table.rows[1].cells, ['采购人员', '需求部门', '技术人员', '供应商代表']):
        cell.text = label
    for cell in table.rows[2].cells:
        cell.text = '\n\n'


def _default_delivery_requirement(project, items):
    if project.get('delivery_requirement'):
        return f"本次采购产品交付要求为{project['delivery_requirement']}。"
    dates = [item.get('required_delivery_date') for item in items if item.get('required_delivery_date')]
    if dates:
        return f"本次采购产品最晚交付时间为{max(dates)}前。"
    return '本次采购产品交付时间按采购需求及谈判确认结果执行。'


def _default_target_price_note(project, items, quotes):
    parts = []
    if project.get('target_price_minor') is not None:
        parts.append(f"本项目目标价格不超过{_money(project['target_price_minor'])}元。")
    elif project.get('budget_minor') is not None:
        parts.append(f"本项目预算金额为{_money(project['budget_minor'])}元。")

    confirmed_quotes = [quote for quote in quotes if quote.get('total_amount_minor') is not None]
    if confirmed_quotes:
        lowest = min(int(quote['total_amount_minor']) for quote in confirmed_quotes)
        parts.append(f"已确认供应商报价中最低总价为{_money(lowest)}元。")

    try:
        from services import historical_price_service
        history_notes = []
        seen = set()
        for item in items:
            item_name = item.get('item_name') or ''
            if not item_name or item_name in seen:
                continue
            seen.add(item_name)
            history = historical_price_service.price_assistance(item_name)
            if history.get('median_minor') is not None:
                history_notes.append(
                    f"{item_name}历史成交中位单价为{_money(history['median_minor'])}元"
                )
            if len(history_notes) >= 3:
                break
        if history_notes:
            parts.append('历史价格参考：' + '；'.join(history_notes) + '。')
    except Exception:
        pass

    if parts:
        return ''.join(parts)
    return '本次目标价格以项目预算、历史采购价格、供应商报价及市场行情为谈判依据。'


def _default_quote_round_note(quotes):
    if not quotes:
        return '拟一轮报价，谈判会上确认最终报价；如报价与目标价偏离较大，可视情况组织补充报价。'
    rounds = sorted({int(quote.get('quote_round') or 1) for quote in quotes})
    return f"已完成{len(rounds)}轮报价导入，本次谈判会上确认最终报价。"


def _default_evaluation_plan():
    return '\n'.join([
        '1）甲方技术专家认可供应商配套能力。',
        '2）供应商交付周期满足甲方要求。',
        '3）供应商报价与目标价格或历史参考价格偏离度绝对值过大，或明显低于产品本身成本价格，考虑否决。',
    ])


def negotiation_plan_defaults(project_id):
    project = procurement_store.get_project(project_id)
    if not project:
        raise ValueError('采购项目不存在')
    items = procurement_store.list_project_items(project_id)
    suppliers = procurement_store.list_project_suppliers(project_id)
    quotes = procurement_store.list_quotes(project_id)
    method_label = PROCUREMENT_METHOD_LABELS.get(
        project.get('purchase_method'), project.get('purchase_method') or ''
    )
    supplier_names = '、'.join(row['supplier_name'] for row in suppliers) or '候选供应商'
    background = _strip_text(project.get('remark'))
    if not background:
        background = (
            f"{project['project_name']}，本项目拟采用{method_label or '谈判'}方式组织采购。"
            f"结合采购需求、{supplier_names}响应情况和目标价格制定本预案。"
        )
    return {
        'project': project,
        'items': items,
        'suppliers': suppliers,
        'quotes': quotes,
        'plan': {
            'title': f"{project['project_name']}谈判预案",
            'filename': f"{project['project_name']}谈判预案.docx",
            'purchase_requirement': f"关于{project['project_name']}采购申请单。",
            'request_no': project['project_no'],
            'project_background': background,
            'delivery_requirement': _default_delivery_requirement(project, items),
            'target_price_note': _default_target_price_note(project, items, quotes),
            'quote_round_note': _default_quote_round_note(quotes),
            'evaluation_plan': _default_evaluation_plan(),
            'fixed_asset': '否',
            'purchase_method_label': method_label,
        },
    }


def _negotiation_plan_from_form(defaults, form):
    plan = dict(defaults['plan'])
    if not form:
        return plan
    for key in [
        'title', 'filename', 'purchase_requirement', 'request_no', 'project_background',
        'delivery_requirement', 'target_price_note', 'quote_round_note',
        'evaluation_plan', 'fixed_asset',
    ]:
        value = _strip_text(form.get(key))
        if value:
            plan[key] = value
    return plan


def generate_negotiation_plan(project_id, form=None, return_info=False):
    defaults = negotiation_plan_defaults(project_id)
    project = defaults['project']
    items = defaults['items']
    if not items:
        raise ValueError('请先录入采购明细')
    plan = _negotiation_plan_from_form(defaults, form)

    document = Document()
    _set_default_font(document)
    _title(document, plan['title'])

    document.add_heading('一、原则', level=2)
    document.add_paragraph('1.采购需求：')
    document.add_paragraph(plan['purchase_requirement'])
    document.add_paragraph(f"编号：{plan['request_no']}。")
    document.add_paragraph('2.项目背景：')
    document.add_paragraph(plan['project_background'])

    document.add_heading('二、采购标的', level=2)
    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    for cell, label in zip(table.rows[0].cells, ['序号', '产品名称', '图代号', '是否固定资产', '数量', '单位', '备注']):
        cell.text = label
    for index, item in enumerate(items, start=1):
        cells = table.add_row().cells
        remark = item.get('remark') or item.get('technical_requirement') or ''
        values = [
            index,
            item['item_name'],
            item.get('drawing_no') or item.get('spec_model') or '',
            plan['fixed_asset'],
            item.get('quantity_text') or '',
            item.get('unit') or '',
            remark,
        ]
        for cell, value in zip(cells, values):
            cell.text = str(value)

    document.add_heading('三、生产周期要求', level=2)
    document.add_paragraph(plan['delivery_requirement'])

    document.add_heading('四、目标价格：', level=2)
    document.add_paragraph('金额：' + plan['target_price_note'])

    document.add_heading('五、报价轮次', level=2)
    document.add_paragraph(plan['quote_round_note'])

    document.add_heading('六、评价方案', level=2)
    for line in plan['evaluation_plan'].splitlines():
        text = line.strip()
        if text:
            document.add_paragraph(text)

    filename = _safe_docx_filename(plan.get('filename'), f"{project['project_name']}谈判预案")
    path = _save_and_register(document, project, 'negotiation_plan', filename)
    if return_info:
        return {'path': path, 'download_name': filename}
    return path


def generate_inquiry_letter(project_id):
    project = procurement_store.get_project(project_id)
    items = procurement_store.list_project_items(project_id)
    suppliers = procurement_store.list_project_suppliers(project_id)
    if not project:
        raise ValueError('采购项目不存在')
    if not items:
        raise ValueError('请先录入采购明细')
    document = Document()
    _set_default_font(document)
    _title(document, '询价函')
    document.add_paragraph(f"项目编号：{project['project_no']}")
    document.add_paragraph(f"项目名称：{project['project_name']}")
    document.add_paragraph(f"需求部门：{project.get('demand_department') or '-'}")
    document.add_paragraph('现就以下采购内容进行询价，请按随附标准报价模板完整填报并在有效期内反馈。')
    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    for cell, label in zip(table.rows[0].cells, ['序号', '物资名称', '规格型号', '图号', '数量', '单位', '要求交期']):
        cell.text = label
    for item in items:
        cells = table.add_row().cells
        values = [item['line_no'], item['item_name'], item.get('spec_model') or '', item.get('drawing_no') or '',
                  item['quantity_text'], item['unit'], item.get('required_delivery_date') or '']
        for cell, value in zip(cells, values):
            cell.text = str(value)
    document.add_heading('商务要求', level=2)
    document.add_paragraph(f"交付地点：{project.get('delivery_place') or '-'}")
    document.add_paragraph(f"交付周期：{project.get('delivery_requirement') or '-'}")
    document.add_paragraph(f"付款条件：{project.get('payment_requirement') or '-'}")
    if suppliers:
        document.add_paragraph('候选供应商：' + '、'.join(row['supplier_name'] for row in suppliers))
    path = _save_and_register(document, project, 'inquiry', f"{project['project_no']}_询价函.docx")
    if project['status'] == 'draft':
        procurement_store.transition_project_status(project_id, 'documents_ready', '已生成询价函')
    return path


def generate_clarification_letter(project_id, supplier_id=None):
    project = procurement_store.get_project(project_id)
    questions = procurement_store.list_clarifications(project_id)
    if not project:
        raise ValueError('采购项目不存在')
    if supplier_id:
        questions = [row for row in questions if row.get('supplier_id') == int(supplier_id)]
    if not questions:
        raise ValueError('尚未生成澄清问题')
    document = Document()
    _set_default_font(document)
    _title(document, '报价澄清函' if supplier_id else '报价澄清问题清单')
    document.add_paragraph(f"项目：{project['project_no']} / {project['project_name']}")
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    for cell, label in zip(table.rows[0].cells, ['序号', '供应商', '相关物资', '问题类型', '澄清问题', '供应商回复']):
        cell.text = label
    for index, question in enumerate(questions, start=1):
        cells = table.add_row().cells
        values = [index, question.get('supplier_name') or '通用', question.get('item_name') or '',
                  question['question_type'], question['question_text'], question.get('answer_text') or '']
        for cell, value in zip(cells, values):
            cell.text = str(value)
    supplier_suffix = f"_{questions[0].get('supplier_name')}" if supplier_id else ''
    return _save_and_register(
        document, project, 'clarification',
        f"{project['project_no']}{supplier_suffix}_{'澄清函' if supplier_id else '澄清问题清单'}.docx"
    )


def generate_award_recommendation(project_id):
    project = procurement_store.get_project(project_id)
    award = procurement_store.get_latest_award(project_id)
    if not project:
        raise ValueError('采购项目不存在')
    if not award:
        raise ValueError('尚未确认成交建议')
    document = Document()
    _set_default_font(document)
    _title(document, '成交建议')
    document.add_paragraph(f"项目编号：{project['project_no']}")
    document.add_paragraph(f"项目名称：{project['project_name']}")
    document.add_paragraph(f"推荐成交供应商：{award.get('supplier_summary') or award['supplier_name']}")
    document.add_paragraph(f"推荐成交金额：人民币 {_money(award['recommended_amount_minor'])} 元")
    document.add_heading('成交明细', level=2)
    split = bool(award.get('is_split'))
    table = document.add_table(rows=1, cols=7 if split else 6)
    table.style = 'Table Grid'
    labels = ['序号', '物资名称', '规格型号', '数量', '单价', '金额']
    if split:
        labels.insert(1, '成交供应商')
    for cell, label in zip(table.rows[0].cells, labels):
        cell.text = label
    for index, item in enumerate(award['items'], start=1):
        cells = table.add_row().cells
        values = [index, item['item_name'], item.get('spec_model') or '',
                  f"{item['quantity_text']} {item['unit']}", _money(item['unit_price_minor']),
                  _money(item['amount_minor'])]
        if split:
            values.insert(1, item.get('supplier_name') or '')
        for cell, value in zip(cells, values):
            cell.text = str(value)
    sections = [
        ('推荐理由', award.get('reason_summary')),
        ('价格合理性说明', award.get('price_reason')),
        ('技术响应说明', award.get('technical_reason')),
        ('商务响应说明', award.get('commercial_reason')),
        ('交付保障说明', award.get('delivery_reason')),
        ('未选择最低价说明', award.get('lowest_price_not_selected_reason')),
        ('风险提示', award.get('risk_note')),
        ('合同签订注意事项', award.get('contract_notice')),
    ]
    for heading, text in sections:
        if text:
            document.add_heading(heading, level=2)
            document.add_paragraph(text)
    return _save_and_register(
        document, project, 'award', f"{project['project_no']}_成交建议.docx"
    )


def export_project_items(project_id):
    project = procurement_store.get_project(project_id)
    items = procurement_store.list_project_items(project_id)
    if not project:
        raise ValueError('采购项目不存在')
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = '采购明细'
    headers = ['物资名称', '规格型号', '图号/代号', '数量', '单位', '要求交付日期', '技术要求', '备注']
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(color='FFFFFF', bold=True)
        cell.fill = PatternFill('solid', fgColor='1D4ED8')
    for item in items:
        sheet.append([
            item['item_name'], item.get('spec_model') or '', item.get('drawing_no') or '',
            Decimal(item['quantity_text']), item['unit'], item.get('required_delivery_date') or '',
            item.get('technical_requirement') or '', item.get('remark') or '',
        ])
    for column, width in zip('ABCDEFGH', [24, 18, 16, 12, 10, 16, 32, 24]):
        sheet.column_dimensions[column].width = width
    path = procurement_file_service.target_path(
        project, 'quote_template', f"{project['project_no']}_采购明细.xlsx"
    )
    workbook.save(path)
    procurement_store.register_project_file(
        project_id, 'project_items', procurement_file_service.relative_path(path), path.name,
        procurement_file_service.sha256_file(path), path.stat().st_size,
    )
    return str(path)


def generate_negotiation_minutes(project_id):
    from services import negotiation_service
    view = negotiation_service.negotiation_view(project_id)
    project = view['project']
    if not view['rounds'] and not view['suppliers']:
        raise ValueError('尚无谈判或报价记录')
    document = Document()
    _set_default_font(document)
    _title(document, '谈判纪要')
    document.add_paragraph(f"项目编号：{project['project_no']}")
    document.add_paragraph(f"项目名称：{project['project_name']}")
    for round_item in view['rounds']:
        document.add_heading(f"第 {round_item['round_no']} 轮谈判", level=2)
        if round_item.get('meeting_date'):
            document.add_paragraph('谈判日期：' + round_item['meeting_date'])
        if round_item.get('summary'):
            document.add_paragraph(round_item['summary'])
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        for cell, label in zip(table.rows[0].cells, ['供应商', '报价金额', '交期', '付款条件', '承诺']):
            cell.text = label
        for commitment in round_item['commitments']:
            cells = table.add_row().cells
            values = [commitment['supplier_name'], _money(commitment.get('quote_amount_minor')),
                      commitment.get('delivery_period') or '', commitment.get('payment_terms') or '',
                      commitment.get('commitment') or '']
            for cell, value in zip(cells, values):
                cell.text = str(value)
    document.add_heading('报价变化摘要', level=2)
    for supplier in view['suppliers']:
        document.add_paragraph(
            f"{supplier['supplier_name']}：首轮 {_money(supplier['first_amount_minor'])} 元，"
            f"最新 {_money(supplier['latest_amount_minor'])} 元，降幅 {supplier['reduction_percent']}%。"
        )
    _append_negotiation_signature_box(document)
    return _save_and_register(
        document, project, 'negotiation', f"{project['project_no']}_谈判纪要.docx"
    )


def export_final_commitments(project_id):
    from services import negotiation_service
    view = negotiation_service.negotiation_view(project_id)
    if not view['suppliers']:
        raise ValueError('尚无供应商报价')
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = '最终承诺'
    headers = ['供应商', '首轮报价', '最新报价', '降价金额', '降幅(%)', '最终交期', '最终付款条件', '承诺说明']
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(color='FFFFFF', bold=True)
        cell.fill = PatternFill('solid', fgColor='1D4ED8')
    latest_commitments = {}
    for round_item in view['rounds']:
        for item in round_item['commitments']:
            latest_commitments[item['supplier_id']] = item
    for supplier in view['suppliers']:
        commitment = latest_commitments.get(supplier['supplier_id'], {})
        quote = supplier['latest_quote'] or {}
        first_amount = Decimal(supplier['first_amount_minor']) / 100 if supplier['first_amount_minor'] is not None else ''
        latest_amount = Decimal(supplier['latest_amount_minor']) / 100 if supplier['latest_amount_minor'] is not None else ''
        reduction_amount = Decimal(supplier['reduction_minor']) / 100 if supplier['first_amount_minor'] is not None else ''
        reduction_percent = Decimal(supplier['reduction_percent']) if supplier['first_amount_minor'] is not None else ''
        sheet.append([
            supplier['supplier_name'], first_amount,
            latest_amount, reduction_amount,
            reduction_percent, commitment.get('delivery_period') or quote.get('delivery_period') or '',
            commitment.get('payment_terms') or quote.get('payment_terms') or '', commitment.get('commitment') or '',
        ])
    project = view['project']
    path = procurement_file_service.target_path(
        project, 'negotiation', f"{project['project_no']}_最终承诺表.xlsx"
    )
    workbook.save(path)
    procurement_store.register_project_file(
        project_id, 'negotiation', procurement_file_service.relative_path(path), path.name,
        procurement_file_service.sha256_file(path), path.stat().st_size,
    )
    return str(path)


def generate_erp_oa_summary(project_id):
    project = procurement_store.get_project(project_id)
    award = procurement_store.get_latest_award(project_id)
    if not project:
        raise ValueError('采购项目不存在')
    if not award:
        raise ValueError('请先确认成交建议')
    workbook = Workbook()
    summary = workbook.active
    summary.title = 'ERP_OA摘要'
    summary.append(['字段', '值'])
    for cell in summary[1]:
        cell.font = Font(color='FFFFFF', bold=True)
        cell.fill = PatternFill('solid', fgColor='1D4ED8')
    fields = [
        ('项目编号', project['project_no']), ('项目名称', project['project_name']),
        ('需求部门', project.get('demand_department') or ''), ('经办人', project.get('owner') or ''),
        ('采购方式', PROCUREMENT_METHOD_LABELS.get(
            project.get('purchase_method'), project.get('purchase_method') or ''
        )),
        ('成交供应商', award.get('supplier_summary') or award['supplier_name']),
        ('成交金额', Decimal(award['recommended_amount_minor']) / 100),
        ('交付地点', project.get('delivery_place') or ''),
        ('交付周期', award.get('delivery_period') or project.get('delivery_requirement') or ''),
        ('付款条件', award.get('payment_terms') or project.get('payment_requirement') or ''),
        ('质保期', award.get('warranty_period') or ''),
        ('合同注意事项', award.get('contract_notice') or ''),
    ]
    for field, value in fields:
        summary.append([field, value])
    summary.column_dimensions['A'].width = 22
    summary.column_dimensions['B'].width = 60
    items_sheet = workbook.create_sheet('成交明细')
    items_sheet.append(['供应商', '物资名称', '规格型号', '数量', '单位', '单价', '金额'])
    for item in award['items']:
        items_sheet.append([
            item.get('supplier_name') or award['supplier_name'], item['item_name'],
            item.get('spec_model') or '', Decimal(item['quantity_text']), item['unit'],
            Decimal(item['unit_price_minor']) / 100, Decimal(item['amount_minor']) / 100,
        ])
    project_path = procurement_file_service.target_path(
        project, 'contract', f"{project['project_no']}_ERP_OA填报摘要.xlsx"
    )
    workbook.save(project_path)
    procurement_store.register_project_file(
        project_id, 'erp_oa_summary', procurement_file_service.relative_path(project_path),
        project_path.name, procurement_file_service.sha256_file(project_path), project_path.stat().st_size,
    )
    return str(project_path)


def generate_project_archive(project_id):
    project = procurement_store.get_project(project_id)
    if not project:
        raise ValueError('采购项目不存在')
    files = procurement_store.list_project_files(project_id)
    manifest = {
        'schema_version': '1.0', 'project': project,
        'items': procurement_store.list_project_items(project_id),
        'suppliers': procurement_store.list_project_suppliers(project_id),
        'quotes': procurement_store.list_quotes(project_id),
        'comparison': procurement_store.get_latest_comparison(project_id),
        'clarifications': procurement_store.list_clarifications(project_id),
        'negotiation_rounds': procurement_store.list_negotiation_rounds(project_id),
        'award': procurement_store.get_latest_award(project_id),
        'contract_links': procurement_store.get_project_contract_links(project_id),
        'files': files,
    }
    path = procurement_file_service.target_path(
        project, 'archive', f"{project['project_no']}_完整归档包.zip"
    )
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as archive:
        archive.writestr('manifest.json', json.dumps(manifest, ensure_ascii=False, indent=2, default=str))
        used_names = set()
        for record in files:
            if record['file_type'] == 'archive':
                continue
            try:
                source = procurement_file_service.absolute_path(record['relative_path'])
            except ValueError:
                continue
            if not source.is_file():
                continue
            name = f"{record['file_type']}/{record.get('original_name') or source.name}"
            if name in used_names:
                name = f"{record['file_type']}/v{record['version']}_{record.get('original_name') or source.name}"
            used_names.add(name)
            archive.write(source, name)
    procurement_store.register_project_file(
        project_id, 'archive', procurement_file_service.relative_path(path), path.name,
        procurement_file_service.sha256_file(path), path.stat().st_size,
    )
    return str(path)
