"""Secure preview and atomic persistence for externally authored contracts."""

from __future__ import annotations

import logging
import os
import re
import uuid
from dataclasses import asdict, dataclass, field as dataclass_field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import payment_extractor
from services.isolated_process import run_isolated_worker
from utils.field_utils import normalize_date, parse_number
from utils.file_digest import sha256_file
from utils.generation_utils import parse_contract_classification
from utils.security import validate_office_archive


_FIELD_LABELS = {
    'contract_no': ('合同编号', '合同号'),
    'title': ('合同名称', '合同标题'),
    'counterparty': ('对方单位', '乙方', '供应商', '卖方', '承包人', '服务方'),
    'amount': ('合同总价款', '合同总价', '合同金额', '含税总价', '总金额', '合同价款'),
    'sign_date': ('合同签订日期', '签订日期', '签署日期'),
    'expiry_date': ('合同到期日期', '到期日期', '有效期至', '终止日期'),
}
_FIELD_LIMITS = {
    'contract_no': 80,
    'title': 200,
    'counterparty': 120,
    'amount': 120,
    'sign_date': 80,
    'expiry_date': 80,
}
MAX_IMPORT_PARAGRAPHS = 5_000
MAX_IMPORT_TABLES = 200
MAX_IMPORT_TABLE_ROWS = 5_000
MAX_IMPORT_TABLE_COLUMNS = 100
CONTRACT_IMPORT_PARSE_TIMEOUT_SECONDS = 60
CONTRACT_IMPORT_PARSE_MEMORY_MB = 1536


@dataclass(frozen=True)
class FieldDiagnostic:
    field: str
    value: Any
    confidence: str
    evidence: str


@dataclass(frozen=True)
class ContractImportPreview:
    original_filename: str
    source_sha256: str
    summary: dict[str, Any]
    diagnostics: list[dict[str, Any]]
    plans: list[dict[str, Any]]
    rules: list[dict[str, Any]]
    warnings: list[str]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(frozen=True)
class ContractImportRequest:
    staging_path: str
    original_filename: str
    source_sha256: str
    summary: dict[str, Any]
    plans: list[dict[str, Any]]
    rules: list[dict[str, Any]] = dataclass_field(default_factory=list)


@dataclass(frozen=True)
class ContractImportResult:
    contract_id: int
    output_path: str
    plan_count: int
    rule_count: int = 0


class ContractImportService:
    def __init__(
        self,
        *,
        ledger_store,
        uploads_dir,
        output_dir,
        max_upload_bytes=50 * 1024 * 1024,
        replace_file=os.replace,
        after_commit=None,
        isolate_preview=True,
    ):
        self.ledger_store = ledger_store
        self.uploads_dir = Path(uploads_dir).resolve()
        self.output_dir = Path(output_dir).resolve()
        self.max_upload_bytes = int(max_upload_bytes)
        self.replace_file = replace_file
        self.after_commit = after_commit or (lambda _result: None)
        self.isolate_preview = bool(isolate_preview)

    @staticmethod
    def sha256_file(path) -> str:
        return sha256_file(path)

    @staticmethod
    def _clean(value, limit=200) -> str:
        text = re.sub(r'\s+', ' ', str(value or '')).strip(' \t|：:，,；;。')
        return text[:limit]

    @staticmethod
    def _confidence(score: int) -> str:
        if score >= 90:
            return 'high'
        if score >= 70:
            return 'medium'
        return 'low'

    @staticmethod
    def _candidate(candidates, field, value, score, evidence):
        limit = _FIELD_LIMITS[field]
        value = ContractImportService._clean(value, limit)
        if value and not re.search(r'\{[^{}\r\n]{1,100}\}', value):
            candidates[field].append({
                'value': value,
                'score': score,
                'evidence': ContractImportService._clean(evidence, 160),
            })

    def _extract_document(self, path):
        document = Document(path)
        candidates = {field: [] for field in _FIELD_LABELS}
        paragraphs = []
        tables = []
        payment_blocks = []
        warnings = []
        try:
            document_paragraphs = document.paragraphs
            if len(document_paragraphs) > MAX_IMPORT_PARAGRAPHS:
                warnings.append('合同段落过多，仅分析前 5000 段')
            for index, paragraph in enumerate(
                document_paragraphs[:MAX_IMPORT_PARAGRAPHS]
            ):
                text = self._clean(paragraph.text, 2000)
                if not text:
                    continue
                paragraphs.append(text)
                payment_blocks.append({
                    'kind': 'paragraph', 'index': index + 1, 'text': text,
                })
                if index < 25 and '合同' in text and len(text) <= 100:
                    style_name = str(getattr(paragraph.style, 'name', '') or '').lower()
                    centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                    if style_name.startswith(('heading', 'title', '标题')) or centered:
                        self._candidate(
                            candidates, 'title', text,
                            84 if centered else 78,
                            f'标题段落：{text}',
                        )
                self._line_candidates(candidates, text, 76)

            document_tables = document.tables
            if len(document_tables) > MAX_IMPORT_TABLES:
                warnings.append('合同表格过多，仅分析前 200 个表格')
            total_rows = 0
            for table_index, table in enumerate(
                document_tables[:MAX_IMPORT_TABLES], start=1
            ):
                table_rows = []
                for row_index, row in enumerate(table.rows, start=1):
                    cells = [
                        self._clean(cell.text, 2000)
                        for cell in row.cells[:MAX_IMPORT_TABLE_COLUMNS]
                    ]
                    total_rows += 1
                    if not any(cells):
                        continue
                    table_rows.append(' | '.join(value for value in cells if value))
                    payment_blocks.append({
                        'kind': 'table_row',
                        'table_index': table_index,
                        'row_index': row_index,
                        'cells': cells,
                        'text': ' | '.join(cells),
                    })
                    for cell_index, cell in enumerate(cells):
                        normalized = cell.strip(' ：:')
                        for field, labels in _FIELD_LABELS.items():
                            if normalized not in labels:
                                continue
                            next_value = next(
                                (value for value in cells[cell_index + 1:] if value), ''
                            )
                            self._candidate(
                                candidates, field, next_value, 96,
                                f'表格{table_index}第{row_index}行：'
                                + ' | '.join(value for value in cells if value),
                            )
                    self._line_candidates(
                        candidates,
                        ' | '.join(value for value in cells if value),
                        86,
                    )
                    if total_rows >= MAX_IMPORT_TABLE_ROWS:
                        warnings.append('合同表格行数过多，仅分析前 5000 行')
                        break
                tables.extend(table_rows)
                if total_rows >= MAX_IMPORT_TABLE_ROWS:
                    break
        finally:
            del document
        return paragraphs, tables, candidates, warnings, payment_blocks

    def _line_candidates(self, candidates, text, score):
        for field, labels in _FIELD_LABELS.items():
            for label in labels:
                pattern = (
                    re.escape(label)
                    + r'(?:\s*[（(][^）)]{0,20}[）)])?\s*[：:]\s*([^|；;。\n]{1,'
                    + str(_FIELD_LIMITS[field])
                    + r'})'
                )
                match = re.search(pattern, text)
                if match:
                    self._candidate(candidates, field, match.group(1), score, text)

    @staticmethod
    def _best(candidates, field, warnings):
        values = candidates.get(field) or []
        if not values:
            return None
        values.sort(key=lambda item: (-item['score'], len(str(item['value']))))
        distinct = []
        for item in values:
            if item['value'] not in [existing['value'] for existing in distinct]:
                distinct.append(item)
        if len(distinct) > 1:
            warnings.append(f'{field} 识别到多个候选值，请人工核对')
        return distinct[0]

    def preview_file(self, path, original_filename) -> ContractImportPreview:
        source = Path(path).resolve()
        if source.suffix.lower() != '.docx':
            raise ValueError('仅支持 .docx 格式的合同')
        if not source.is_file():
            raise ValueError('上传的合同文件不存在')
        if source.stat().st_size <= 0:
            raise ValueError('上传的合同文件为空')
        if source.stat().st_size > self.max_upload_bytes:
            raise ValueError('合同文件超过允许的上传大小')
        validate_office_archive(str(source))

        source_sha256 = self.sha256_file(source)
        duplicate = self.ledger_store.get_contract_by_source_sha256(source_sha256)
        if duplicate:
            error = ValueError('该合同文件已导入')
            error.contract_id = duplicate['id']
            raise error
        if self.isolate_preview:
            payload = run_isolated_worker(
                _contract_import_preview_worker,
                (str(source), original_filename, self.max_upload_bytes),
                timeout=CONTRACT_IMPORT_PARSE_TIMEOUT_SECONDS,
                label='外部合同解析',
                memory_limit_mb=CONTRACT_IMPORT_PARSE_MEMORY_MB,
            )
            return ContractImportPreview(**payload)

        (
            paragraphs, table_lines, candidates, warnings, payment_blocks
        ) = self._extract_document(source)
        diagnostics = []
        selected = {}
        for field in _FIELD_LABELS:
            candidate = self._best(candidates, field, warnings)
            if candidate:
                selected[field] = candidate['value']
                diagnostics.append(asdict(FieldDiagnostic(
                    field=field,
                    value=candidate['value'],
                    confidence=self._confidence(candidate['score']),
                    evidence=candidate['evidence'],
                )))

        filename_title = Path(original_filename or source.name).stem
        title = self._clean(selected.get('title') or filename_title, 200) or '未命名合同'
        amount = parse_number(selected.get('amount'))
        sign_date = normalize_date(selected.get('sign_date'))
        expiry_date = normalize_date(selected.get('expiry_date'))
        if selected.get('amount') and amount is None:
            warnings.append('合同金额候选值无法解析，请人工填写')
        if selected.get('sign_date') and not sign_date:
            warnings.append('签订日期候选值无法解析，请人工填写')
        if selected.get('expiry_date') and not expiry_date:
            warnings.append('到期日期候选值无法解析，请人工填写')

        summary = {
            'contract_no': self._clean(selected.get('contract_no'), 80),
            'title': title,
            'counterparty': self._clean(selected.get('counterparty'), 120),
            'amount': amount,
            'sign_date': sign_date,
            'expiry_date': expiry_date,
            'owner': '',
            'status': 'draft',
            'project_name': '',
            'subsystem_name': '',
            'coverage_mode': '',
            'coverage_not_applicable': False,
            'coverage_start': None,
            'coverage_end': None,
        }
        extraction = payment_extractor.extract_payment_items(
            payment_blocks,
            contract_amount=amount,
            sign_date=sign_date,
        )
        normalized_plans = []
        for plan in extraction.plans:
            normalized_plans.append({
                **plan,
                'confirm_status': 'pending',
                'payment_status': 'unpaid',
                'paid_amount': 0,
                'paid_date': '',
            })
        warnings.extend(extraction.warnings)
        document_text = '\n'.join(paragraphs + table_lines)
        if not document_text:
            warnings.append('合同中没有可识别的正文或表格，请人工填写台账信息')
        return ContractImportPreview(
            original_filename=os.path.basename(original_filename or source.name)[:255],
            source_sha256=source_sha256,
            summary=summary,
            diagnostics=diagnostics,
            plans=normalized_plans,
            rules=extraction.rules,
            warnings=warnings,
        )

    @staticmethod
    def _within(root: Path, candidate: Path) -> bool:
        try:
            candidate.relative_to(root)
            return True
        except ValueError:
            return False

    def finalize(self, request: ContractImportRequest) -> ContractImportResult:
        staging = Path(request.staging_path).resolve()
        if not self._within(self.uploads_dir, staging):
            raise ValueError('合同暂存路径无效')
        if not staging.is_file() or staging.suffix.lower() != '.docx':
            raise ValueError('待导入合同文件不存在')
        validate_office_archive(str(staging))
        actual_sha256 = self.sha256_file(staging)
        expected_sha256 = str(request.source_sha256 or '').strip().lower()
        if not re.fullmatch(r'[0-9a-f]{64}', expected_sha256):
            raise ValueError('合同文件摘要无效，请重新上传')
        if actual_sha256.lower() != expected_sha256:
            raise ValueError('合同文件在复核期间发生变化，请重新上传')
        duplicate = self.ledger_store.get_contract_by_source_sha256(actual_sha256)
        if duplicate:
            error = ValueError('该合同文件已导入')
            error.contract_id = duplicate['id']
            raise error

        title = self._clean(request.summary.get('title'), 200)
        if not title:
            raise ValueError('合同名称不能为空')
        status = str(request.summary.get('status') or 'draft').strip()
        if status not in self.ledger_store.CONTRACT_STATUSES:
            raise ValueError('合同状态无效')
        classification = parse_contract_classification(request.summary)
        if len(request.plans or []) > 30:
            raise ValueError('导入时付款计划不能超过 30 条')
        if len(request.rules or []) > 60:
            raise ValueError('导入时付款规则不能超过 60 条')

        self.output_dir.mkdir(parents=True, exist_ok=True)
        output_path = self.output_dir / f'imported_{uuid.uuid4().hex}.docx'
        job_id = uuid.uuid4().hex
        summary = {
            **request.summary,
            **classification,
            'title': title,
            'status': status,
            'record_origin': 'imported',
            'template_name': '',
            'original_filename': os.path.basename(request.original_filename or '')[:255],
            'source_sha256': actual_sha256,
        }
        file_moved = False
        committed = False
        self.ledger_store.create_generation_job(
            job_id, str(output_path), str(staging)
        )
        self.ledger_store.update_generation_job(job_id, 'staged')
        try:
            with self.ledger_store.get_conn() as conn:
                contract_id, plan_count = self.ledger_store.create_contract_with_plans(
                    summary, {}, str(output_path), request.plans, conn=conn,
                    rules=request.rules,
                )
                self.replace_file(str(staging), str(output_path))
                file_moved = True
                self.ledger_store.update_generation_job(
                    job_id,
                    'file_moved',
                    contract_id=contract_id,
                    conn=conn,
                )
            committed = True
            result = ContractImportResult(
                contract_id, str(output_path), plan_count, len(request.rules or [])
            )
        except Exception as exc:
            if file_moved and not committed and output_path.exists():
                try:
                    self.replace_file(str(output_path), str(staging))
                except OSError:
                    try:
                        output_path.unlink()
                    except OSError:
                        logging.getLogger('contract_tool').error(
                            'Failed to remove uncompensated import output: %s',
                            output_path,
                            exc_info=True,
                        )
            if isinstance(exc, ValueError) and not hasattr(exc, 'contract_id'):
                duplicate = self.ledger_store.get_contract_by_source_sha256(actual_sha256)
                if duplicate:
                    exc.contract_id = duplicate['id']
            if not committed:
                try:
                    self.ledger_store.update_generation_job(
                        job_id, 'failed', error=str(exc)
                    )
                except Exception:
                    logging.getLogger('contract_tool').error(
                        'Failed to record contract import failure for job %s',
                        job_id,
                        exc_info=True,
                    )
            raise

        # Once the transaction commits, the file and ledger row are durable.
        # A terminal-marker failure is reconciled safely during the next start.
        try:
            self.after_commit(result)
            self.ledger_store.update_generation_job(job_id, 'completed')
        except Exception:
            logging.getLogger('contract_tool').error(
                'Contract import committed but job finalization failed: %s',
                job_id,
                exc_info=True,
            )
        return result


class _NoDuplicateLedgerStore:
    @staticmethod
    def get_contract_by_source_sha256(_source_sha256):
        return None


def _contract_import_preview_worker(
    path,
    original_filename,
    max_upload_bytes,
    result_queue,
):
    source = Path(path)
    service = ContractImportService(
        ledger_store=_NoDuplicateLedgerStore(),
        uploads_dir=source.parent,
        output_dir=source.parent,
        max_upload_bytes=max_upload_bytes,
        isolate_preview=False,
    )
    preview = service.preview_file(source, original_filename)
    result_queue.put(('ok', preview.to_dict()))
