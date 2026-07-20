"""Non-standard Excel, Word, and PDF quotation mapping pipeline."""

from __future__ import annotations

import os
from datetime import date, datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP

from docx import Document
from openpyxl import load_workbook

import procurement_store
from services import procurement_file_service
from utils.logger import get_logger


MAPPING_FIELDS = [
    ('project_item_id', '项目明细ID'), ('line_no', '行号'), ('item_name', '物资名称'),
    ('spec_model', '规格型号'), ('drawing_no', '图号/代号'), ('quantity', '数量'),
    ('unit', '单位'), ('unit_price', '单价'), ('amount', '金额'),
    ('delivery_period', '分项交期'), ('technical_deviation', '技术偏离'),
    ('commercial_deviation', '商务偏离'), ('remark', '备注'),
]


def _json_value(value):
    if isinstance(value, (date, datetime)):
        return value.strftime('%Y-%m-%d')
    if value is None:
        return ''
    if isinstance(value, (int, float, str, bool)):
        return value
    return str(value)


def _extract_excel(path):
    workbook = load_workbook(path, data_only=False, read_only=True)
    try:
        tables = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [_json_value(value) for value in row]
                if any(value not in ('', None) for value in values):
                    rows.append(values)
                if len(rows) >= 5000:
                    break
            if rows:
                tables.append({'name': sheet.title, 'rows': rows})
        return tables, []
    finally:
        workbook.close()


def _extract_word(path):
    document = Document(path)
    tables = []
    for index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append({'name': f'表格{index}', 'rows': rows})
    return tables, []


def _extract_pdf(path):
    diagnostics = []
    try:
        import pdfplumber
    except ImportError:
        return [], ['缺少 pdfplumber，安装 requirements.txt 后可解析文本型 PDF']
    tables = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            for table_index, raw_table in enumerate(page.extract_tables() or [], start=1):
                rows = [[str(cell or '').strip() for cell in row] for row in raw_table if row]
                if rows:
                    tables.append({'name': f'第{page_number}页表格{table_index}', 'rows': rows})
    if tables:
        return tables, diagnostics
    try:
        import pytesseract
        import pypdfium2 as pdfium
        tesseract_cmd = os.environ.get('CT_TESSERACT_CMD', '').strip()
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        pdf = pdfium.PdfDocument(path)
        images = [page.render(scale=3).to_pil() for page in pdf]
        rows = []
        for page_number, image in enumerate(images, start=1):
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            page_rows = [[part.strip() for part in line.split('\t')]
                         for line in text.splitlines() if line.strip()]
            rows.extend(page_rows)
        if rows:
            tables.append({'name': 'OCR识别结果', 'rows': rows})
    except Exception as exc:
        diagnostics.append(f'未识别到 PDF 表格；扫描件 OCR 需要本机 Tesseract 与 Poppler：{exc}')
    return tables, diagnostics


def create_mapping_job(project_id, supplier_id, quote_round, file_storage):
    project = procurement_store.get_project(project_id)
    supplier = procurement_store.get_project_supplier(supplier_id)
    if not project or not supplier or supplier['project_id'] != project_id:
        raise ValueError('采购项目或候选供应商不存在')
    extension = os.path.splitext(file_storage.filename or '')[1].lower()
    extractors = {'.xlsx': ('excel', _extract_excel), '.docx': ('word', _extract_word), '.pdf': ('pdf', _extract_pdf)}
    if extension not in extractors:
        raise ValueError('非标准报价仅支持 .xlsx、.docx 或 .pdf')
    saved = procurement_file_service.save_upload(project, 'supplier_quote', file_storage)
    source_type, extractor = extractors[extension]
    try:
        tables, diagnostics = extractor(saved['absolute_path'])
        if not tables:
            raise ValueError('；'.join(diagnostics) or '文件中未识别到可映射的表格')
        return procurement_store.create_mapping_job({
            'project_id': project_id, 'supplier_id': supplier_id, 'quote_round': int(quote_round),
            'source_type': source_type, 'original_name': saved['original_name'],
            'relative_path': saved['relative_path'], 'file_sha256': saved['sha256'],
            'source': {'tables': tables, 'diagnostics': diagnostics, 'size_bytes': saved['size_bytes']},
        })
    except Exception:
        try:
            os.remove(saved['absolute_path'])
        except OSError:
            get_logger().warning(
                '非标准报价解析失败后无法删除上传文件: %s',
                saved['absolute_path'],
                exc_info=True,
            )
        raise


def _decimal(value, label, errors, row_number, required=True):
    raw = str(value or '').replace(',', '').strip()
    if not raw and not required:
        return None
    try:
        number = Decimal(raw)
    except InvalidOperation:
        errors.append(f'第 {row_number} 行{label}格式无效')
        return None
    if number < 0:
        errors.append(f'第 {row_number} 行{label}不能为负数')
        return None
    return number


def _minor(number):
    return int((number * 100).quantize(Decimal('1'), rounding=ROUND_HALF_UP))


def _cell(row, mapping, key):
    index = mapping.get(key)
    if index is None or index < 0 or index >= len(row):
        return ''
    return row[index]


def map_to_import_job(mapping_job_id, form):
    job = procurement_store.get_mapping_job(mapping_job_id)
    if not job:
        raise ValueError('字段映射任务不存在')
    table_name = str(form.get('table_name') or '')
    table = next((item for item in job['source']['tables'] if item['name'] == table_name), None)
    if not table:
        raise ValueError('请选择有效的工作表或表格')
    try:
        header_row = max(0, int(form.get('header_row', 1)) - 1)
    except ValueError as exc:
        raise ValueError('表头行号无效') from exc
    if header_row >= len(table['rows']):
        raise ValueError('表头行号超出文件范围')
    mapping = {}
    for key, _ in MAPPING_FIELDS:
        raw = str(form.get(f'map_{key}', '')).strip()
        mapping[key] = int(raw) if raw else None
    if mapping['item_name'] is None and mapping['project_item_id'] is None and mapping['line_no'] is None:
        raise ValueError('至少映射项目明细ID、行号或物资名称之一')
    if mapping['unit_price'] is None:
        raise ValueError('必须映射单价列')

    project_items = procurement_store.list_project_items(job['project_id'])
    by_id = {item['id']: item for item in project_items}
    by_line = {item['line_no']: item for item in project_items}
    by_name = {(item['item_name'].strip(), (item.get('spec_model') or '').strip()): item for item in project_items}
    parsed = []
    errors = []
    warnings = list(job['source'].get('diagnostics') or [])
    seen = set()
    for source_index, row in enumerate(table['rows'][header_row + 1:], start=header_row + 2):
        if not any(str(value or '').strip() for value in row):
            continue
        item = None
        raw_id = _cell(row, mapping, 'project_item_id')
        raw_line = _cell(row, mapping, 'line_no')
        if str(raw_id).strip():
            try:
                item = by_id.get(int(raw_id))
            except ValueError:
                get_logger().debug('报价映射中的项目明细 ID 无效: %r', raw_id)
        if item is None and str(raw_line).strip():
            try:
                item = by_line.get(int(raw_line))
            except ValueError:
                get_logger().debug('报价映射中的行号无效: %r', raw_line)
        if item is None:
            name = str(_cell(row, mapping, 'item_name') or '').strip()
            spec = str(_cell(row, mapping, 'spec_model') or '').strip()
            item = by_name.get((name, spec)) or next(
                (candidate for candidate in project_items if candidate['item_name'].strip() == name), None
            )
        if not item:
            warnings.append(f'第 {source_index} 行无法匹配项目明细，已忽略')
            continue
        if item['id'] in seen:
            errors.append(f'第 {source_index} 行重复匹配项目明细：{item["item_name"]}')
            continue
        seen.add(item['id'])
        quantity_raw = _cell(row, mapping, 'quantity') or item['quantity_text']
        quantity = _decimal(quantity_raw, '数量', errors, source_index)
        unit_price = _decimal(_cell(row, mapping, 'unit_price'), '单价', errors, source_index)
        if quantity is None or unit_price is None:
            continue
        amount = quantity * unit_price
        stated_amount = _decimal(_cell(row, mapping, 'amount'), '金额', errors, source_index, required=False)
        if stated_amount is not None and _minor(stated_amount) != _minor(amount):
            warnings.append(f'第 {source_index} 行金额与数量×单价不一致，已按系统重算')
        parsed.append({
            'project_item_id': item['id'], 'line_no': item['line_no'], 'item_name': item['item_name'],
            'spec_model': item.get('spec_model') or '', 'drawing_no': item.get('drawing_no') or '',
            'quantity_text': format(quantity.normalize(), 'f'),
            'unit': str(_cell(row, mapping, 'unit') or item['unit']).strip(),
            'unit_price_minor': _minor(unit_price), 'amount_minor': _minor(amount),
            'delivery_period': str(_cell(row, mapping, 'delivery_period') or '').strip(),
            'technical_deviation': str(_cell(row, mapping, 'technical_deviation') or '').strip(),
            'commercial_deviation': str(_cell(row, mapping, 'commercial_deviation') or '').strip(),
            'remark': str(_cell(row, mapping, 'remark') or '').strip(),
        })
    if not parsed:
        errors.append('映射后没有可导入的报价明细')
    missing = sorted(set(by_id) - seen)
    if missing:
        warnings.append(f'缺少 {len(missing)} 项项目明细')
    tax_raw = str(form.get('tax_rate') or '').strip()
    tax_bps = None
    if tax_raw:
        try:
            tax = Decimal(tax_raw)
            if tax < 0 or tax > 100:
                raise InvalidOperation
            tax_bps = int(tax * 100)
        except InvalidOperation:
            errors.append('税率格式无效')
    payload = {
        'header': {
            'quote_round': job['quote_round'], 'quote_date': str(form.get('quote_date') or ''),
            'quote_valid_until': str(form.get('quote_valid_until') or ''),
            'total_amount_minor': sum(item['amount_minor'] for item in parsed), 'currency': 'CNY',
            'tax_rate_bps': tax_bps, 'price_basis': str(form.get('price_basis') or 'tax_inclusive'),
            'delivery_period': str(form.get('delivery_period') or ''),
            'payment_terms': str(form.get('payment_terms') or ''),
            'warranty_period': str(form.get('warranty_period') or ''),
            'package_transport': str(form.get('package_transport') or ''),
            'technical_deviation': '', 'commercial_deviation': '',
        },
        'items': parsed, 'missing_project_item_ids': missing,
        'size_bytes': int(job['source'].get('size_bytes') or 0),
    }
    # 仅存储映射相关字段，排除 CSRF token 等敏感信息
    _SENSITIVE_FORM_KEYS = {'csrf_token'}
    metadata = {k: v for k, v in dict(form).items() if k not in _SENSITIVE_FORM_KEYS}
    procurement_store.update_mapping_job(mapping_job_id, mapping, metadata, 'invalid' if errors else 'parsed')
    import_job_id = procurement_store.create_import_job({
        'project_id': job['project_id'], 'supplier_id': job['supplier_id'],
        'quote_round': job['quote_round'], 'original_name': job['original_name'],
        'relative_path': job['relative_path'], 'file_sha256': job['file_sha256'],
        'parser_version': 'mapping-1.0', 'payload': payload, 'errors': errors, 'warnings': warnings,
    })
    return import_job_id
