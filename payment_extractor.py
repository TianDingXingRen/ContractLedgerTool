"""Deterministic payment-clause extraction.

The module deliberately separates contractual rules from actionable payment
instances.  Rules describe *how* a payment is calculated; plans are created
only for one-off clauses whose amount can be resolved from the contract.  A
recurring clause such as "每次投产通知金额的 30%" remains a rule until a
business event supplies the notification amount.

No probabilistic model is used.  Regular expressions perform tokenisation and
small state-machine style helpers bind tokens inside one payment node.
"""

from __future__ import annotations

import hashlib
import json
import math
import re
from dataclasses import dataclass
from datetime import datetime, timedelta
from decimal import Decimal
from docx import Document

from utils.field_utils import normalize_date


MAX_EXTRACTED_PLANS = 30
MAX_EXTRACTED_RULES = 60
EXTRACTOR_VERSION = 'payment-rules-v2'

PAYMENT_KEYWORDS = (
    '付款', '支付', '付清', '结算', '预付款', '预付', '定金', '尾款', '余款',
    '质保金', '保证金', '货款', '价款', '款项', '验收', '发票', '开票',
    '货到', '到货', '投产通知', '生产通知', '排产通知',
)
STRONG_PAYMENT_KEYWORDS = (
    '付款', '支付', '付清', '结算', '预付款', '定金', '尾款', '余款',
    '质保金', '保证金', '货款', '价款', '款项',
)
PAYMENT_ACTION_KEYWORDS = (
    '付款', '支付', '付清', '结算', '预付款', '预付', '定金', '尾款',
    '余款', '质保金', '保证金', '一次性总付', '分期支付', '扣留',
)
TRIGGER_ONLY_KEYWORDS = (
    '验收', '发票', '开票', '货到', '到货', '投产通知', '生产通知',
    '排产通知', '对账',
)
PLANISH_NO_MONEY_KEYWORDS = (
    '一次性总付', '分期支付', '预付款', '尾款', '余款', '质保金',
    '付款期限', '结算方式', '投产通知', '生产通知', '排产通知',
)
EXCLUDE_PAYMENT_KEYWORDS = (
    '不得', '报销', '赔偿', '违约金', '罚金', '滞纳金', '责任方',
)
PRICE_SUMMARY_KEYWORDS = (
    '合同总价款', '合同款项', '订购', '不含税金额', '税额', '增值税',
    '税率', '含税', '大写', '小写',
)

EVENT_LABELS = {
    'contract_signed': '合同签订',
    'effective': '合同生效',
    'invoice_received': '收到发票',
    'arrival': '到货',
    'shipment': '发货',
    'delivery': '交付',
    'acceptance': '验收合格',
    'warranty_end': '质保期满',
    'production_notice_issued': '投产通知下达',
    'settlement_confirmed': '结算确认',
    'fixed_date': '固定日期',
    'other': '其他',
}

EVENT_PATTERNS = (
    ('production_notice_issued', r'(?:收到|接到|下达|签收|确认)?(?:投产|生产|排产)通知'),
    ('contract_signed', r'(?:合同)?(?:签订|签署|盖章)'),
    ('effective', r'合同生效|生效'),
    ('invoice_received', r'收到[^，；。]{0,16}发票|发票(?:开具|送达|提交)|开票'),
    ('acceptance', r'(?:验收合格|通过验收|终验|初验)'),
    ('arrival', r'(?:货到|到货)'),
    ('shipment', r'发货'),
    ('delivery', r'交付'),
    ('warranty_end', r'(?:质保期|保修期)(?:届满|期满|满)'),
    ('settlement_confirmed', r'(?:结算|对账)(?:完成|确认|审核)'),
)

AMOUNT_BASIS_PATTERNS = (
    ('production_notice_total', r'(?:该|本|当)?(?:次|份)?(?:投产|生产|排产)通知(?:内|中|所列|项下)?(?:的)?(?:产品|货物)?总?价(?:款|金额)?'),
    ('batch_delivery_total', r'(?:该|本|当)?(?:批|批次)(?:产品|货物)?(?:的)?总?(?:价款|金额|货款)'),
    ('accepted_product_total', r'(?:已|实际)?验收(?:合格)?(?:产品|货物)?(?:的)?总?(?:价款|金额)'),
    ('settlement_amount', r'(?:当期|本期|该期)?(?:结算款|结算金额)'),
    ('invoice_amount', r'(?:发票|开票)(?:含税)?金额'),
    ('remaining_contract_amount', r'(?:剩余|余下|未付)(?:合同)?(?:价款|款项|金额)|尾款|余款'),
    ('contract_total_tax_exclusive', r'(?:不含税|税前)(?:合同)?(?:总价|价款|金额)'),
    ('contract_total_tax_inclusive', r'(?:含税)?(?:合同总价款|合同总价|合同金额|合同价款|总合同额)'),
)

REPEAT_KEYWORDS = ('每次', '每批', '各批', '各批次', '逐批', '每份通知', '每一批')
PAYMENT_SECTION_HEADING = re.compile(
    r'^(?:第[一二三四五六七八九十\d]+条\s*)?'
    r'(?:合同价款及支付|价款及支付|付款方式|支付方式|结算方式|付款条件|发票与付款|投产及付款)[:：]?$'
)
GENERIC_SECTION_HEADING = re.compile(
    r'^(?:第[一二三四五六七八九十\d]+条|[一二三四五六七八九十]+、)\s*[^，。；]{2,30}$'
)
ACTION_PATTERN = re.compile(r'一次性总付|分期支付|支付|付款|付清|预付|结算|扣留')
NUMBERED_BOUNDARY = re.compile(
    r'(?=(?:（\s*[一二三四五六七八九十\d]+\s*）|'
    r'\(\s*\d+\s*\)|[①②③④⑤⑥⑦⑧⑨⑩]|'
    r'第[一二三四五六七八九十\d]+(?:期|笔|阶段)))'
)


@dataclass(frozen=True)
class PaymentExtractionResult:
    plans: list[dict]
    rules: list[dict]
    warnings: list[str]


def extract_docx_blocks(path):
    """Return paragraphs and table rows without flattening their boundaries."""
    doc = Document(path)
    try:
        blocks = []
        for index, para in enumerate(doc.paragraphs, start=1):
            text = (para.text or '').strip()
            if text:
                blocks.append({'kind': 'paragraph', 'index': index, 'text': text})
        for table_index, table in enumerate(doc.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                cells = []
                for cell in row.cells:
                    value = ''.join(p.text or '' for p in cell.paragraphs).strip()
                    cells.append(value)
                if any(cells):
                    blocks.append({
                        'kind': 'table_row',
                        'table_index': table_index,
                        'row_index': row_index,
                        'cells': cells,
                        'text': ' | '.join(cells),
                    })
        return blocks
    finally:
        del doc


def extract_docx_text(path):
    return '\n'.join(block['text'] for block in extract_docx_blocks(path))


def extract_payment_plans(text, contract_amount=None, sign_date=''):
    """Compatibility API returning only actionable one-off plan drafts."""
    return extract_payment_items(text, contract_amount, sign_date).plans


def extract_payment_items(source, contract_amount=None, sign_date=''):
    """Extract versioned contractual rules and resolvable payment plans."""
    blocks = _mark_payment_context(_coerce_blocks(source))
    full_text = '\n'.join(block['text'] for block in blocks)
    selected_option = _selected_payment_option(full_text)
    rules = []
    plans = []
    warnings = []

    for block in blocks:
        snippets = _payment_snippets(
            block['text'], payment_context=bool(block.get('payment_context'))
        )
        for snippet_index, snippet in enumerate(snippets):
            if selected_option == 2 and any(k in snippet for k in ('一次性总付', '一次总付', '一次付清')):
                continue
            if selected_option == 1 and '分期支付' in snippet:
                continue
            group_key = _fingerprint(
                f"{block.get('kind')}|{block.get('index', block.get('row_index', 0))}|{snippet}"
            )
            for segment_index, segment in enumerate(_split_segments(snippet)):
                parsed = _parse_rule(
                    segment,
                    contract_amount=contract_amount,
                    sign_date=sign_date,
                    group_key=group_key,
                    ordinal=segment_index,
                    block=block,
                )
                if not parsed:
                    continue
                parsed_rules, draft_plans = parsed
                rules.extend(parsed_rules)
                plans.extend(draft_plans)

    rules = _propagate_group_context(_dedupe_rules(rules))
    rules = _validate_rule_groups(rules)[:MAX_EXTRACTED_RULES]
    rule_by_fp = {rule['rule_fingerprint']: rule for rule in rules}
    for plan in plans:
        rule = rule_by_fp.get(plan.get('rule_fingerprint'))
        if rule:
            _sync_plan_status(plan, rule)
    plans = [
        plan for plan in plans
        if (rule_by_fp.get(plan.get('rule_fingerprint')) or {}).get('repeat_mode') != 'each_event'
    ]
    plans = _trim_plans(_dedupe_plans(plans), contract_amount)[:MAX_EXTRACTED_PLANS]
    if len(rules) >= MAX_EXTRACTED_RULES:
        warnings.append('付款规则数量达到安全上限，请人工核对合同结构')
    return PaymentExtractionResult(plans=plans, rules=rules, warnings=warnings)


def _coerce_blocks(source):
    if isinstance(source, str):
        return [
            {'kind': 'paragraph', 'index': index, 'text': line.strip()}
            for index, line in enumerate(source.splitlines() or [source], start=1)
            if line.strip()
        ]
    blocks = []
    for index, item in enumerate(source or [], start=1):
        if isinstance(item, dict):
            text = str(item.get('text') or '').strip()
            if text:
                blocks.append({**item, 'text': text})
        else:
            text = str(item or '').strip()
            if text:
                blocks.append({'kind': 'paragraph', 'index': index, 'text': text})
    return blocks


def _selected_payment_option(text):
    match = re.search(r'采用以下第\s*([一二三四五六七八九十\d]+)\s*种方式', text)
    if not match:
        return None
    parsed = _parse_cn_number(match.group(1))
    return int(parsed) if parsed else None


def _payment_snippets(text, payment_context=False):
    text = re.sub(r'[\t\u3000]+', ' ', str(text or ''))
    rough = []
    for numbered in NUMBERED_BOUNDARY.split(text):
        rough.extend(re.split(r'[。；;\n\r]+', numbered))
    snippets = []
    for chunk in rough:
        chunk = re.sub(r'\s+', ' ', chunk).strip(' ，,|')
        if PAYMENT_SECTION_HEADING.fullmatch(chunk):
            continue
        if len(chunk) >= 2 and _is_candidate_clause(
            chunk, allow_bare_money=payment_context
        ):
            snippets.append(chunk)
    return snippets


def _action_matches(text):
    matches = []
    for match in ACTION_PATTERN.finditer(text):
        suffix = text[match.end():match.end() + 3]
        if suffix.startswith(('条件', '方式', '条款', '计划')):
            continue
        matches.append(match)
    return matches


def _split_segments(snippet):
    """Split around payment actions while keeping condition lists intact."""
    actions = _action_matches(snippet)
    if len(actions) >= 2:
        boundaries = []
        for current, following in zip(actions, actions[1:]):
            between = snippet[current.end():following.start()]
            delimiter = re.search(r'[，,]', between)
            if delimiter:
                boundaries.append(current.end() + delimiter.start())
        if boundaries:
            pieces = []
            start = 0
            for boundary in boundaries:
                pieces.append(snippet[start:boundary].strip(' ，,'))
                start = boundary + 1
            pieces.append(snippet[start:].strip(' ，,'))
            useful = [piece for piece in pieces if _is_candidate_clause(piece)]
            if len(useful) >= 2:
                return useful

    ratios = list(re.finditer(r'\d+(?:\.\d+)?\s*%', snippet))
    if len(ratios) >= 2:
        pieces = [part.strip() for part in re.split(r'[，,]', snippet) if part.strip()]
        ratio_pieces = [part for part in pieces if _extract_ratios(part)]
        if len(ratio_pieces) >= 2:
            prefix = next((part for part in pieces if '方式' in part and not _extract_ratios(part)), '')
            return [f'{prefix} {part}'.strip() if prefix else part for part in ratio_pieces]
    return [snippet]


def _is_candidate_clause(text, allow_bare_money=False):
    if any(k in text for k in EXCLUDE_PAYMENT_KEYWORDS):
        return False
    if _looks_like_price_summary(text):
        return False
    has_money = bool(_extract_ratios(text) or _extract_amounts(text))
    has_action = bool(_action_matches(text)) or any(k in text for k in PAYMENT_ACTION_KEYWORDS)
    has_strong = any(k in text for k in STRONG_PAYMENT_KEYWORDS)
    has_trigger = any(k in text for k in TRIGGER_ONLY_KEYWORDS)
    has_planish = any(k in text for k in PLANISH_NO_MONEY_KEYWORDS)
    has_dynamic_basis = _detect_amount_basis(text)[0] != 'unknown'
    if allow_bare_money and has_money:
        return True
    if has_money and (has_action or has_trigger or has_dynamic_basis):
        return True
    if has_planish or (has_action and (has_strong or has_trigger)):
        return True
    return False


def _mark_payment_context(blocks):
    """Mark blocks following a payment-section heading without flattening them."""
    marked = [dict(block) for block in blocks]
    remaining = 0
    for block in marked:
        text = str(block.get('text') or '').strip()
        if PAYMENT_SECTION_HEADING.fullmatch(text):
            block['payment_context'] = True
            remaining = 12
            continue
        if remaining > 0 and GENERIC_SECTION_HEADING.fullmatch(text):
            remaining = 0
        if remaining > 0:
            block['payment_context'] = True
            remaining -= 1
        elif _is_candidate_clause(text):
            block['payment_context'] = False
    return marked


def _looks_like_price_summary(text):
    if not any(k in text for k in PRICE_SUMMARY_KEYWORDS):
        return False
    if _action_matches(text):
        return False
    if any(k in text for k in ('每次', '每批', '各批次')):
        return False
    return bool(_extract_amounts(text) or _extract_ratios(text))


def _parse_rule(segment, contract_amount, sign_date, group_key, ordinal, block):
    ratios = _extract_ratios(segment)
    amounts = _extract_amounts(segment)
    explicit_date = _extract_date(segment, sign_date)
    conditions = _detect_conditions(segment, explicit_date)
    trigger_days = _extract_days(segment)
    basis, basis_text = _detect_amount_basis(segment)
    repeat_mode = 'each_event' if _is_recurring(segment, basis, conditions) else 'once'
    scope = _scope_for_rule(repeat_mode, basis, conditions)
    reasons = []

    if len(ratios) > 1:
        reasons.append('NODE_BOUNDARY_AMBIGUOUS')
    if not ratios and not amounts:
        reasons.append('AMOUNT_MISSING')
    if not conditions and not explicit_date:
        reasons.append('TRIGGER_MISSING')

    condition_logic = _condition_logic(segment, conditions)
    if len(conditions) > 1 and condition_logic == 'OTHER':
        reasons.append('CONDITION_LOGIC_AMBIGUOUS')

    if basis == 'unknown' and repeat_mode == 'each_event':
        reasons.append('AMOUNT_BASIS_MISSING')
    elif basis == 'unknown' and ratios and contract_amount is not None:
        basis = 'contract_total_tax_inclusive'
        basis_text = '合同金额（根据上下文推定）'
        reasons.append('AMOUNT_BASIS_INFERRED')

    ratio_values = ratios or [None]
    rule_plans = []
    rules_for_node = []
    for ratio_index, ratio in enumerate(ratio_values):
        explicit_amount = _paired_amount(amounts, ratio_index, len(ratio_values))
        calculated_amount = _calculated_amount(ratio, basis, contract_amount)
        conflict = _amounts_conflict(explicit_amount, calculated_amount)
        item_reasons = list(reasons)
        if conflict:
            item_reasons.append('EXPLICIT_AMOUNT_MISMATCH')
        parse_status = _parse_status(item_reasons)
        due_amount = None if conflict else (
            explicit_amount if explicit_amount is not None else calculated_amount
        )
        event_codes = [item['code'] for item in conditions]
        trigger_event = _condition_label(conditions, condition_logic, explicit_date)
        due_date = explicit_date
        legacy_event = event_codes[0] if event_codes else ('fixed_date' if explicit_date else 'other')
        if not due_date and sign_date and legacy_event in ('contract_signed', 'effective'):
            due_date = _add_days(sign_date, trigger_days or 0)
        fingerprint = _fingerprint(
            f'{group_key}|{ordinal}|{ratio_index}|{segment}|{ratio}|{explicit_amount}|{basis}|{repeat_mode}'
        )
        phase_name = _phase_name(segment, ratio_index)
        rule = {
            'group_key': group_key,
            'phase_name': phase_name,
            'rule_type': 'recurring' if repeat_mode == 'each_event' else 'conditional',
            'scope': scope,
            'trigger_event_type': legacy_event,
            'trigger_event': trigger_event,
            'trigger_days': trigger_days,
            'due_date': due_date,
            'conditions': event_codes,
            'conditions_json': json.dumps(event_codes, ensure_ascii=False),
            'condition_logic': condition_logic,
            'amount_basis': basis,
            'amount_basis_text': basis_text,
            'ratio': ratio,
            'explicit_amount': explicit_amount,
            'calculated_amount': calculated_amount,
            'repeat_mode': repeat_mode,
            'source_text': segment,
            'source_block': _block_label(block),
            'rule_fingerprint': fingerprint,
            'source_fingerprint': _fingerprint(segment),
            'extractor_version': EXTRACTOR_VERSION,
            'rule_version': 1,
            'parse_status': parse_status,
            'reason_codes': item_reasons,
            'reason_codes_json': json.dumps(item_reasons, ensure_ascii=False),
            'confirm_status': 'pending',
            'user_modified': 0,
        }
        rules_for_node.append(rule)
        if repeat_mode == 'once' and parse_status not in ('conflict', 'unsupported'):
            if ratio is not None or due_amount is not None:
                plan = _make_plan(
                    segment=segment,
                    phase_name=phase_name,
                    trigger_event=legacy_event,
                    trigger_days=trigger_days,
                    due_date=due_date,
                    ratio=ratio,
                    due_amount=due_amount,
                    confidence=_status_confidence(parse_status),
                )
                plan.update({
                    'trigger_event': trigger_event,
                    'amount_basis': basis,
                    'amount_basis_text': basis_text,
                    'explicit_amount': explicit_amount,
                    'calculated_amount': calculated_amount,
                    'conditions_json': rule['conditions_json'],
                    'condition_logic': condition_logic,
                    'repeat_mode': repeat_mode,
                    'parse_status': parse_status,
                    'reason_codes_json': rule['reason_codes_json'],
                    'rule_fingerprint': fingerprint,
                    'extractor_version': EXTRACTOR_VERSION,
                    'user_modified': 0,
                })
                rule_plans.append(plan)
    return rules_for_node, rule_plans


def _parse_segment(segment, contract_amount, sign_date):
    """Legacy private API retained for existing callers and tests."""
    parsed = _parse_rule(
        segment, contract_amount, sign_date,
        group_key=_fingerprint(segment), ordinal=0,
        block={'kind': 'paragraph', 'index': 1, 'text': segment},
    )
    return parsed[1] if parsed else []


def _make_plan(segment, phase_name, trigger_event, trigger_days, due_date,
               ratio, due_amount, confidence):
    payment_type = 'fixed_date' if due_date and trigger_event == 'fixed_date' else 'conditional'
    return {
        'phase_name': phase_name,
        'payment_type': payment_type,
        'trigger_event': EVENT_LABELS.get(trigger_event, '其他'),
        'trigger_days': trigger_days,
        'due_date': due_date,
        'ratio': ratio,
        'due_amount': due_amount,
        'paid_amount': 0,
        'condition_text': segment,
        'source_text': segment,
        'confidence': confidence,
        'confirm_status': 'pending',
        'payment_status': 'unpaid',
    }


def _extract_ratios(text):
    ratios = []
    for match in re.finditer(r'(\d+(?:\.\d+)?)\s*%', text):
        context = text[max(0, match.start() - 8):min(len(text), match.end() + 8)]
        if any(k in context for k in ('增值税', '税率', '税额')):
            continue
        ratios.append(float(match.group(1)))
    for match in re.finditer(r'百分之([零一二两三四五六七八九十百千万亿\d.]+)', text):
        parsed = _parse_cn_number(match.group(1))
        if parsed is not None:
            ratios.append(float(parsed))
    return ratios


def _extract_amounts(text):
    amounts = []
    pattern = r'(?:人民币|RMB|[￥¥])?\s*([0-9][0-9,]*(?:\.\d+)?)\s*(亿元|万元|千元|百元|元)'
    for match in re.finditer(pattern, text, re.I):
        prefix = text[max(0, match.start() - 10):match.start()]
        if any(kw in prefix for kw in ('已付', '已支付', '已预付', '实付', '扣减')):
            continue
        raw = match.group(1).replace(',', '')
        if len(raw.replace('.', '')) > 15:
            continue
        try:
            value = float(raw)
        except (ValueError, OverflowError):
            continue
        if not math.isfinite(value):
            continue
        unit = match.group(2)
        if unit == '亿元':
            value *= 100000000
        elif unit == '万元':
            value *= 10000
        elif unit == '千元':
            value *= 1000
        elif unit == '百元':
            value *= 100
        amounts.append(round(value, 2))
    return amounts


def _extract_days(text):
    # Date behaviour intentionally remains compatible with the original engine.
    match = re.search(r'(\d+)\s*(?:个)?(?:工作日|日|天)\s*内?', text)
    return int(match.group(1)) if match else None


def _extract_date(text, sign_date=''):
    result = normalize_date(text)
    if result:
        return result
    match = re.search(r'(\d{1,2})月(\d{1,2})日', text)
    if match:
        year = (sign_date or '')[:4] or str(datetime.now().year)
        month, day = match.groups()
        try:
            return datetime(int(year), int(month), int(day)).strftime('%Y-%m-%d')
        except ValueError:
            return ''
    return ''


def _detect_conditions(text, explicit_date=''):
    found = []
    for code, pattern in EVENT_PATTERNS:
        match = re.search(pattern, text)
        if match:
            found.append({'code': code, 'start': match.start(), 'end': match.end()})
    found.sort(key=lambda item: item['start'])
    if not found and explicit_date:
        found.append({'code': 'fixed_date', 'start': 0, 'end': 0})
    return found


def _condition_logic(text, conditions):
    if len(conditions) <= 1:
        return 'SINGLE'
    connectors = []
    for left, right in zip(conditions, conditions[1:]):
        connectors.append(text[left['end']:right['start']])
    joined = ''.join(connectors)
    if any(word in joined for word in ('或', '任一')):
        return 'OR'
    if any(word in joined for word in ('且', '并', '及', '同时', '和')):
        return 'AND'
    return 'OTHER'


def _condition_label(conditions, logic, explicit_date):
    if not conditions:
        return EVENT_LABELS['fixed_date'] if explicit_date else EVENT_LABELS['other']
    labels = [EVENT_LABELS.get(item['code'], item['code']) for item in conditions]
    connector = {'AND': '且', 'OR': '或'}.get(logic, '、')
    return connector.join(labels)


def _detect_event(text, explicit_date=''):
    conditions = _detect_conditions(text, explicit_date)
    return conditions[0]['code'] if conditions else 'other'


def _detect_amount_basis(text):
    for basis, pattern in AMOUNT_BASIS_PATTERNS:
        match = re.search(pattern, text)
        if match:
            return basis, match.group(0)
    return 'unknown', ''


def _is_recurring(text, basis, conditions):
    event_codes = {item['code'] for item in conditions}
    recurring_scope = basis in {'production_notice_total', 'batch_delivery_total'}
    recurring_event = 'production_notice_issued' in event_codes
    return any(keyword in text for keyword in REPEAT_KEYWORDS) and (
        recurring_scope or recurring_event
    )


def _scope_for_rule(repeat_mode, basis, conditions):
    event_codes = {item['code'] for item in conditions}
    if basis == 'production_notice_total' or 'production_notice_issued' in event_codes:
        return 'production_notice'
    if basis == 'batch_delivery_total':
        return 'delivery_batch'
    if basis == 'settlement_amount':
        return 'settlement_period'
    return 'contract' if repeat_mode == 'once' else 'other'


def _paired_amount(amounts, index, ratio_count):
    if not amounts:
        return None
    if ratio_count == 1:
        return amounts[-1]
    if len(amounts) == ratio_count:
        return amounts[index]
    return None


def _calculated_amount(ratio, basis, contract_amount):
    if ratio is None or contract_amount is None:
        return None
    if basis not in {'contract_total_tax_inclusive', 'contract_total_tax_exclusive'}:
        return None
    try:
        return round(float(contract_amount) * float(ratio) / 100, 2)
    except (TypeError, ValueError, OverflowError):
        return None


def _amounts_conflict(explicit_amount, calculated_amount):
    if explicit_amount is None or calculated_amount is None:
        return False
    tolerance = max(1.0, abs(calculated_amount) * 0.01)
    return abs(explicit_amount - calculated_amount) > tolerance


def _parse_status(reasons):
    if 'EXPLICIT_AMOUNT_MISMATCH' in reasons or 'RATIO_SUM_EXCEEDS_100' in reasons:
        return 'conflict'
    if 'AMOUNT_MISSING' in reasons:
        return 'unsupported'
    if reasons:
        return 'partial'
    return 'exact'


def _status_confidence(status):
    return {'exact': 'high', 'partial': 'medium'}.get(status, 'low')


def _phase_name(text, index):
    if '预付款' in text or '预付' in text or '定金' in text:
        return '预付款'
    if '投产通知' in text or '生产通知' in text or '排产通知' in text:
        return '投产通知款'
    if '到货' in text or '货到' in text:
        return '到货款'
    if '验收' in text:
        return '验收款'
    if '质保' in text:
        return '质保金'
    if '尾款' in text or '余款' in text:
        return '尾款'
    if '进度' in text:
        return '进度款'
    if index:
        return f'第{index + 1}期款'
    return '付款'


def _confidence(text, ratio, amount, due_date, trigger_event):
    has_money = ratio is not None or amount is not None
    has_condition = trigger_event not in ('other', '') or bool(due_date)
    if has_money and has_condition:
        return 'high'
    if has_money:
        return 'medium'
    return 'low'


def _safe_float(val, default=0.0):
    if val is None:
        return default
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def _trim_plans(plans, contract_amount):
    """Remove false positives without silently dropping monetary conflicts."""
    cleaned = []
    for plan in plans:
        source = plan.get('source_text') or ''
        if plan.get('due_amount') is None and plan.get('ratio') is None:
            continue
        if _looks_like_price_summary(source):
            continue
        cleaned.append(plan)
    return cleaned


def _validate_rule_groups(rules):
    groups = {}
    for rule in rules:
        key = (rule.get('group_key'), rule.get('scope'), rule.get('amount_basis'))
        groups.setdefault(key, []).append(rule)
    for group_rules in groups.values():
        ratios = [Decimal(str(rule['ratio'])) for rule in group_rules if rule.get('ratio') is not None]
        if len(ratios) < 2 or sum(ratios) <= Decimal('100.01'):
            continue
        for rule in group_rules:
            reasons = list(rule.get('reason_codes') or [])
            if 'RATIO_SUM_EXCEEDS_100' not in reasons:
                reasons.append('RATIO_SUM_EXCEEDS_100')
            rule['reason_codes'] = reasons
            rule['reason_codes_json'] = json.dumps(reasons, ensure_ascii=False)
            rule['parse_status'] = 'conflict'
    return rules


def _propagate_group_context(rules):
    """Carry one batch/notification scope to later phases in the same clause."""
    groups = {}
    for rule in rules:
        groups.setdefault(rule.get('group_key'), []).append(rule)
    for group_rules in groups.values():
        anchor = next((
            rule for rule in group_rules
            if rule.get('repeat_mode') == 'each_event'
            and rule.get('scope') in {'production_notice', 'delivery_batch'}
            and rule.get('amount_basis') != 'unknown'
        ), None)
        if not anchor:
            continue
        for rule in group_rules:
            if rule is anchor or rule.get('scope') != 'contract':
                continue
            source = rule.get('source_text') or ''
            if not any(word in source for word in (
                '该批', '本批', '当批', '到货', '验收', '质保', '尾款', '余款'
            )):
                continue
            rule['scope'] = anchor['scope']
            rule['repeat_mode'] = 'each_event'
            rule['rule_type'] = 'recurring'
            if rule.get('amount_basis') in ('unknown', 'contract_total_tax_inclusive'):
                rule['amount_basis'] = anchor['amount_basis']
                rule['amount_basis_text'] = f"沿用同组规则：{anchor.get('amount_basis_text') or anchor['amount_basis']}"
                rule['calculated_amount'] = None
            reasons = [
                code for code in (rule.get('reason_codes') or [])
                if code != 'AMOUNT_BASIS_INFERRED'
            ]
            rule['reason_codes'] = reasons
            rule['reason_codes_json'] = json.dumps(reasons, ensure_ascii=False)
            rule['parse_status'] = _parse_status(reasons)
    return rules


def _sync_plan_status(plan, rule):
    plan['parse_status'] = rule['parse_status']
    plan['reason_codes_json'] = rule['reason_codes_json']
    plan['confidence'] = _status_confidence(rule['parse_status'])
    if rule['parse_status'] in ('conflict', 'unsupported'):
        plan['due_amount'] = None


def _dedupe_rules(rules):
    result = []
    seen = set()
    for rule in rules:
        key = rule.get('rule_fingerprint')
        if key in seen:
            continue
        seen.add(key)
        result.append(rule)
    return result


def _dedupe_plans(plans):
    result = []
    seen = set()
    for plan in plans:
        key = plan.get('rule_fingerprint') or (
            plan.get('phase_name'), plan.get('ratio'), plan.get('due_amount'),
            plan.get('due_date'), plan.get('source_text'),
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(plan)
    return result


def _block_label(block):
    if block.get('kind') == 'table_row':
        return f"表格{block.get('table_index', '')}第{block.get('row_index', '')}行"
    return f"段落{block.get('index', '')}"


def _fingerprint(value):
    normalized = re.sub(r'\s+', '', str(value or '')).strip()
    return hashlib.sha256(normalized.encode('utf-8')).hexdigest()


def _add_days(date_text, days):
    try:
        base = datetime.strptime(date_text[:10], '%Y-%m-%d')
        return (base + timedelta(days=int(days or 0))).strftime('%Y-%m-%d')
    except (ValueError, TypeError):
        return ''


def _parse_cn_number(text):
    """Parse Chinese integer text, including section units 万 and 亿."""
    if re.fullmatch(r'\d+(?:\.\d+)?', text):
        return float(text)
    digit_map = {
        '零': 0, '一': 1, '二': 2, '两': 2, '三': 3, '四': 4,
        '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
    }
    unit_map = {'十': 10, '百': 100, '千': 1000}
    section_units = {'万': 10000, '亿': 100000000}
    result = 0
    section = 0
    digit = 0
    for char in text:
        if char in digit_map:
            digit = digit_map[char]
        elif char in unit_map:
            section += (digit or 1) * unit_map[char]
            digit = 0
        elif char in section_units:
            section = (section + digit) * section_units[char]
            result += section
            section = 0
            digit = 0
        else:
            return None
    result += section + digit
    return float(result) if result > 0 else None
