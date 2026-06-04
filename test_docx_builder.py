import unittest

from docx import Document

import docx_builder


class DocxBuilderFormattingTests(unittest.TestCase):
    def test_paragraph_placeholder_replacement_preserves_run_styles(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        before = paragraph.add_run('前')
        before.bold = True
        marker = paragraph.add_run('{字段}')
        marker.italic = True
        after = paragraph.add_run('后')
        after.underline = True

        docx_builder.apply_text_field(
            doc,
            {'type': 'paragraph', 'body_index': 0, 'placeholder': '{字段}'},
            '值',
        )

        self.assertEqual(paragraph.text, '前值后')
        self.assertEqual([run.text for run in paragraph.runs], ['前', '值', '后'])
        self.assertTrue(paragraph.runs[0].bold)
        self.assertTrue(paragraph.runs[1].italic)
        self.assertTrue(paragraph.runs[2].underline)

    def test_split_placeholder_replacement_preserves_surrounding_runs(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        before = paragraph.add_run('前')
        before.bold = True
        marker_start = paragraph.add_run('{字')
        marker_start.italic = True
        marker_end = paragraph.add_run('段}')
        marker_end.underline = True
        after = paragraph.add_run('后')
        after.bold = True

        docx_builder.apply_text_field(
            doc,
            {'type': 'paragraph', 'body_index': 0, 'placeholder': '{字段}'},
            '值',
        )

        self.assertEqual(paragraph.text, '前值后')
        self.assertEqual([run.text for run in paragraph.runs], ['前', '值', '', '后'])
        self.assertTrue(paragraph.runs[0].bold)
        self.assertTrue(paragraph.runs[1].italic)
        self.assertTrue(paragraph.runs[2].underline)
        self.assertTrue(paragraph.runs[3].bold)

    def test_table_cell_placeholder_replacement_preserves_run_styles(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        before = paragraph.add_run('前')
        before.bold = True
        marker = paragraph.add_run('{字段}')
        marker.italic = True
        after = paragraph.add_run('后')
        after.underline = True

        docx_builder.apply_text_field(
            doc,
            {
                'type': 'table_cell',
                'table_index': 0,
                'row_index': 0,
                'col_index': 0,
                'placeholder': '{字段}',
            },
            '值',
        )

        self.assertEqual(paragraph.text, '前值后')
        self.assertEqual([run.text for run in paragraph.runs], ['前', '值', '后'])
        self.assertTrue(paragraph.runs[0].bold)
        self.assertTrue(paragraph.runs[1].italic)
        self.assertTrue(paragraph.runs[2].underline)

    def test_table_field_marker_cell_preserves_surrounding_text(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = '产品'
        paragraph = table.cell(1, 0).paragraphs[0]
        label = paragraph.add_run('项:')
        label.bold = True
        marker = paragraph.add_run('{产品}')
        marker.italic = True

        docx_builder.apply_table_field(
            doc,
            {
                'location': {'table_index': 0, 'template_row_index': 1},
                'columns': [{'key': 'product', 'label': '产品'}],
            },
            [{'product': '产品A'}],
        )

        result_paragraph = table.cell(1, 0).paragraphs[0]
        self.assertEqual(result_paragraph.text, '项:产品A')
        self.assertEqual([run.text for run in result_paragraph.runs], ['项:', '产品A'])
        self.assertTrue(result_paragraph.runs[0].bold)
        self.assertTrue(result_paragraph.runs[1].italic)


if __name__ == '__main__':
    unittest.main()
