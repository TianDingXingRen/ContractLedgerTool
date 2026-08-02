import tempfile
import unittest
from pathlib import Path

from docx import Document

from utils.contract_preview import MAX_PARAGRAPHS, build_preview_blocks, build_preview_model


class ContractPreviewTests(unittest.TestCase):
    def test_build_preview_blocks_preserves_paragraphs_and_tables(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / 'template.docx'
            doc = Document()
            doc.add_paragraph('合同名称：{合同名称}')
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = '乙方'
            table.cell(0, 1).text = '{乙方单位名称}'
            table.cell(1, 0).text = '序号'
            table.cell(1, 1).text = '{物资名称}'
            doc.save(path)

            fields = [
                {
                    'id': 1,
                    'key': 'title',
                    'label': '合同名称',
                    'field_type': 'text',
                    'location': {'type': 'paragraph', 'body_index': 0, 'placeholder': '{合同名称}'},
                },
                {
                    'id': 2,
                    'key': 'counterparty',
                    'label': '乙方单位名称',
                    'field_type': 'text',
                    'location': {'type': 'table_cell', 'table_index': 0, 'row_index': 0, 'col_index': 1, 'placeholder': '{乙方单位名称}'},
                },
                {
                    'id': 3,
                    'key': 'items',
                    'label': '明细',
                    'field_type': 'table',
                    'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1},
                    'columns': [
                        {'key': '序号', 'label': '序号', 'field_type': 'text'},
                        {'key': '物资名称', 'label': '物资名称', 'field_type': 'text'},
                    ],
                },
            ]

            blocks = build_preview_blocks(str(path), fields)

        self.assertEqual(blocks[0]['type'], 'paragraph')
        self.assertEqual(blocks[0]['parts'][1]['field_id'], 1)
        self.assertIn('format', blocks[0])
        self.assertEqual(blocks[1]['type'], 'table')
        self.assertIn('grid', blocks[1])
        self.assertEqual(blocks[1]['rows'][0]['cells'][1]['parts'][0]['field_id'], 2)
        self.assertEqual(blocks[1]['rows'][0]['cells'][1]['col_span'], 1)
        self.assertEqual(blocks[1]['rows'][1]['repeat_field_id'], 3)
        self.assertEqual(blocks[1]['rows'][1]['cells'][1]['parts'][0]['kind'], 'table_column')

    def test_build_preview_model_reports_missing_source(self):
        model = build_preview_model('missing-source.docx', [])

        self.assertEqual(model['blocks'], [])
        self.assertTrue(model['warnings'])

    def test_build_preview_model_reports_truncation(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / 'long-template.docx'
            doc = Document()
            for index in range(MAX_PARAGRAPHS + 2):
                doc.add_paragraph(f'paragraph {index}')
            doc.save(path)

            model = build_preview_model(str(path), [])

        self.assertEqual(len(model['blocks']), MAX_PARAGRAPHS)
        self.assertTrue(model['warnings'])


if __name__ == '__main__':
    unittest.main()
