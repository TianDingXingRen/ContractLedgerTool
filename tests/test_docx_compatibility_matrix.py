"""Cross-application DOCX structure compatibility matrix.

The fixtures model OOXML structures commonly emitted by Microsoft Word and
WPS Office. Real desktop conversion is covered by scripts/office_compatibility_check.py.
"""

import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt

import docx_builder
from utils.security import validate_office_archive


def _paragraph_text(document):
    return '\n'.join(paragraph.text for paragraph in document.paragraphs)


def test_word_split_run_placeholder_round_trip(tmp_path):
    source = tmp_path / 'word-split-runs.docx'
    document = Document()
    paragraph = document.add_paragraph('合同编号：')
    paragraph.add_run('{合同')
    emphasized = paragraph.add_run('编号')
    emphasized.bold = True
    paragraph.add_run('}')
    document.save(source)

    validate_office_archive(source)
    reopened = Document(source)
    docx_builder.apply_text_field(
        reopened,
        {'type': 'paragraph', 'body_index': 0, 'placeholder': '{合同编号}'},
        'WORD-2026-001', '合同编号', 'contract_no',
    )
    output = tmp_path / 'word-split-runs-output.docx'
    reopened.save(output)

    final = Document(output)
    assert 'WORD-2026-001' in final.paragraphs[0].text
    assert '{合同编号}' not in final.paragraphs[0].text
    validate_office_archive(output)


def test_wps_style_cjk_fonts_headers_footers_and_sections_are_preserved(tmp_path):
    source = tmp_path / 'wps-layout.docx'
    document = Document()
    normal = document.styles['Normal']
    normal.font.name = 'SimSun'
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    document.sections[0].header.paragraphs[0].text = '采购合同—内部资料'
    document.sections[0].footer.paragraphs[0].text = '第 1 页'
    document.add_paragraph('对方单位：{对方单位}')
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.header.is_linked_to_previous = False
    second.header.paragraphs[0].text = '合同附件'
    document.add_paragraph('附件说明：中文、English、￥、①')
    document.save(source)

    reopened = Document(source)
    docx_builder.apply_text_field(
        reopened,
        {'type': 'paragraph', 'body_index': 0, 'placeholder': '{对方单位}'},
        '上海精工设备有限公司', '对方单位', 'counterparty',
    )
    output = tmp_path / 'wps-layout-output.docx'
    reopened.save(output)

    final = Document(output)
    assert len(final.sections) == 2
    assert final.sections[0].header.paragraphs[0].text == '采购合同—内部资料'
    assert final.sections[0].footer.paragraphs[0].text == '第 1 页'
    assert final.sections[1].header.paragraphs[0].text == '合同附件'
    assert '上海精工设备有限公司' in _paragraph_text(final)
    assert '中文、English、￥、①' in _paragraph_text(final)
    assert final.styles['Normal']._element.rPr.rFonts.get(qn('w:eastAsia')) == '宋体'


def test_merged_header_and_repeating_rows_round_trip(tmp_path):
    source = tmp_path / 'merged-table.docx'
    document = Document()
    table = document.add_table(rows=2, cols=3)
    merged = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    merged.text = '物资信息'
    table.rows[0].cells[2].text = '金额'
    table.rows[1].cells[0].text = '{name}'
    table.rows[1].cells[1].text = '{quantity}'
    table.rows[1].cells[2].text = '{amount}'
    document.save(source)

    reopened = Document(source)
    docx_builder.apply_table_field(reopened, {
        'key': 'items',
        'label': '采购明细',
        'field_type': 'table',
        'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1},
        'columns': [
            {'key': 'name', 'label': '品名'},
            {'key': 'quantity', 'label': '数量'},
            {'key': 'amount', 'label': '金额'},
        ],
    }, [
        {'name': '轴承', 'quantity': '10', 'amount': '1200.00'},
        {'name': '联轴器', 'quantity': '2', 'amount': '680.00'},
        {'name': '密封圈', 'quantity': '50', 'amount': '350.00'},
    ])
    output = tmp_path / 'merged-table-output.docx'
    reopened.save(output)

    final = Document(output)
    rows = final.tables[0].rows
    assert rows[0].cells[0].text == '物资信息'
    assert rows[0].cells[1].text == '物资信息'
    assert [[cell.text for cell in row.cells] for row in rows[1:]] == [
        ['轴承', '10', '1200.00'],
        ['联轴器', '2', '680.00'],
        ['密封圈', '50', '350.00'],
    ]


def test_generated_docx_contains_required_ooxml_parts(tmp_path):
    output = tmp_path / 'generated.docx'
    docx_builder.generate_from_scratch({
        'template_name': '兼容性基准合同',
        'fields': [{'key': 'party', 'label': '对方单位', 'field_type': 'text'}],
    }, {'party': '基准供应商'}, output)

    validate_office_archive(output)
    with zipfile.ZipFile(output) as archive:
        names = set(archive.namelist())
        assert '[Content_Types].xml' in names
        assert '_rels/.rels' in names
        assert 'word/document.xml' in names
        document_xml = archive.read('word/document.xml').decode('utf-8')
    assert '基准供应商' in document_xml
