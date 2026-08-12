import io
import json
import os
import shutil
import subprocess
import uuid
import zipfile
from html.parser import HTMLParser

import pytest
from docx import Document

import excel_bill_service
import field_eval
import ledger_store
import template_def
from utils import helpers
from utils.field_utils import parse_submitted_field_values


def _docx_text(blob):
    doc = Document(io.BytesIO(blob))
    parts = [paragraph.text or '' for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text or '' for cell in row.cells)
    return '\n'.join(parts)


class _InlineScriptCollector(HTMLParser):
    """Collect inline script bodies using HTML parsing, not tag regexes."""

    def __init__(self):
        super().__init__(convert_charrefs=False)
        self.scripts = []
        self._active_attrs = None
        self._active_body = []

    def handle_starttag(self, tag, attrs):
        if tag.lower() == 'script':
            self._active_attrs = dict(attrs)
            self._active_body = []

    def handle_data(self, data):
        if self._active_attrs is not None:
            self._active_body.append(data)

    def handle_endtag(self, tag):
        if tag.lower() != 'script' or self._active_attrs is None:
            return
        self.scripts.append((self._active_attrs, ''.join(self._active_body)))
        self._active_attrs = None
        self._active_body = []


def _assert_inline_scripts_have_valid_syntax(html):
    node = shutil.which('node')
    if not node:
        pytest.skip('Node.js is not installed')
    collector = _InlineScriptCollector()
    collector.feed(html)
    executable_scripts = (
        body for attrs, body in collector.scripts
        if body.strip() and (attrs.get('type') or '').lower() != 'application/json'
    )
    for script in executable_scripts:
        result = subprocess.run(
            [node, '--check', '-'],
            input=script,
            text=True,
            encoding='utf-8',
            capture_output=True,
            check=False,
        )
        assert result.returncode == 0, result.stderr or result.stdout


def test_number_required_and_select_are_validated():
    fields = [
        {
            'id': 0, 'key': 'amount', 'label': '金额', 'field_type': 'number',
            'required': True, 'decimal_places': 2, 'min_value': 0, 'max_value': 100,
        },
        {
            'id': 1, 'key': 'method', 'label': '方式', 'field_type': 'select',
            'required': True, 'options': ['月结', '季结'],
        },
    ]
    values, errors = parse_submitted_field_values(
        fields, {'field_0': '12.5', 'field_1': '月结'}
    )
    assert errors == []
    assert values == {'amount': '12.50', 'method': '月结'}

    _, errors = parse_submitted_field_values(
        fields, {'field_0': '101', 'field_1': '任意值'}
    )
    assert any('不能大于' in error for error in errors)
    assert any('选项无效' in error for error in errors)


def test_table_aggregate_formula_is_supported():
    fields = [
        {
            'id': 0, 'key': 'items', 'label': '明细', 'field_type': 'table',
            'columns': [{'key': 'subtotal', 'label': '小计', 'field_type': 'number'}],
        },
        {
            'id': 1, 'key': 'total', 'label': '合计', 'field_type': 'calculated',
            'formula': 'SUM(items, subtotal)', 'decimal_places': 2,
        },
    ]
    values = {'items': [{'subtotal': '10.5'}, {'subtotal': '4.5'}]}
    errors = helpers.recalculate_scalar_fields(fields, values)
    assert errors == []
    assert float(values['total']) == 15
    assert field_eval.get_calc_deps(fields[1]) == {'items', 'subtotal'}


def test_table_formula_errors_block_preflight_and_generation(app, client):
    columns = [
        {'key': 'qty', 'label': '数量', 'field_type': 'number'},
        {'key': 'divisor', 'label': '除数', 'field_type': 'number'},
        {
            'key': 'ratio', 'label': '比率', 'field_type': 'calculated',
            'formula': 'qty / divisor', 'decimal_places': 2,
        },
    ]
    fields = [{
        'id': 0, 'key': 'items', 'label': '明细', 'field_type': 'table',
        'required': True, 'columns': columns,
    }]

    values = {'items': [{'qty': '10', 'divisor': '2'}]}
    assert helpers.recalculate_table_fields(fields, values) == []
    assert values['items'][0]['ratio'] == '5.00'

    tpl = template_def.TemplateDef.create('表格公式阻断测试', '', fields)
    template_path = tpl.save()
    sid = uuid.uuid4().hex
    helpers.save_session_data(sid, {
        'template_name': tpl.name,
        'template_path': template_path,
        'template_filename': os.path.basename(template_path),
        'stored_name': '',
        'step': 'editor',
    })
    with client.session_transaction() as session:
        session['sid'] = sid
        session['_csrf_token'] = 'formula-token'

    form = {
        'csrf_token': 'formula-token',
        'coverage_mode': 'not_applicable',
        'field_0': json.dumps([{'qty': '10', 'divisor': '0'}]),
        'table_cols_0': json.dumps(columns, ensure_ascii=False),
    }
    preflight = client.post('/generate/preflight', data=form)
    assert preflight.status_code == 400
    assert '除数为零' in preflight.get_json()['blocking'][0]

    generated = client.post('/generate', data=form)
    assert generated.status_code == 400
    assert '公式计算失败' in generated.get_data(as_text=True)
    assert ledger_store.list_contracts(per_page=10)['total'] == 0


def test_table_formulas_recalculate_forward_dependencies_and_ignore_stale_values():
    columns = [
        {
            'key': 'total', 'label': '总计', 'field_type': 'calculated',
            'formula': 'subtotal * 2', 'decimal_places': 2,
        },
        {'key': 'price', 'label': '单价', 'field_type': 'number'},
        {
            'key': 'subtotal', 'label': '小计', 'field_type': 'calculated',
            'formula': 'price + 1', 'decimal_places': 2,
        },
    ]
    fields = [{
        'id': 0, 'key': 'items', 'label': '明细',
        'field_type': 'table', 'columns': columns,
    }]
    values = {'items': [{
        'price': '10',
        'subtotal': '999999',
        'total': '888888',
    }]}

    assert helpers.recalculate_table_fields(fields, values) == []
    assert values['items'][0]['subtotal'] == '11.00'
    assert values['items'][0]['total'] == '22.00'


def test_contract_number_unique_and_batch_history(tmp_db):
    summary = {'contract_no': 'HT-001', 'title': '测试合同'}
    first_id = ledger_store.create_contract(summary, {}, 'a.docx')
    with pytest.raises(ValueError, match='合同编号已存在'):
        ledger_store.create_contract(summary, {}, 'b.docx')

    assert ledger_store.batch_update_status([first_id], 'signed') == 1
    assert ledger_store.soft_delete_contract(first_id) == 1
    assert ledger_store.restore_contract(first_id) == 1
    history = ledger_store.get_contract_history(first_id)
    assert any(row['field'] == 'status' and row['new_value'] == 'signed' for row in history)
    assert len([row for row in history if row['field'] == 'deleted_at']) == 2


def test_contract_number_remains_reserved_while_in_trash(tmp_db):
    contract_id = ledger_store.create_contract(
        {'contract_no': 'TRASH-001', 'title': '回收站编号'}, {}, 'trash.docx'
    )
    ledger_store.soft_delete_contract(contract_id)

    assert ledger_store.contract_no_exists('TRASH-001')
    with pytest.raises(ValueError, match='合同编号已存在'):
        ledger_store.create_contract(
            {'contract_no': 'TRASH-001', 'title': '重复编号'}, {}, 'duplicate.docx'
        )

    assert ledger_store.permanently_delete_contract(contract_id) == 1
    assert not ledger_store.contract_no_exists('TRASH-001')


def test_payment_plan_changes_are_atomic_and_status_is_derived(tmp_db):
    contract_id = ledger_store.create_contract(
        {'contract_no': 'PAY-001', 'title': '付款测试'}, {}, 'pay.docx'
    )
    with pytest.raises(ValueError):
        ledger_store.save_payment_plan_changes(contract_id, [
            {'data': {'phase_name': '首款', 'due_amount': 100, 'paid_amount': 0}},
            {'data': {'phase_name': '尾款', 'due_amount': 100, 'paid_amount': 120}},
        ])
    assert ledger_store.list_payment_plans(contract_id=contract_id) == []

    ledger_store.save_payment_plan_changes(contract_id, [{
        'data': {
            'phase_name': '首款', 'due_amount': 100, 'paid_amount': 40,
            'paid_date': '2026-06-23',
        },
    }])
    plan = ledger_store.list_payment_plans(contract_id=contract_id)[0]
    assert plan['payment_status'] == 'partial'


def test_excel_bill_can_find_non_table3_contract_table(monkeypatch):
    monkeypatch.setattr(excel_bill_service.ledger_store, 'get_contract', lambda _id: {
        'id': 1,
        'template_name': '',
        'values_json': json.dumps({
            'purchase_items': [{'name': '设备', 'quantity': 2}],
            'notes': 'not a table',
        }, ensure_ascii=False),
    })
    detail = excel_bill_service.extract_contract_table(1)
    assert detail['table_key'] == 'purchase_items'
    assert detail['rows'][0]['name'] == '设备'
    assert {column['key'] for column in detail['columns']} == {'name', 'quantity'}


def test_preview_uses_current_values_and_batch_files_remain_downloadable(app, client):
    fields = [
        {
            'id': 0, 'key': 'party_b', 'label': '乙方', 'field_type': 'text',
            'required': True,
            'location': {'type': 'paragraph', 'body_index': 0, 'placeholder': '{乙方}'},
        },
        {
            'id': 1, 'key': 'contract_no', 'label': '合同编号', 'field_type': 'text',
            'required': True,
            'location': {'type': 'paragraph', 'body_index': 1, 'placeholder': '{合同编号}'},
        },
    ]
    tpl = template_def.TemplateDef.create('计划实现测试模板', '', fields)
    template_path = tpl.save()
    template_filename = os.path.basename(template_path)
    sid = uuid.uuid4().hex
    helpers.save_session_data(sid, {
        'template_name': tpl.name,
        'template_path': template_path,
        'template_filename': template_filename,
        'stored_name': '',
        'step': 'editor',
    })
    with client.session_transaction() as session:
        session['sid'] = sid
        session['_csrf_token'] = 'plan-token'

    preview = client.post(
        f'/template/{template_filename}/preview',
        data={
            'csrf_token': 'plan-token',
            'field_0': '当前预览乙方',
            'field_1': 'PREVIEW-001',
        },
    )
    assert preview.status_code == 200
    assert '当前预览乙方' in _docx_text(preview.get_data())
    preview.close()
    assert ledger_store.list_contracts(per_page=10)['total'] == 0

    batch = client.post('/generate-batch', data={
        'csrf_token': 'plan-token',
        'coverage_mode': 'not_applicable',
        'field_0': '',
        'field_1': 'BATCH-001',
        'batch_counterparties': '甲公司\n乙公司',
        'batch_field_key': 'party_b',
    })
    assert batch.status_code == 200, batch.get_data(as_text=True)
    with zipfile.ZipFile(io.BytesIO(batch.get_data())) as archive:
        assert len(archive.namelist()) == 2
        assert archive.namelist()[0].startswith('001_')
    batch.close()

    contracts = ledger_store.list_contracts(per_page=10)['rows']
    assert len(contracts) == 2
    assert {row['contract_no'] for row in contracts} == {'BATCH-001-001', 'BATCH-001-002'}
    for contract in contracts:
        assert os.path.isfile(contract['docx_path'])
        download = client.get(f'/contracts/{contract["id"]}/download')
        assert download.status_code == 200
        download.close()


def test_runtime_defaults_are_written_under_runtime_data(app):
    paths = app.extensions['runtime_paths']
    filename = excel_bill_service.save_header_default(
        'standard_pr', '持久化测试', {'dept': '测试部门'}
    )
    expected = paths.excel_bill_defaults_dir / filename
    assert expected.is_file()
    assert excel_bill_service.load_header_default(filename)['header_data']['dept'] == '测试部门'


def test_number_field_controls_render_in_template_builder_and_editor(app, client):
    create_page = client.get('/create-template')
    assert create_page.status_code == 200
    create_html = create_page.get_data(as_text=True)
    with open(os.path.join(app.static_folder, 'js', 'template-builder.js'), encoding='utf-8') as f:
        builder_script = f.read()
    assert '<option value="number">数字</option>' in builder_script
    assert 'field_number_min_' in builder_script
    assert 'field_number_max_' in builder_script
    assert 'field_number_decimal_' in builder_script
    _assert_inline_scripts_have_valid_syntax(create_html)

    tpl = template_def.TemplateDef.create('数字字段渲染测试', '', [{
        'id': 0,
        'key': 'amount',
        'label': '合同金额',
        'field_type': 'number',
        'required': True,
        'min_value': 0,
        'max_value': 1000,
        'decimal_places': 2,
    }])
    template_path = tpl.save()
    editor_page = client.get(f'/template/{os.path.basename(template_path)}')
    assert editor_page.status_code == 200
    editor_html = editor_page.get_data(as_text=True)
    assert 'name="field_0"' in editor_html
    assert 'type="number"' in editor_html
    assert 'min="0"' in editor_html
    assert 'max="1000"' in editor_html
    assert 'step="0.01"' in editor_html
    _assert_inline_scripts_have_valid_syntax(editor_html)
