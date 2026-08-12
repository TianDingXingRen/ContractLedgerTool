"""Deterministic end-to-end document generation tests."""

import json
from pathlib import Path

from docx import Document

import docx_builder
import field_eval
import template_def


def _create_fixture(tmp_path: Path):
    source = tmp_path / 'source.docx'
    doc = Document()
    doc.add_paragraph('合同编号：{合同编号}')
    doc.add_paragraph('对方单位：{对方单位}')
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = '品名'
    table.rows[0].cells[1].text = '数量'
    table.rows[0].cells[2].text = '单价'
    table.rows[1].cells[0].text = '{item_name}'
    table.rows[1].cells[1].text = '{quantity}'
    table.rows[1].cells[2].text = '{unit_price}'
    doc.save(source)

    fields = [
        {
            'id': 0,
            'key': 'contract_no',
            'label': '合同编号',
            'field_type': 'text',
            'required': True,
            'location': {'type': 'paragraph', 'body_index': 0, 'placeholder': '{合同编号}'},
        },
        {
            'id': 1,
            'key': 'counterparty',
            'label': '对方单位',
            'field_type': 'text',
            'required': True,
            'location': {'type': 'paragraph', 'body_index': 1, 'placeholder': '{对方单位}'},
        },
        {
            'id': 2,
            'key': 'items',
            'label': '采购明细',
            'field_type': 'table',
            'required': True,
            'columns': [
                {'key': 'item_name', 'label': '品名', 'field_type': 'text'},
                {'key': 'quantity', 'label': '数量', 'field_type': 'number'},
                {'key': 'unit_price', 'label': '单价', 'field_type': 'number'},
            ],
            'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1},
        },
    ]
    template_path = tmp_path / 'fixture.contract-template'
    template_path.write_text(json.dumps({
        'format_version': '1.0',
        'template_name': '集成测试模板',
        'source_docx': source.name,
        'fields': fields,
    }, ensure_ascii=False), encoding='utf-8')
    return source, template_path


def test_complete_document_generation_round_trip(tmp_path):
    source, template_path = _create_fixture(tmp_path)
    tpl = template_def.TemplateDef.load(template_path)
    tpl.validate()
    values = {
        'contract_no': 'INT-20260720-001',
        'counterparty': '集成测试供应商',
        'items': [
            {'item_name': '高强度螺栓', 'quantity': '100', 'unit_price': '8.50'},
            {'item_name': '防松螺母', 'quantity': '100', 'unit_price': '2.30'},
        ],
    }

    doc = Document(source)
    for field in field_eval.sort_fields_by_dependency(tpl.data['fields']):
        if field['field_type'] == 'table':
            docx_builder.apply_table_field(doc, field, values[field['key']])
        else:
            docx_builder.apply_text_field(
                doc,
                field['location'],
                values[field['key']],
                field['label'],
                field['key'],
            )

    output = tmp_path / 'generated.docx'
    doc.save(output)
    reopened = Document(output)
    paragraphs = '\n'.join(p.text for p in reopened.paragraphs)
    table_rows = [[cell.text for cell in row.cells] for row in reopened.tables[0].rows]

    assert 'INT-20260720-001' in paragraphs
    assert '集成测试供应商' in paragraphs
    assert table_rows[1] == ['高强度螺栓', '100', '8.50']
    assert table_rows[2] == ['防松螺母', '100', '2.30']
    assert '{' not in paragraphs
    assert all('{' not in cell for row in table_rows for cell in row)


def test_template_load_deterministically_migrates_legacy_field_ids(tmp_path):
    template_path = tmp_path / 'legacy.contract-template'
    template_path.write_text(json.dumps({
        'format_version': '1.0',
        'template_name': '旧模板',
        'source_docx': '',
        'fields': [
            {'id': '7'},
            {'id': 7},
            {'id': 'bad'},
            {},
            {'id': 2},
        ],
    }), encoding='utf-8')

    first = template_def.TemplateDef.load(template_path)
    second = template_def.TemplateDef.load(template_path)

    assert [field['id'] for field in first.data['fields']] == [7, 0, 1, 3, 2]
    assert first.data['fields'] == second.data['fields']

    migrated_path = tmp_path / 'migrated.contract-template'
    first.save(str(migrated_path))
    migrated = json.loads(migrated_path.read_text(encoding='utf-8'))
    assert [field['id'] for field in migrated['fields']] == [7, 0, 1, 3, 2]
