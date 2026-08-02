# -*- coding: utf-8 -*-
"""Route-level regression tests for the Flask contract tool."""

import io
import tempfile
import unittest
import uuid
import zipfile
from unittest import mock

from docx import Document

import app
import template_def
from utils.session_store import save_session_data


def _docx_text(blob):
    doc = Document(io.BytesIO(blob))
    parts = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text or '')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(''.join(p.text or '' for p in cell.paragraphs))
    return '\n'.join(parts)


class AppFlowTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.original_base_dir = app.BASE_DIR
        self.original_resource_dir = app.RESOURCE_DIR
        app.reset_runtime()
        self.test_app = app.create_app(
            runtime_base_dir=self.tmp.name,
            resource_dir=self.original_resource_dir,
            run_maintenance=False,
            testing=True,
        )
        self.paths = self.test_app.extensions['runtime_paths']

        # 创建测试模板（无 source_docx，使用 generate_from_scratch）
        self._create_test_template()

    def tearDown(self):
        app.reset_runtime()
        self.tmp.cleanup()
        app.configure_runtime_paths(
            self.original_base_dir,
            self.original_resource_dir,
        )

    def _create_test_template(self):
        """创建一个最小测试模板（无需源 DOCX 文件）。"""
        fields = [
            {
                'key': 'party_b',
                'label': '乙方',
                'field_type': 'text',
                'required': True,
                'location': {'type': 'paragraph', 'body_index': 0, 'placeholder': '{乙方}'},
            },
            {
                'key': 'contract_amount',
                'label': '合同金额',
                'field_type': 'text',
                'required': False,
                'location': {'type': 'paragraph', 'body_index': 1, 'placeholder': '{合同金额}'},
            },
        ]
        tpl = template_def.TemplateDef.create('测试模板', '', fields)
        self.test_template_path = tpl.save()
        self.test_tpl = tpl

    def test_batch_generation_from_scratch_template(self):
        """批量生成：使用无源文件的模板（generate_from_scratch 路径）。"""
        form = {
            'batch_counterparties': '测试甲公司\n测试乙公司',
        }
        for index, field in enumerate(self.test_tpl.data.get('fields', [])):
            fid = field.get('id', index)
            form[f'field_{fid}'] = f'测试值{fid}'

        sid = uuid.uuid4().hex
        save_session_data(sid, {
            'template_name': self.test_tpl.name,
            'template_path': self.test_template_path,
            'stored_name': '',
            'step': 'editor',
        }, self.paths)

        with self.test_app.test_client() as client:
            with client.session_transaction() as sess:
                sess['sid'] = sid
                sess['_csrf_token'] = 'test-token'
            form['csrf_token'] = 'test-token'
            response = client.post('/generate-batch', data=form)
            status_code = response.status_code
            response_body = response.get_data()
            response.close()

        self.assertEqual(status_code, 200, f'Expected 200, got {status_code}: {response_body[:300]}')
        with zipfile.ZipFile(io.BytesIO(response_body)) as zf:
            names = zf.namelist()
            self.assertEqual(len(names), 2)
            first_text = _docx_text(zf.read(names[0]))

        self.assertIn('测试甲公司', first_text)

    def test_batch_zip_failure_discards_created_contract(self):
        form = {'batch_counterparties': '测试甲公司'}
        for index, field in enumerate(self.test_tpl.data.get('fields', [])):
            fid = field.get('id', index)
            form[f'field_{fid}'] = f'测试值{fid}'

        sid = uuid.uuid4().hex
        save_session_data(sid, {
            'template_name': self.test_tpl.name,
            'template_path': self.test_template_path,
            'stored_name': '',
            'step': 'editor',
        }, self.paths)

        with self.test_app.test_client() as client, \
                mock.patch(
                    'services.contract_generation_service.create_ledger_record',
                    return_value=123,
                ), \
                mock.patch(
                    'services.contract_batch_generation_service.zipfile.ZipFile.write',
                    side_effect=OSError('disk full'),
                ), \
                mock.patch(
                    'services.contract_batch_generation_service._discard_generated_contract'
                ) as discard:
            with client.session_transaction() as sess:
                sess['sid'] = sid
                sess['_csrf_token'] = 'test-token'
            form['csrf_token'] = 'test-token'
            response = client.post('/generate-batch', data=form)

        self.assertEqual(response.status_code, 500)
        discard.assert_called_once()
        self.assertEqual(discard.call_args.args[0], 123)


if __name__ == '__main__':
    unittest.main()
