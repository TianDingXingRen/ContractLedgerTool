from docx import Document
from openpyxl import load_workbook

import docx_builder
import xlsx_exporter


def test_generated_docx_matches_semantic_snapshot(tmp_path):
    output = tmp_path / 'contract.docx'
    template = {
        'template_name': '采购合同',
        'fields': [
            {'key': 'party', 'label': '供应商', 'field_type': 'text'},
            {
                'key': 'items', 'label': '采购明细', 'field_type': 'table',
                'columns': [
                    {'key': 'name', 'label': '物资名称'},
                    {'key': 'quantity', 'label': '数量'},
                ],
            },
        ],
    }
    docx_builder.generate_from_scratch(
        template,
        {'party': '北辰精工有限公司', 'items': [{'name': '精密轴承', 'quantity': '47'}]},
        output,
    )

    document = Document(output)
    snapshot = {
        'paragraphs': [paragraph.text for paragraph in document.paragraphs if paragraph.text],
        'tables': [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in document.tables
        ],
    }

    assert snapshot == {
        'paragraphs': ['采购合同', '供应商：北辰精工有限公司'],
        'tables': [[['物资名称', '数量'], ['精密轴承', '47']]],
    }


def test_generated_xlsx_matches_semantic_snapshot(tmp_path):
    output = tmp_path / 'payments.xlsx'
    xlsx_exporter.export_payment_plans(output, [{
        'project_name': '西岭升级项目', 'subsystem_name': '姿控分系统',
        'serial_no': 12,
        'coverage_start': 8, 'coverage_end': 16,
        'contract_no': 'HT-2026-071', 'contract_title': '设备采购合同',
        'counterparty': '北辰精工有限公司', 'phase_name': '到货款',
        'due_date': '2026-08-19', 'due_amount': 47200.35, 'paid_amount': 12000.10,
        'condition_text': '到货验收后付款', 'owner': 'Shao', 'remark': '',
    }])

    workbook = load_workbook(output, data_only=False, read_only=True)
    try:
        sheet = workbook.active
        assert [sheet.cell(3, column).value for column in range(1, 17)] == [
            '序号', '项目名称', '所属分系统', '所属发次', '发次范围',
            '合同编号', '合同名称', '对方单位',
            '款项名称', '应付日期', '应付金额', '已付金额', '未付金额',
            '付款条件', '负责人', '备注',
        ]
        assert [sheet.cell(4, column).value for column in range(1, 17)] == [
            1, '西岭升级项目', '姿控分系统', '第 12 发', '第 8–16 发',
            'HT-2026-071', '设备采购合同',
            '北辰精工有限公司', '到货款', '2026-08-19', 47200.35, 12000.1,
            35200.25, '到货验收后付款', 'Shao', None,
        ]
    finally:
        workbook.close()


def test_xlsx_exports_render_explicit_not_applicable_coverage(tmp_path):
    payments_output = tmp_path / 'payments-na.xlsx'
    xlsx_exporter.export_payment_plans(payments_output, [{
        'project_name': '', 'subsystem_name': '',
        'coverage_not_applicable': True,
        'contract_no': 'HT-NA-026', 'contract_title': '通用服务合同',
        'counterparty': '北岳质量技术有限公司', 'phase_name': '验收款',
        'due_amount': 68000, 'paid_amount': 0,
    }])
    workbook = load_workbook(payments_output, data_only=False, read_only=True)
    try:
        assert workbook.active.cell(4, 4).value == '不适用'
        assert workbook.active.cell(4, 5).value == '不适用'
    finally:
        workbook.close()

    historical_output = tmp_path / 'payments-na-historical.xlsx'
    xlsx_exporter.export_payment_plans(historical_output, [{
        'coverage_not_applicable': True,
        'serial_no': 9,
        'contract_no': 'HT-NA-HISTORY',
        'contract_title': '历史关联服务合同',
    }])
    workbook = load_workbook(
        historical_output, data_only=False, read_only=True
    )
    try:
        assert workbook.active.cell(4, 4).value == '第 9 发（历史关联）'
        assert workbook.active.cell(4, 5).value == '不适用'
    finally:
        workbook.close()

    contracts_output = tmp_path / 'contracts-na.xlsx'
    xlsx_exporter.export_contracts(contracts_output, [{
        'coverage_not_applicable': True,
        'contract_no': 'HT-NA-026', 'title': '通用服务合同',
        'counterparty': '北岳质量技术有限公司', 'status': 'active',
    }])
    workbook = load_workbook(contracts_output, data_only=False, read_only=True)
    try:
        assert workbook.active.cell(4, 4).value == '不适用'
    finally:
        workbook.close()
