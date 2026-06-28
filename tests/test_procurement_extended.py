import io
import json
import zipfile

from docx import Document
from openpyxl import Workbook, load_workbook
from werkzeug.datastructures import FileStorage

import procurement_store
from services import (
    award_service, historical_price_service, negotiation_service,
    procurement_project_service, project_document_service,
    quote_mapping_service, quote_service,
)


def _base_project(with_items=True):
    project_id = procurement_project_service.create_project({
        'project_no': 'CG-EXT-0001', 'project_name': '扩展业务功能测试',
        'demand_department': '制造部', 'owner': '采购员乙',
        'delivery_requirement': '30天', 'payment_requirement': '验收后付款',
    })
    if with_items:
        for name, spec, quantity in [('结构件A', 'A-01', '10'), ('结构件B', 'B-02', '5')]:
            procurement_project_service.add_item(project_id, {
                'item_name': name, 'spec_model': spec, 'quantity': quantity, 'unit': '件',
            })
    suppliers = [
        procurement_project_service.add_supplier(project_id, {'supplier_name': name})
        for name in ('供应商甲', '供应商乙')
    ]
    return project_id, suppliers


def _standard_quote(project_id, supplier_id, prices, quote_round=1):
    path = quote_service.generate_quote_template(project_id, supplier_id)
    workbook = load_workbook(path)
    info = workbook['报价信息']
    info['B5'] = quote_round
    info['B6'] = '2026-06-23'
    info['B8'] = 13
    detail = workbook['报价明细']
    detail['H2'], detail['H3'] = prices
    info['B16'] = 10 * prices[0] + 5 * prices[1]
    stream = io.BytesIO()
    workbook.save(stream)
    stream.seek(0)
    job_id = quote_service.create_import_job(
        project_id, supplier_id, quote_round,
        FileStorage(stream=stream, filename=f'standard-{supplier_id}-{quote_round}.xlsx'),
    )
    return quote_service.confirm_import(job_id)


def test_project_items_bulk_paste_excel_import_and_export(app, client):
    project_id, _ = _base_project(with_items=False)
    assert client.get(f'/procurement/projects/{project_id}/items/bulk').status_code == 200
    procurement_project_service.add_items_from_paste(
        project_id,
        '结构件A\tA-01\tT-A\t10\t件\t2026-08-01\t按图加工\t首批\n'
        '结构件B\tB-02\tT-B\t5\t件\t2026-08-10\t无\t',
    )
    assert len(procurement_store.list_project_items(project_id)) == 2
    export_path = project_document_service.export_project_items(project_id)
    workbook = load_workbook(export_path, data_only=True)
    assert workbook['采购明细']['A2'].value == '结构件A'

    workbook2 = Workbook()
    sheet = workbook2.active
    sheet.append(['物资名称', '规格型号', '图号/代号', '数量', '单位', '要求交付日期'])
    sheet.append(['结构件C', 'C-03', 'T-C', 3, '套', '2026-09-01'])
    stream = io.BytesIO()
    workbook2.save(stream)
    stream.seek(0)
    procurement_project_service.add_items_from_excel(
        project_id, FileStorage(stream=stream, filename='items.xlsx')
    )
    assert len(procurement_store.list_project_items(project_id)) == 3


def test_nonstandard_excel_and_word_mapping_import(app, client):
    project_id, suppliers = _base_project()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = '供应商报价'
    sheet.append(['名称', '型号', '数量', '单位', '报价'])
    sheet.append(['结构件A', 'A-01', 10, '件', 88])
    sheet.append(['结构件B', 'B-02', 5, '件', 166])
    stream = io.BytesIO()
    workbook.save(stream)
    stream.seek(0)
    mapping_id = quote_mapping_service.create_mapping_job(
        project_id, suppliers[0], 1,
        FileStorage(stream=stream, filename='custom.xlsx'),
    )
    assert client.get(f'/procurement/quote-mappings/{mapping_id}').status_code == 200
    import_id = quote_mapping_service.map_to_import_job(mapping_id, {
        'table_name': '供应商报价', 'header_row': '1', 'map_item_name': '0',
        'map_spec_model': '1', 'map_quantity': '2', 'map_unit': '3',
        'map_unit_price': '4', 'tax_rate': '13', 'price_basis': 'tax_inclusive',
    })
    quote_id = quote_service.confirm_import(import_id)
    assert len(procurement_store.get_quote_items(quote_id)) == 2

    document = Document()
    table = document.add_table(rows=1, cols=5)
    for cell, value in zip(table.rows[0].cells, ['物资', '规格', '数量', '单位', '单价']):
        cell.text = value
    for values in [('结构件A', 'A-01', '10', '件', '90'), ('结构件B', 'B-02', '5', '件', '170')]:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    word_stream = io.BytesIO()
    document.save(word_stream)
    word_stream.seek(0)
    word_mapping = quote_mapping_service.create_mapping_job(
        project_id, suppliers[1], 1,
        FileStorage(stream=word_stream, filename='custom.docx'),
    )
    word_import = quote_mapping_service.map_to_import_job(word_mapping, {
        'table_name': '表格1', 'header_row': '1', 'map_item_name': '0',
        'map_spec_model': '1', 'map_quantity': '2', 'map_unit': '3',
        'map_unit_price': '4', 'tax_rate': '13', 'price_basis': 'tax_inclusive',
    })
    assert quote_service.confirm_import(word_import)


def test_negotiation_analysis_minutes_and_commitments(app, client):
    project_id, suppliers = _base_project()
    _standard_quote(project_id, suppliers[0], [100, 200], 1)
    _standard_quote(project_id, suppliers[0], [90, 180], 2)
    negotiation_service.save_round(project_id, {
        'round_no': '1', 'meeting_date': '2026-06-23', 'summary': '完成首轮谈判',
        f'commitment_{suppliers[0]}': '保证按期交付',
        f'delivery_{suppliers[0]}': '25天',
        f'payment_{suppliers[0]}': '验收后30天',
    })
    view = negotiation_service.negotiation_view(project_id)
    assert view['suppliers'][0]['reduction_minor'] > 0
    assert view['rounds'][0]['commitments'][0]['commitment'] == '保证按期交付'
    assert client.get(f'/procurement/projects/{project_id}/negotiation').status_code == 200
    minutes = project_document_service.generate_negotiation_minutes(project_id)
    assert '谈判纪要' in '\n'.join(p.text for p in Document(minutes).paragraphs)
    commitments = project_document_service.export_final_commitments(project_id)
    assert load_workbook(commitments)['最终承诺']['A2'].value == '供应商甲'


def test_split_award_history_erp_summary_and_archive(app, client):
    project_id, suppliers = _base_project()
    _standard_quote(project_id, suppliers[0], [80, 220], 1)
    _standard_quote(project_id, suppliers[1], [100, 180], 1)
    rows, _ = award_service.split_award_options(project_id)
    form = {'award_mode': 'split', 'reason_summary': '按分项最低有效价拆分成交'}
    selected_suppliers = []
    for row in rows:
        option = row['options'][0]
        form[f"selection_{row['item']['id']}"] = str(option['id'])
        selected_suppliers.append(option['supplier_name'])
    award_service.create_split_award(project_id, form)
    award = procurement_store.get_latest_award(project_id)
    assert award['is_split'] == 1
    assert set(selected_suppliers) == {'供应商甲', '供应商乙'}
    sheet = award_service.build_contract_data_sheet(project_id)
    payload = json.loads(sheet['payload_json'])
    assert set(payload['suppliers']) == {'供应商甲', '供应商乙'}

    history = historical_price_service.price_assistance('结构件')
    assert history['count'] == 2
    assert history['suggested_target_minor'] is not None
    assert client.get('/procurement/history-prices?q=结构件').status_code == 200
    summary_path = project_document_service.generate_erp_oa_summary(project_id)
    assert load_workbook(summary_path)['成交明细'].max_row == 3
    archive_path = project_document_service.generate_project_archive(project_id)
    with zipfile.ZipFile(archive_path) as archive:
        assert 'manifest.json' in archive.namelist()
        manifest = json.loads(archive.read('manifest.json').decode('utf-8'))
    assert manifest['project']['project_no'] == 'CG-EXT-0001'
