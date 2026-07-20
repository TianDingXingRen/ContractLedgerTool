# -*- coding: utf-8 -*-
"""Optional browser-side smoke tests.

These run when Playwright and its browser binaries are installed. The regular
unit suite skips them automatically on lightweight environments.
"""

import json
import os
import tempfile
import threading
import unittest

from docx import Document
from werkzeug.serving import make_server

import app as app_module
import ledger_store

try:
    from playwright.sync_api import sync_playwright
except Exception:  # pragma: no cover - optional dependency
    sync_playwright = None


class BrowserUiSmokeTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        if sync_playwright is None:
            raise unittest.SkipTest('Playwright is not installed')
        cls.runtime = tempfile.TemporaryDirectory()
        cls.test_app = app_module.create_app(
            runtime_base_dir=cls.runtime.name,
            resource_dir=app_module.RESOURCE_DIR,
            run_maintenance=False,
            testing=True,
        )
        cls._provision_contract_template()
        cls.server = make_server('127.0.0.1', 0, cls.test_app)
        cls.base_url = f'http://127.0.0.1:{cls.server.server_port}'
        cls.thread = threading.Thread(target=cls.server.serve_forever, daemon=True)
        cls.thread.start()

    @classmethod
    def _provision_contract_template(cls):
        paths = cls.test_app.extensions['runtime_paths']
        source_name = 'browser-e2e-source.docx'
        source_path = paths.uploads_dir / source_name
        doc = Document()
        doc.add_heading('浏览器端到端测试合同', level=1)
        doc.add_paragraph('合同编号：{合同编号}')
        doc.add_paragraph('对方单位：{对方单位}')
        doc.add_paragraph('数量：{数量}')
        doc.add_paragraph('单价：{单价}')
        doc.add_paragraph('合同金额：{合同金额}')
        doc.add_paragraph('签订日期：{签订日期}')
        doc.save(source_path)

        fields = []
        definitions = [
            (0, 'contract_no', '合同编号', 'text', 1, '{合同编号}'),
            (1, 'counterparty', '对方单位', 'text', 2, '{对方单位}'),
            (2, 'quantity', '数量', 'number', 3, '{数量}'),
            (3, 'unit_price', '单价', 'number', 4, '{单价}'),
            (5, 'sign_date', '签订日期', 'text', 6, '{签订日期}'),
        ]
        for field_id, key, label, field_type, body_index, placeholder in definitions:
            fields.append({
                'id': field_id,
                'key': key,
                'label': label,
                'field_type': field_type,
                'required': True,
                'location': {
                    'type': 'paragraph',
                    'body_index': body_index,
                    'placeholder': placeholder,
                },
            })
        fields.insert(4, {
            'id': 4,
            'key': 'amount',
            'label': '合同金额',
            'field_type': 'calculated',
            'required': True,
            'formula': 'quantity * unit_price',
            'decimal_places': 2,
            'depends_on': ['quantity', 'unit_price'],
            'location': {
                'type': 'paragraph',
                'body_index': 5,
                'placeholder': '{合同金额}',
            },
        })
        template_path = paths.templates_dir / 'browser-e2e.contract-template'
        template_path.write_text(json.dumps({
            'format_version': '1.0',
            'template_name': '浏览器端到端测试模板',
            'source_docx': source_name,
            'fields': fields,
        }, ensure_ascii=False, indent=2), encoding='utf-8')

    @classmethod
    def tearDownClass(cls):
        if hasattr(cls, 'server'):
            cls.server.shutdown()
        if hasattr(cls, 'thread'):
            cls.thread.join(timeout=5)
        app_module.reset_runtime()
        if hasattr(cls, 'runtime'):
            cls.runtime.cleanup()

    def test_operations_pages_render_in_browser(self):
        try:
            with sync_playwright() as playwright:
                browser = playwright.chromium.launch(headless=True)
                page = browser.new_page(viewport={'width': 1280, 'height': 800})
                page.goto(f'{self.base_url}/diagnostics', wait_until='networkidle')
                page.locator('[data-testid="copy-diagnostics"]').wait_for()
                page.locator('[data-testid="refresh-diagnostics"]').wait_for()

                page.goto(f'{self.base_url}/backups', wait_until='networkidle')
                page.locator('[data-testid="create-backup"]').wait_for()
                page.get_by_role(
                    'button', name='选择完整数据包 ZIP 文件'
                ).wait_for()

                page.set_viewport_size({'width': 1024, 'height': 800})
                page.goto(
                    f'{self.base_url}/template/browser-e2e.contract-template',
                    wait_until='networkidle',
                )
                input_box = page.locator('.editor-input-column').bounding_box()
                preview_box = page.locator('.editor-assist-panel').bounding_box()
                self.assertIsNotNone(input_box)
                self.assertIsNotNone(preview_box)
                self.assertLess(input_box['y'], preview_box['y'])
                browser.close()
        except Exception as exc:
            if 'Executable doesn' in str(exc) or 'playwright install' in str(exc):
                self.skipTest('Playwright browser binaries are not installed')
            raise

    def test_full_contract_generation_flow(self):
        """Select template, fill fields, download DOCX, and verify ledger UI."""
        try:
            with sync_playwright() as playwright:
                browser = playwright.chromium.launch(headless=True)
                page = browser.new_page(viewport={'width': 1440, 'height': 960})
                page.goto(f'{self.base_url}/templates', wait_until='networkidle')
                page.locator(
                    '[data-testid="template-open"]'
                    '[data-template-filename="browser-e2e.contract-template"]'
                ).click()
                page.wait_for_url('**/template/browser-e2e.contract-template')

                values = {
                    'contract_no': 'E2E-20260720-001',
                    'counterparty': '端到端测试供应商',
                    'quantity': '2',
                    'unit_price': '64000.25',
                    'sign_date': '2026-07-20',
                }
                for key, value in values.items():
                    page.locator(f'input[data-field-key="{key}"]').fill(value)

                self.assertEqual(page.locator('#calc_4').input_value(), '128000.50')
                self.assertEqual(page.locator('#calc_input_4').input_value(), '128000.50')

                with page.expect_download(timeout=30_000) as download_info:
                    page.locator('#generateBtn').click()
                download = download_info.value
                self.assertTrue(download.suggested_filename.endswith('.docx'))
                self.assertIsNone(download.failure())

                result = page.locator('[data-testid="generation-result"]')
                result.wait_for(state='visible')
                detail_url = page.locator('#resultDetailLink').get_attribute('href')
                self.assertTrue(detail_url)
                page.locator('#resultDetailLink').click()
                page.wait_for_url('**/contracts/*')
                self.assertIn('E2E-20260720-001', page.locator('body').inner_text())

                page.goto(f'{self.base_url}/contracts', wait_until='networkidle')
                ledger_text = page.locator('[data-testid="contract-list-view"]').inner_text()
                self.assertIn('E2E-20260720-001', ledger_text)
                self.assertIn('端到端测试供应商', ledger_text)

                contracts = ledger_store.list_contracts()['rows']
                self.assertEqual(len(contracts), 1)
                self.assertEqual(contracts[0]['contract_no'], 'E2E-20260720-001')
                contract = ledger_store.get_contract(contracts[0]['id'])
                self.assertTrue(os.path.isfile(contract['docx_path']))
                generated = Document(contract['docx_path'])
                generated_text = '\n'.join(
                    paragraph.text for paragraph in generated.paragraphs
                )
                self.assertIn('合同金额：128000.50', generated_text)
                browser.close()
        except Exception as exc:
            if 'Executable doesn' in str(exc) or 'playwright install' in str(exc):
                self.skipTest('Playwright browser binaries are not installed')
            raise


if __name__ == '__main__':
    unittest.main()
