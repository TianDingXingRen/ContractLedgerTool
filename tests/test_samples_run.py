"""使用4个样本合同的自动化功能测试"""
import io
import os
import sys
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document

PASS = FAIL = 0
_results = []

def _docx_text(blob):
    doc = Document(io.BytesIO(blob))
    parts = [p.text or '' for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(''.join(p.text or '' for p in c.paragraphs))
    return '\n'.join(parts)

def check(label, condition, detail=''):
    global PASS, FAIL
    if condition:
        PASS += 1; _results.append(f'  [OK] {label}')
    else:
        FAIL += 1; _results.append(f'  [FAIL] {label}: {detail}')

def section(title):
    _results.append(f'\n-- {title} --')

def summary():
    print('\n' + '=' * 60)
    for r in _results: print(r)
    print('=' * 60)
    print(f'  PASS={PASS}  FAIL={FAIL}  TOTAL={PASS+FAIL}')
    return FAIL == 0
