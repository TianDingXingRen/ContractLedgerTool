import io
from concurrent.futures import ThreadPoolExecutor
import json
import os
import uuid
import zipfile
from datetime import date

import pytest
from openpyxl import load_workbook
from docx import Document
from docx.oxml.ns import qn
from werkzeug.datastructures import FileStorage

import ledger_store
import procurement_store
import template_def
from services import (
    award_service, comparison_service, negotiation_service, procurement_project_service,
    procurement_file_service, project_document_service, quote_service,
)
from utils import helpers


def _non_empty_runs(document):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                yield run
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():
                            yield run


def _assert_docx_uses_fangsong(document):
    style_fonts = document.styles['Normal']._element.rPr.rFonts
    assert style_fonts.get(qn('w:eastAsia')) == '仿宋'
    for run in _non_empty_runs(document):
        r_pr = run._element.rPr
        assert r_pr is not None
        assert r_pr.rFonts.get(qn('w:eastAsia')) == '仿宋'


def _project_with_items_and_suppliers():
    project_id = procurement_project_service.create_project({
        'project_no': 'CG-TEST-0001',
        'project_name': '结构件加工竞争性谈判',
        'purchase_method': 'competitive_negotiation',
        'demand_department': '研发部',
        'owner': '采购员甲',
        'budget_amount': '10000',
        'target_price': '9000',
        'delivery_place': '北京',
        'delivery_requirement': '30天',
        'payment_requirement': '验收后30天付款',
    })
    procurement_project_service.add_item(project_id, {
        'item_name': '结构件A', 'spec_model': 'A-01', 'drawing_no': 'T-A',
        'quantity': '10', 'unit': '件', 'required_delivery_date': '2026-08-01',
    })
    procurement_project_service.add_item(project_id, {
        'item_name': '结构件B', 'spec_model': 'B-02', 'drawing_no': 'T-B',
        'quantity': '5', 'unit': '件', 'required_delivery_date': '2026-08-10',
    })
    suppliers = []
    for index, name in enumerate(('供应商A', '供应商B', '供应商C'), start=1):
        suppliers.append(procurement_project_service.add_supplier(project_id, {
            'supplier_name': name, 'contact_person': name[-1] + '联系人',
            'contact_phone': f'1380000000{index}',
            'email': f'supplier{index}@example.test',
            'direct_support_experience': name[-1] + '直接配套经验',
            'aerospace_support_experience': name[-1] + '航空航天配套经验',
            'qualifications': name[-1] + '资质',
            'remark': name[-1] + '其他信息',
        }))
    return project_id, suppliers


def _filled_quote_file(project_id, supplier_id, prices, quote_round=1, missing_last=False):
    path = quote_service.generate_quote_template(project_id, supplier_id)
    workbook = load_workbook(path)
    info = workbook['报价信息']
    info['B5'] = quote_round
    info['B6'] = '2026-06-23'
    info['B7'] = '2026-07-23'
    info['B10'] = '30天'
    info['B11'] = '验收后30天付款'
    detail = workbook['报价明细']
    detail['H2'] = prices[0]
    detail['H3'] = prices[1]
    if missing_last:
        detail.delete_rows(3, 1)
    total = 10 * prices[0] + (0 if missing_last else 5 * prices[1])
    info['B16'] = total
    stream = io.BytesIO()
    workbook.save(stream)
    stream.seek(0)
    return FileStorage(stream=stream, filename=f'quote-{supplier_id}.xlsx')


def _import_quote(project_id, supplier_id, prices, missing_last=False, quote_round=1):
    file_storage = _filled_quote_file(
        project_id, supplier_id, prices, quote_round=quote_round, missing_last=missing_last
    )
    job_id = quote_service.create_import_job(project_id, supplier_id, quote_round, file_storage)
    job = procurement_store.get_import_job(job_id)
    assert job['errors'] == []
    quote_id = quote_service.confirm_import(job_id)
    assert quote_id
    return quote_id, job


def test_multi_round_quote_uses_latest_confirmed_round(app):
    project_id, suppliers = _project_with_items_and_suppliers()
    first_quote, _ = _import_quote(project_id, suppliers[0], [100, 200], quote_round=1)
    second_quote, second_job = _import_quote(project_id, suppliers[0], [90, 180], quote_round=2)
    assert quote_service.confirm_import(second_job['id']) == second_quote
    latest = procurement_store.get_latest_quotes(project_id)
    assert len(latest) == 1
    assert latest[0]['id'] == second_quote
    assert latest[0]['quote_round'] == 2
    assert latest[0]['total_amount_minor'] < procurement_store.get_quote(first_quote)['total_amount_minor']


def test_procurement_schema_crud_and_constraints(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    project = procurement_project_service.project_detail(project_id)
    assert project['project_no'] == 'CG-TEST-0001'
    assert project['budget_minor'] == 1_000_000
    assert len(project['items']) == 2
    assert len(project['suppliers']) == 3
    assert client.get(f'/procurement/projects/{project_id}').status_code == 200
    assert client.get(
        f'/procurement/projects/{project_id}/items/{project["items"][0]["id"]}/edit'
    ).status_code == 200
    assert client.get(
        f'/procurement/projects/{project_id}/suppliers/{suppliers[0]}/edit'
    ).status_code == 200

    with pytest.raises(ValueError, match='项目编号已存在'):
        procurement_project_service.create_project({
            'project_no': 'CG-TEST-0001', 'project_name': '重复项目',
        })

    procurement_project_service.transition(project_id, 'documents_ready')
    assert procurement_store.get_project(project_id)['status'] == 'documents_ready'
    try:
        procurement_project_service.transition(project_id, 'contract_created')
        assert False, '非法状态跳转应被拒绝'
    except ValueError as exc:
        message = str(exc)
        assert '询价文件已准备' in message
        assert '合同已生成' in message
        assert 'documents_ready' not in message
        assert 'contract_created' not in message

    procurement_store.delete_project_supplier(project_id, suppliers[-1])
    assert len(procurement_store.list_project_suppliers(project_id)) == 2


def test_automatic_project_numbers_are_unique_under_concurrency(app):
    def create(index):
        project_id = procurement_project_service.create_project({
            'project_name': f'并发采购项目 {index}',
        })
        return procurement_store.get_project(project_id)['project_no']

    with ThreadPoolExecutor(max_workers=8) as executor:
        numbers = list(executor.map(create, range(12)))

    assert len(numbers) == len(set(numbers))


def test_inline_item_and_supplier_add_return_to_entry_sections(app, client):
    project_id = procurement_project_service.create_project({
        'project_no': 'CG-TEST-ANCHOR',
        'project_name': '连续录入跳转测试',
    })
    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'inline-token'

    item_response = client.post(
        f'/procurement/projects/{project_id}/items',
        data={
            'csrf_token': 'inline-token',
            'item_name': '测试物资',
            'quantity': '2',
            'unit': '件',
        },
        follow_redirects=False,
    )
    assert item_response.status_code == 302
    assert item_response.headers['Location'].endswith(f'/procurement/projects/{project_id}#items')

    supplier_response = client.post(
        f'/procurement/projects/{project_id}/suppliers',
        data={
            'csrf_token': 'inline-token',
            'supplier_name': '连续录入供应商',
            'direct_support_experience': '  直接配套项目甲  ',
            'aerospace_support_experience': '  有，曾配套航天结构件  ',
            'qualifications': '  质量体系认证  ',
            'remark': '  其他补充信息  ',
        },
        follow_redirects=False,
    )
    assert supplier_response.status_code == 302
    assert supplier_response.headers['Location'].endswith(f'/procurement/projects/{project_id}#suppliers')
    supplier = procurement_store.list_project_suppliers(project_id)[0]
    assert supplier['direct_support_experience'] == '直接配套项目甲'
    assert supplier['aerospace_support_experience'] == '有，曾配套航天结构件'
    assert supplier['qualifications'] == '质量体系认证'
    assert supplier['remark'] == '其他补充信息'

    detail_html = client.get(
        f'/procurement/projects/{project_id}'
    ).get_data(as_text=True)
    assert all(label in detail_html for label in (
        '直接配套经验', '是否有航空航天配套经验', '资质', '其他',
    ))
    assert '直接配套项目甲' in detail_html

    edit_html = client.get(
        f'/procurement/projects/{project_id}/suppliers/{supplier["id"]}/edit'
    ).get_data(as_text=True)
    assert 'name="direct_support_experience"' in edit_html
    assert 'name="aerospace_support_experience"' in edit_html
    assert 'name="qualifications"' in edit_html
    assert '<span class="label-text">其他</span>' in edit_html
    assert '<span class="label-text">备注</span>' not in edit_html

    update_response = client.post(
        f'/procurement/projects/{project_id}/suppliers/{supplier["id"]}/edit',
        data={
            'csrf_token': 'inline-token',
            'supplier_name': '连续录入供应商',
            'direct_support_experience': '  更新后的直接配套经验  ',
            'aerospace_support_experience': '  更新后的航空航天经验  ',
            'qualifications': '  更新后的资质  ',
            'remark': '  更新后的其他信息  ',
        },
        follow_redirects=False,
    )
    assert update_response.status_code == 302
    updated_supplier = procurement_store.get_project_supplier(supplier['id'])
    assert updated_supplier['direct_support_experience'] == '更新后的直接配套经验'
    assert updated_supplier['aerospace_support_experience'] == '更新后的航空航天经验'
    assert updated_supplier['qualifications'] == '更新后的资质'
    assert updated_supplier['remark'] == '更新后的其他信息'


def test_standard_quote_import_page_downloads_supplier_template(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    page = client.get(f'/procurement/projects/{project_id}/quotes/import')
    assert page.status_code == 200
    html = page.get_data(as_text=True)
    assert '下载标准报价模板' in html
    assert f'/procurement/projects/{project_id}/quote-template' in html

    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'quote-template-token'
    response = client.post(
        f'/procurement/projects/{project_id}/quote-template',
        data={
            'csrf_token': 'quote-template-token',
            'supplier_id': str(suppliers[0]),
        },
    )
    assert response.status_code == 200
    assert response.headers['Content-Type'].startswith(
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response.close()


def test_confirmed_quote_can_be_edited_and_deleted_from_project(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    quote_id, job = _import_quote(project_id, suppliers[0], [100, 200])
    quote = procurement_store.get_quote(quote_id)
    quote_items = procurement_store.get_quote_items(quote_id)
    original_file = procurement_store.get_project_file(quote['original_file_id'])
    original_path = procurement_file_service.absolute_path(original_file['relative_path'])
    assert original_path.is_file()
    comparison_service.run_comparison(project_id, 20)

    detail_html = client.get(
        f'/procurement/projects/{project_id}'
    ).get_data(as_text=True)
    assert f'/procurement/projects/{project_id}/quotes/{quote_id}/edit' in detail_html
    assert f'/procurement/projects/{project_id}/quotes/{quote_id}/delete' in detail_html

    edit_page = client.get(
        f'/procurement/projects/{project_id}/quotes/{quote_id}/edit'
    )
    assert edit_page.status_code == 200
    assert '编辑供应商报价' in edit_page.get_data(as_text=True)
    assert client.get(
        f'/procurement/projects/{project_id + 999}/quotes/{quote_id}/edit'
    ).status_code == 404

    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'quote-edit-token'
    form = {
        'csrf_token': 'quote-edit-token',
        'quote_date': '2026-06-24',
        'quote_valid_until': '2026-08-24',
        'tax_rate': '9.5',
        'price_basis': 'tax_exclusive',
        'delivery_period': '20天',
        'payment_terms': '验收后60天付款',
        'warranty_period': '两年',
        'package_transport': '含包装运输',
        'technical_deviation': '整体无偏离',
        'commercial_deviation': '',
    }
    for index, item in enumerate(quote_items):
        form[f'unit_price_{item["id"]}'] = str(110 + index * 100)
        form[f'delivery_period_{item["id"]}'] = f'{15 + index}天'
        form[f'technical_deviation_{item["id"]}'] = ''
        form[f'commercial_deviation_{item["id"]}'] = ''
        form[f'remark_{item["id"]}'] = f'复核行{index + 1}'
    update_response = client.post(
        f'/procurement/projects/{project_id}/quotes/{quote_id}/edit',
        data=form,
        follow_redirects=False,
    )
    assert update_response.status_code == 302
    assert update_response.headers['Location'].endswith(
        f'/procurement/projects/{project_id}#quotes'
    )
    updated = procurement_store.get_quote(quote_id)
    assert updated['total_amount_minor'] == 215_000
    assert updated['tax_rate_bps'] == 950
    assert updated['price_basis'] == 'tax_exclusive'
    assert procurement_store.get_quote_items(quote_id)[0]['remark'] == '复核行1'
    assert procurement_store.get_latest_comparison(project_id) is None
    assert original_path.is_file()

    invalid_form = dict(form)
    invalid_form['tax_rate'] = 'NaN'
    invalid_response = client.post(
        f'/procurement/projects/{project_id}/quotes/{quote_id}/edit',
        data=invalid_form,
    )
    assert invalid_response.status_code == 400
    assert '税率必须是 0 到 100 之间的有限数值' in invalid_response.get_data(as_text=True)
    assert procurement_store.get_quote(quote_id)['tax_rate_bps'] == 950

    rejected_delete = client.post(
        f'/procurement/projects/{project_id}/quotes/{quote_id}/delete',
        data={},
    )
    assert rejected_delete.status_code == 400
    assert procurement_store.get_quote(quote_id)

    delete_response = client.post(
        f'/procurement/projects/{project_id}/quotes/{quote_id}/delete',
        data={'csrf_token': 'quote-edit-token'},
        follow_redirects=False,
    )
    assert delete_response.status_code == 302
    assert delete_response.headers['Location'].endswith(
        f'/procurement/projects/{project_id}#quotes'
    )
    assert procurement_store.get_quote(quote_id) is None
    assert procurement_store.get_import_job(job['id'])['status'] == 'cancelled'
    assert not original_path.exists()


def test_pdf_quote_attachment_uploads_and_rejects_non_pdf(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'pdf-token'

    response = client.post(
        f'/procurement/projects/{project_id}/quotes/pdf',
        data={
            'csrf_token': 'pdf-token',
            'supplier_id': str(suppliers[0]),
            'quote_round': '2',
            'file': (io.BytesIO(b'%PDF-1.4\n%%EOF'), 'quote.pdf'),
        },
        content_type='multipart/form-data',
        follow_redirects=False,
    )
    assert response.status_code == 302
    files = [
        row for row in procurement_store.list_project_files(project_id)
        if row['file_type'] == 'supplier_quote_pdf'
    ]
    assert files
    assert files[0]['original_name'] == '第2轮_供应商A_quote.pdf'
    assert procurement_file_service.absolute_path(files[0]['relative_path']).is_file()

    rejected = client.post(
        f'/procurement/projects/{project_id}/quotes/pdf',
        data={
            'csrf_token': 'pdf-token',
            'supplier_id': str(suppliers[0]),
            'quote_round': '1',
            'file': (io.BytesIO(b'not pdf'), 'quote.txt'),
        },
        content_type='multipart/form-data',
    )
    assert rejected.status_code == 400
    assert 'PDF 报价单仅支持 .pdf 格式' in rejected.get_data(as_text=True)

    fake_pdf = client.post(
        f'/procurement/projects/{project_id}/quotes/pdf',
        data={
            'csrf_token': 'pdf-token',
            'supplier_id': str(suppliers[0]),
            'quote_round': '1',
            'file': (io.BytesIO(b'not actually a pdf'), 'quote.pdf'),
        },
        content_type='multipart/form-data',
    )
    assert fake_pdf.status_code == 400
    files_after_reject = [
        row for row in procurement_store.list_project_files(project_id)
        if row['file_type'] == 'supplier_quote_pdf'
    ]
    assert len(files_after_reject) == 1


def test_quote_import_comparison_clarification_and_award(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    _import_quote(project_id, suppliers[0], [100, 200])
    _, missing_job = _import_quote(project_id, suppliers[1], [130, 210], missing_last=True)
    _import_quote(project_id, suppliers[2], [70, 190])
    assert missing_job['warnings']
    preview = client.get(f'/procurement/quote-imports/{missing_job["id"]}')
    assert preview.status_code == 200
    preview_html = preview.get_data(as_text=True)
    assert '已导入' in preview_html
    assert '>confirmed<' not in preview_html

    run_id = comparison_service.run_comparison(project_id, 20)
    assert run_id
    comparison = procurement_store.get_latest_comparison(project_id)
    result_types = {row['result_type'] for row in comparison['results']}
    assert 'missing_item' in result_types
    assert 'high_price' in result_types
    assert 'low_price' in result_types
    assert client.get(f'/procurement/projects/{project_id}/comparison').status_code == 200

    created = comparison_service.generate_clarifications(project_id)
    assert created > 0
    questions = procurement_store.list_clarifications(project_id)
    assert any(row['question_type'] == 'missing_item' for row in questions)
    clarification_path = project_document_service.generate_clarification_letter(project_id)
    assert '报价澄清问题清单' in '\n'.join(p.text for p in Document(clarification_path).paragraphs)

    try:
        award_service.create_award(project_id, suppliers[0], {
            'reason_summary': '未说明为何不选最低价',
        })
        assert False, '非最低有效价必须填写原因'
    except ValueError as exc:
        assert '最低总价' in str(exc)

    recommendation_id = award_service.create_award(project_id, suppliers[2], {
        'reason_summary': '价格合理且技术响应完整',
        'price_reason': '总价最低',
    })
    assert recommendation_id
    award = procurement_store.get_latest_award(project_id)
    assert award['supplier_id'] == suppliers[2]
    assert client.get(f'/procurement/projects/{project_id}/award').status_code == 200
    award_path = project_document_service.generate_award_recommendation(project_id)
    assert '成交建议' in '\n'.join(p.text for p in Document(award_path).paragraphs)
    detail_html = client.get(f'/procurement/projects/{project_id}').get_data(as_text=True)
    assert 'ERP/OA 摘要' not in detail_html
    sheet = award_service.build_contract_data_sheet(project_id)
    payload = json.loads(sheet['payload_json'])
    assert payload['supplier']['name'] == '供应商C'
    assert len(payload['items']) == 2
    file_types = {row['file_type'] for row in procurement_store.list_project_files(project_id)}
    assert {'quote_template', 'supplier_quote', 'clarification', 'award'} <= file_types


def test_inquiry_document_contains_project_items(app):
    project_id, _ = _project_with_items_and_suppliers()
    path = project_document_service.generate_inquiry_letter(project_id)
    document = Document(path)
    all_text = '\n'.join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert '询价函' in all_text
    assert '结构件A' in all_text
    assert '结构件B' in all_text
    _assert_docx_uses_fangsong(document)
    assert procurement_store.get_project(project_id)['status'] == 'documents_ready'


def test_negotiation_plan_prefills_generates_word_and_archives(app, client):
    project_id, _ = _project_with_items_and_suppliers()
    page = client.get(f'/procurement/projects/{project_id}/negotiation/plan')
    assert page.status_code == 200
    html = page.get_data(as_text=True)
    assert '谈判预案' in html
    assert '结构件A' in html
    assert '目标价格' in html
    assert '生成文件名' in html
    assert '备选供应商信息' in html
    assert all(name in html for name in ('供应商A', '供应商B', '供应商C'))
    assert all(label in html for label in (
        '直接配套经验', '是否有航空航天配套经验', '资质', '其他',
    ))
    assert all(value in html for value in (
        'A直接配套经验', 'A航空航天配套经验', 'A资质', 'A其他信息',
    ))
    assert 'A联系人' not in html
    assert 'supplier1@example.test' not in html

    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'plan-token'
    response = client.post(
        f'/procurement/projects/{project_id}/negotiation/plan',
        data={
            'csrf_token': 'plan-token',
            'filename': '自定义谈判预案.docx',
            'project_background': '按附件模板生成，减少重复填写。',
            'fixed_asset': '否',
        },
        follow_redirects=False,
    )
    assert response.status_code == 200
    assert response.headers['Content-Type'].startswith(
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response.close()

    files = [
        row for row in procurement_store.list_project_files(project_id)
        if row['file_type'] == 'negotiation_plan'
    ]
    assert files
    assert files[-1]['original_name'] == '自定义谈判预案.docx'
    path = procurement_file_service.absolute_path(files[-1]['relative_path'])
    document = Document(path)
    all_text = '\n'.join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert '结构件加工竞争性谈判谈判预案' in all_text
    assert '按附件模板生成，减少重复填写。' in all_text
    assert '结构件A' in all_text
    assert all(name in all_text for name in ('供应商A', '供应商B', '供应商C'))
    assert '三、备选供应商信息' in all_text
    assert '四、生产周期要求' in all_text
    assert '五、目标价格' in all_text
    assert '六、报价轮次' in all_text
    assert '七、评价方案' in all_text
    supplier_table = next(
        table for table in document.tables
        if table.rows[0].cells[1].text == '备选供应商名称'
    )
    assert [cell.text for cell in supplier_table.rows[0].cells] == [
        '序号', '备选供应商名称', '直接配套经验',
        '是否有航空航天配套经验', '资质', '其他',
    ]
    assert [row.cells[1].text for row in supplier_table.rows[1:]] == [
        '供应商A', '供应商B', '供应商C',
    ]
    assert [cell.text for cell in supplier_table.rows[1].cells] == [
        '1', '供应商A', 'A直接配套经验', 'A航空航天配套经验',
        'A资质', 'A其他信息',
    ]
    supplier_table_text = '\n'.join(
        cell.text for row in supplier_table.rows for cell in row.cells
    )
    assert 'A联系人' not in supplier_table_text
    assert 'supplier1@example.test' not in supplier_table_text
    _assert_docx_uses_fangsong(document)


@pytest.mark.parametrize('supplier_names', [
    [],
    ['单一备选供应商'],
    ['供应商甲', '供应商乙', '超长供应商名称' * 18],
])
def test_negotiation_plan_supplier_section_handles_cardinality_and_wrapping(
    app, supplier_names,
):
    project_id = procurement_project_service.create_project({
        'project_no': 'CG-SUPPLIER-DOCX',
        'project_name': '供应商章节测试',
        'purchase_method': 'competitive_negotiation',
    })
    procurement_project_service.add_item(project_id, {
        'item_name': '测试产品', 'quantity': '1', 'unit': '件',
    })
    for name in supplier_names:
        procurement_project_service.add_supplier(project_id, {
            'supplier_name': name,
        })

    path = project_document_service.generate_negotiation_plan(project_id)
    document = Document(path)
    supplier_table = next(
        table for table in document.tables
        if table.rows[0].cells[1].text == '备选供应商名称'
    )
    exported_names = [row.cells[1].text for row in supplier_table.rows[1:]]
    if supplier_names:
        assert exported_names == supplier_names
        assert all(
            cell.text == '—'
            for row in supplier_table.rows[1:]
            for cell in row.cells[2:]
        )
    else:
        assert exported_names == ['暂无备选供应商信息']
    assert [cell.text for cell in supplier_table.rows[0].cells] == [
        '序号', '备选供应商名称', '直接配套经验',
        '是否有航空航天配套经验', '资质', '其他',
    ]
    assert len(supplier_table.columns) == 6
    assert supplier_table.autofit is False
    _assert_docx_uses_fangsong(document)


def test_procurement_to_contract_prefills_editor_and_links_ledger(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    _import_quote(project_id, suppliers[0], [100, 200])
    award_service.create_award(project_id, suppliers[0], {
        'reason_summary': '综合评审通过',
    })
    assert client.get(f'/procurement/projects/{project_id}/to-contract').status_code == 200

    fields = [
        {'id': 0, 'key': 'project_name', 'label': '项目名称', 'field_type': 'text',
         'required': True, 'location': {'type': 'paragraph', 'body_index': 0}},
        {'id': 1, 'key': 'supplier', 'label': '供应商名称', 'field_type': 'text',
         'required': True, 'location': {'type': 'paragraph', 'body_index': 1}},
        {'id': 2, 'key': 'amount', 'label': '合同金额', 'field_type': 'number',
         'required': True, 'decimal_places': 2, 'location': {'type': 'paragraph', 'body_index': 2}},
    ]
    tpl = template_def.TemplateDef.create('采购转合同测试模板', '', fields)
    template_path = tpl.save()
    data = award_service.prepare_editor_session(
        project_id,
        os.path.basename(template_path),
        app.extensions['runtime_paths'],
    )
    assert data['fields'][0]['default_value'] == '结构件加工竞争性谈判'
    assert data['fields'][1]['default_value'] == '供应商A'
    assert data['project_name'] == '结构件加工竞争性谈判'

    sid = uuid.uuid4().hex
    helpers.save_session_data(sid, data)
    with client.session_transaction() as flask_session:
        flask_session['sid'] = sid
        flask_session['_csrf_token'] = 'procurement-token'
    editor_page = client.get('/editor')
    assert editor_page.status_code == 200
    editor_html = editor_page.get_data(as_text=True)
    assert 'value="供应商A"' in editor_html
    assert 'value="结构件加工竞争性谈判"' in editor_html
    assert 'id="batchToggle"' not in editor_html

    before_count = ledger_store.list_contracts(per_page=100)['total']
    preflight = client.post('/generate/preflight', data={
        'csrf_token': 'procurement-token',
        '_generation_mode': 'batch',
    })
    assert preflight.status_code == 400
    assert '仅支持单份生成' in preflight.get_json()['blocking'][0]
    batch = client.post('/generate-batch', data={'csrf_token': 'procurement-token'})
    assert batch.status_code == 400
    assert '仅支持单份生成' in batch.get_data(as_text=True)
    assert ledger_store.list_contracts(per_page=100)['total'] == before_count

    response = client.post('/generate', data={
        'csrf_token': 'procurement-token',
        'project_name': '结构件加工竞争性谈判',
        'field_0': '结构件加工竞争性谈判',
        'field_1': '供应商A',
        'field_2': '2000.00',
    })
    assert response.status_code == 200, response.get_data(as_text=True)
    contract_id = int(response.headers['X-Contract-Id'])
    response.close()
    links = procurement_store.get_project_contract_links(project_id)
    assert links[0]['contract_id'] == contract_id
    assert procurement_store.get_project(project_id)['status'] == 'contract_created'
    assert ledger_store.get_contract(contract_id)['counterparty'] == '供应商A'


def test_workflow_jump_records_skipped_stages(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'jump-token'
    response = client.post(
        f'/procurement/projects/{project_id}/workflow/jump',
        data={
            'csrf_token': 'jump-token',
            'target_stage': 'negotiation',
            'note': '项目已线下询价，直接补录谈判',
        },
        follow_redirects=False,
    )
    assert response.status_code == 302
    assert response.headers['Location'].endswith(f'/procurement/projects/{project_id}/negotiation')
    assert procurement_store.get_project(project_id)['status'] == 'negotiating'
    events = procurement_store.list_project_audit_events(project_id, actions=['workflow_jump'])
    assert events
    assert events[0]['after']['target_stage'] == 'negotiation'
    assert 'quotes' in events[0]['after']['skipped_stages']


def test_workflow_enter_does_not_require_skip_note_or_change_status(app, client):
    project_id, _suppliers = _project_with_items_and_suppliers()
    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'enter-token'

    response = client.post(
        f'/procurement/projects/{project_id}/workflow/jump',
        data={
            'csrf_token': 'enter-token',
            'target_stage': 'award',
            'mode': 'enter',
        },
        follow_redirects=False,
    )
    assert response.status_code == 302
    assert response.headers['Location'].endswith(f'/procurement/projects/{project_id}/award')
    assert procurement_store.get_project(project_id)['status'] == 'draft'

    award_page = client.get(f'/procurement/projects/{project_id}/award')
    award_html = award_page.get_data(as_text=True)
    assert '尚无已确认的结构化报价' in award_html
    assert '导入报价' in award_html

    try:
        procurement_project_service.jump_to_stage(project_id, 'award', '')
    except ValueError as exc:
        assert '跳过前置环节时需要填写原因' in str(exc)
    else:
        raise AssertionError('跳过前置环节必须填写原因')


def test_single_source_skips_comparison_as_required_stage(app, client):
    project_id = procurement_project_service.create_project({
        'project_no': 'CG-SINGLE-0001',
        'project_name': '单一来源项目',
        'purchase_method': 'single_source',
    })
    procurement_project_service.add_item(project_id, {
        'item_name': '专用件',
        'quantity': '1',
        'unit': '件',
    })
    procurement_project_service.add_supplier(project_id, {'supplier_name': '唯一供应商'})

    workflow = procurement_project_service.build_workflow_view(project_id)
    comparison = next(stage for stage in workflow['stages'] if stage['key'] == 'comparison')
    negotiation = next(stage for stage in workflow['stages'] if stage['key'] == 'negotiation')
    assert comparison['status'] == 'not_applicable'
    assert '比价与澄清' not in negotiation['missing_labels']

    page_html = client.get('/procurement/projects/new').get_data(as_text=True)
    assert '单一来源' in page_html
    list_html = client.get('/procurement/projects').get_data(as_text=True)
    assert '单一来源' in list_html
    assert 'single_source' not in list_html


def test_direct_contract_session_and_generated_contract_ref(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    fields = [
        {'id': 0, 'key': 'project_name', 'label': '项目名称', 'field_type': 'text',
         'required': True, 'location': {'type': 'paragraph', 'body_index': 0}},
        {'id': 1, 'key': 'supplier', 'label': '供应商名称', 'field_type': 'text',
         'required': True, 'location': {'type': 'paragraph', 'body_index': 1}},
    ]
    tpl = template_def.TemplateDef.create('直接合同测试模板', '', fields)
    template_path = tpl.save()

    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'direct-token'
    response = client.post(
        f'/procurement/projects/{project_id}/direct-contract',
        data={'csrf_token': 'direct-token', 'template_filename': os.path.basename(template_path)},
        follow_redirects=False,
    )
    assert response.status_code == 302
    assert response.headers['Location'].endswith('/editor')

    editor_page = client.get('/editor')
    editor_html = editor_page.get_data(as_text=True)
    assert 'value="结构件加工竞争性谈判"' in editor_html
    assert 'value="供应商A"' in editor_html

    response = client.post('/generate', data={
        'csrf_token': 'direct-token',
        'project_name': '结构件加工竞争性谈判',
        'field_0': '结构件加工竞争性谈判',
        'field_1': '供应商A',
    })
    assert response.status_code == 200, response.get_data(as_text=True)
    contract_id = int(response.headers['X-Contract-Id'])
    response.close()
    links = procurement_store.get_project_contract_links(project_id)
    assert any(row['contract_id'] == contract_id and row['source_type'] == 'direct_contract' for row in links)


def test_direct_procurement_batch_links_every_contract(app, client):
    project_id, _suppliers = _project_with_items_and_suppliers()
    fields = [
        {'id': 0, 'key': 'project_name', 'label': '项目名称', 'field_type': 'text'},
        {'id': 1, 'key': 'supplier', 'label': '供应商名称', 'field_type': 'text'},
        {'id': 2, 'key': 'contract_no', 'label': '合同编号', 'field_type': 'text'},
    ]
    tpl = template_def.TemplateDef.create('直接采购批量合同模板', '', fields)
    template_path = tpl.save()
    data = procurement_project_service.prepare_direct_contract_session(
        project_id,
        os.path.basename(template_path),
        app.extensions['runtime_paths'],
    )
    sid = uuid.uuid4().hex
    helpers.save_session_data(sid, data)
    with client.session_transaction() as flask_session:
        flask_session['sid'] = sid
        flask_session['_csrf_token'] = 'direct-batch-token'

    editor_html = client.get('/editor').get_data(as_text=True)
    assert 'id="batchToggle"' in editor_html

    response = client.post('/generate-batch', data={
        'csrf_token': 'direct-batch-token',
        'project_name': data['project_name'],
        'field_0': data['project_name'],
        'field_1': '',
        'field_2': 'DIRECT-BATCH',
        'batch_counterparties': '供应商甲\n供应商乙',
        'batch_field_key': 'supplier',
    })
    assert response.status_code == 200, response.get_data(as_text=True)
    with zipfile.ZipFile(io.BytesIO(response.get_data())) as archive:
        assert len(archive.namelist()) == 2
    response.close()

    links = procurement_store.get_project_contract_links(project_id)
    assert len(links) == 2
    assert {row['source_type'] for row in links} == {'direct_contract'}
    assert {row['contract_no'] for row in links} == {
        'DIRECT-BATCH-001', 'DIRECT-BATCH-002',
    }


def test_direct_procurement_link_failure_discards_generated_records(app, client, monkeypatch):
    project_id, _suppliers = _project_with_items_and_suppliers()
    fields = [
        {'id': 0, 'key': 'supplier', 'label': '供应商名称', 'field_type': 'text'},
        {'id': 1, 'key': 'contract_no', 'label': '合同编号', 'field_type': 'text'},
    ]
    tpl = template_def.TemplateDef.create('采购关联失败回滚模板', '', fields)
    template_path = tpl.save()
    data = procurement_project_service.prepare_direct_contract_session(
        project_id,
        os.path.basename(template_path),
        app.extensions['runtime_paths'],
    )
    sid = uuid.uuid4().hex
    helpers.save_session_data(sid, data)
    with client.session_transaction() as flask_session:
        flask_session['sid'] = sid
        flask_session['_csrf_token'] = 'link-failure-token'

    def fail_link(*_args, **_kwargs):
        raise RuntimeError('simulated procurement link failure')

    monkeypatch.setattr(procurement_store, 'add_contract_ref', fail_link)

    single = client.post('/generate', data={
        'csrf_token': 'link-failure-token',
        'field_0': '供应商甲',
        'field_1': 'LINK-FAIL-SINGLE',
    })
    assert single.status_code == 500
    assert ledger_store.list_contracts(per_page=100)['total'] == 0

    batch = client.post('/generate-batch', data={
        'csrf_token': 'link-failure-token',
        'field_0': '',
        'field_1': 'LINK-FAIL-BATCH',
        'batch_counterparties': '供应商甲\n供应商乙',
        'batch_field_key': 'supplier',
    })
    assert batch.status_code == 500
    assert ledger_store.list_contracts(per_page=100)['total'] == 0

    output_dir = app.extensions['runtime_paths'].output_dir
    assert not list(output_dir.glob(f'{sid}_*_output.docx'))
    assert not list(output_dir.glob(f'{sid}_batch_*.docx'))
    assert not list(output_dir.glob(f'{sid}_*_batch.zip'))


def test_supplier_delete_cleans_temporary_quote_jobs_and_files(app):
    project_id = procurement_store.create_project({
        'project_no': 'CG-SUPPLIER-CLEAN', 'project_name': '供应商清理测试',
    })
    supplier_id = procurement_store.add_project_supplier(project_id, {
        'supplier_name': '待删除供应商',
    })
    project = procurement_store.get_project(project_id)
    import_path = procurement_file_service.target_path(
        project, 'supplier_quote', 'invalid.xlsx'
    )
    mapping_path = procurement_file_service.target_path(
        project, 'supplier_quote', 'mapping.xlsx'
    )
    import_path.write_bytes(b'invalid')
    mapping_path.write_bytes(b'mapping')
    import_job_id = procurement_store.create_import_job({
        'project_id': project_id,
        'supplier_id': supplier_id,
        'quote_round': 1,
        'original_name': 'invalid.xlsx',
        'relative_path': procurement_file_service.relative_path(import_path),
        'file_sha256': 'invalid-hash',
        'payload': {},
        'errors': ['文件无效'],
    })
    mapping_job_id = procurement_store.create_mapping_job({
        'project_id': project_id,
        'supplier_id': supplier_id,
        'quote_round': 1,
        'source_type': 'xlsx',
        'original_name': 'mapping.xlsx',
        'relative_path': procurement_file_service.relative_path(mapping_path),
        'file_sha256': 'mapping-hash',
        'source': {'tables': [{'name': '报价表', 'rows': []}]},
    })

    procurement_project_service.delete_supplier(project_id, supplier_id)

    assert procurement_store.get_project_supplier(supplier_id) is None
    assert procurement_store.get_import_job(import_job_id) is None
    assert procurement_store.get_mapping_job(mapping_job_id) is None
    assert not import_path.exists()
    assert not mapping_path.exists()


def test_negotiation_can_start_without_quotes(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'negotiation-token'
    page = client.get(f'/procurement/projects/{project_id}/negotiation')
    html = page.get_data(as_text=True)
    assert '无报价，可手工录入' in html
    response = client.post(
        f'/procurement/projects/{project_id}/negotiation',
        data={
            'csrf_token': 'negotiation-token',
            'round_no': '1',
            'meeting_date': '2026-06-25',
            'summary': '直接进入谈判',
            f'amount_{suppliers[0]}': '8800.50',
            f'delivery_{suppliers[0]}': '20天',
            f'payment_{suppliers[0]}': '验收后付款',
            f'commitment_{suppliers[0]}': '按期交付',
        },
        follow_redirects=False,
    )
    assert response.status_code == 302
    rounds = procurement_store.list_negotiation_rounds(project_id)
    assert rounds[0]['commitments'][0]['quote_amount_minor'] == 880050
    minutes_path = project_document_service.generate_negotiation_minutes(project_id)
    minutes = Document(minutes_path)
    minutes_text = '\n'.join(
        [paragraph.text for paragraph in minutes.paragraphs]
        + [cell.text for table in minutes.tables for row in table.rows for cell in row.cells]
    )
    assert '谈判人员签字' in minutes_text
    assert '供应商代表' in minutes_text
    _assert_docx_uses_fangsong(minutes)


def test_negotiation_round_can_be_edited_from_existing_record(app, client):
    project_id, suppliers = _project_with_items_and_suppliers()
    negotiation_service.save_round(project_id, {
        'round_no': '1',
        'meeting_date': '2026-06-25',
        'summary': '原始谈判记录',
        f'amount_{suppliers[0]}': '8800.50',
        f'delivery_{suppliers[0]}': '20天',
        f'payment_{suppliers[0]}': '验收后付款',
        f'commitment_{suppliers[0]}': '按期交付',
    })
    page = client.get(f'/procurement/projects/{project_id}/negotiation?round_no=1')
    assert page.status_code == 200
    html = page.get_data(as_text=True)
    assert '编辑第 1 轮谈判' in html
    assert 'value="8800.50"' in html
    assert '原始谈判记录' in html

    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = 'edit-round-token'
    response = client.post(
        f'/procurement/projects/{project_id}/negotiation',
        data={
            'csrf_token': 'edit-round-token',
            'round_no': '1',
            'meeting_date': '2026-06-26',
            'summary': '更新后的谈判记录',
            f'amount_{suppliers[0]}': '8700',
            f'delivery_{suppliers[0]}': '18天',
            f'payment_{suppliers[0]}': '到货验收后付款',
            f'commitment_{suppliers[0]}': '提前交付',
        },
        follow_redirects=False,
    )
    assert response.status_code == 302
    rounds = procurement_store.list_negotiation_rounds(project_id)
    assert len(rounds) == 1
    assert rounds[0]['meeting_date'] == '2026-06-26'
    assert rounds[0]['summary'] == '更新后的谈判记录'
    edited = rounds[0]['commitments'][0]
    assert edited['quote_amount_minor'] == 870000
    assert edited['delivery_period'] == '18天'
    assert edited['commitment'] == '提前交付'


def test_payment_due_soon_crosses_month_boundary(app, client, monkeypatch):
    import routes.payments_bp as payments_bp

    class FixedDate(date):
        @classmethod
        def today(cls):
            return cls(2026, 6, 25)

    monkeypatch.setattr(payments_bp, 'date', FixedDate)
    contract_id = ledger_store.create_contract_with_plans(
        {'contract_no': 'PAY-CROSS-001', 'title': '跨月付款测试'}, {}, 'cross.docx',
        [{
            'phase_name': '到货款',
            'confirm_status': 'confirmed',
            'payment_status': 'unpaid',
            'due_date': '2026-07-01',
            'due_amount': 1000,
            'paid_amount': 0,
        }],
    )
    assert contract_id
    response = client.get('/payment-plans')
    html = response.get_data(as_text=True)
    assert '2026-07-01' in html
    assert '即将到期' in html


def test_procurement_routes_render_and_reject_missing_csrf(app, client):
    response = client.get('/procurement', follow_redirects=False)
    assert response.status_code == 302
    assert response.headers['Location'].endswith('/procurement/projects')

    response = client.get('/procurement/projects')
    assert response.status_code == 200
    assert '采购前置工作台' in response.get_data(as_text=True)
    assert client.get('/procurement/history-prices').status_code == 200
    assert client.get('/procurement/projects/new').status_code == 200
    project_id, _ = _project_with_items_and_suppliers()
    detail = client.get(f'/procurement/projects/{project_id}')
    assert detail.status_code == 200
    detail_html = detail.get_data(as_text=True)
    assert 'data-testid="procurement-workflow"' in detail_html
    assert '直接生成合同' in detail_html
    assert '谈判预案' in detail_html
    list_html = client.get('/procurement/projects').get_data(as_text=True)
    assert '竞争性谈判' in list_html
    assert 'competitive_negotiation' not in list_html
    assert client.get(f'/procurement/projects/{project_id}/quotes/import').status_code == 200
    comparison_html = client.get(f'/procurement/projects/{project_id}/comparison').get_data(as_text=True)
    assert '返回项目' in comparison_html
    response = client.post('/procurement/projects/new', data={'project_name': '无令牌'})
    assert response.status_code == 400


def test_procurement_schema_is_idempotent_and_preserves_contracts(app):
    contract_id = ledger_store.create_contract(
        {'contract_no': 'PROC-MIG-001', 'title': '迁移保护测试'}, {}, 'migration.docx'
    )
    procurement_store.init_db()
    procurement_store.init_db()
    assert ledger_store.get_contract(contract_id)['contract_no'] == 'PROC-MIG-001'
    with ledger_store.get_conn() as conn:
        version = conn.execute('SELECT MAX(version) FROM procurement_schema_version').fetchone()[0]
        assert version == procurement_store.schema.CURRENT_SCHEMA_VERSION
