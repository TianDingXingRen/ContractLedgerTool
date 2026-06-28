# -*- coding: utf-8 -*-
import os, tempfile, unittest
from docx import Document
from docx.oxml.ns import qn
import docx_builder, field_eval
from utils.security import MAX_TABLE_ROWS


class ParagraphReplacementTests(unittest.TestCase):
    def setUp(self):
        self.doc = Document()
        self.tmpdir = tempfile.mkdtemp()
    def tearDown(self):
        import shutil; shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_simple_placeholder_replacement(self):
        p = self.doc.add_paragraph('甲方：{甲方名称}')
        docx_builder.apply_text_field(self.doc,
            {'type': 'paragraph', 'body_index': 0, 'placeholder': '{甲方名称}'},
            'test_company', '甲方名称', 'party_a')
        text = ''.join(t.text or '' for t in p._p.iter(qn('w:t')))
        self.assertIn('test_company', text)
        self.assertNotIn('{甲方名称}', text)

    def test_no_placeholder_fallback_to_label(self):
        p = self.doc.add_paragraph('甲方：party_a_value')
        docx_builder.apply_text_field(self.doc,
            {'type': 'paragraph', 'body_index': 0, 'placeholder': ''},
            'test_company', '甲方名称', 'party_a')
        text = ''.join(t.text or '' for t in p._p.iter(qn('w:t')))
        self.assertIn('test_company', text)

    def test_body_index_out_of_range_raises(self):
        with self.assertRaises(ValueError):
            docx_builder.apply_text_field(self.doc,
                {'type': 'paragraph', 'body_index': 99, 'placeholder': '{x}'},
                'v', 'l', 'k')


class TableOperationTests(unittest.TestCase):
    def setUp(self):
        self.doc = Document()
        self.tmpdir = tempfile.mkdtemp()
    def tearDown(self):
        import shutil; shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _make_table(self):
        t = self.doc.add_table(rows=3, cols=3)
        t.rows[0].cells[0].text = 'product'
        t.rows[0].cells[1].text = 'qty'
        t.rows[0].cells[2].text = 'price'
        t.rows[1].cells[0].text = '{product}'
        t.rows[1].cells[1].text = '{qty}'
        t.rows[1].cells[2].text = '{price}'
        return t

    def test_clone_row_with_data(self):
        self._make_table()
        fd = {'key': 'items', 'label': 'items', 'field_type': 'table',
              'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1},
              'columns': [{'key': 'product', 'label': 'product'},
                          {'key': 'qty', 'label': 'qty'},
                          {'key': 'price', 'label': 'price'}]}
        docx_builder.apply_table_field(self.doc, fd, [
            {'product': 'A', 'qty': '10', 'price': '100'},
            {'product': 'B', 'qty': '20', 'price': '200'}])
        self.assertEqual(len(self.doc.tables[0].rows), 4)  # header + template(as data) + clone + empty row
        self.assertIn('A', self.doc.tables[0].rows[1].cells[0].text)

    def test_empty_data_clears_markers_without_removing_row(self):
        t = self._make_table()
        n = len(t.rows)
        fd = {'key': 'items', 'label': 'items', 'field_type': 'table',
              'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1},
              'columns': [{'key': 'product', 'label': 'product'},
                          {'key': 'qty', 'label': 'qty'},
                          {'key': 'price', 'label': 'price'}]}
        docx_builder.apply_table_field(self.doc, fd, [])
        self.assertEqual(len(self.doc.tables[0].rows), n)
        self.assertNotIn('{product}', self.doc.tables[0].rows[1].cells[0].text)

    def test_empty_table_does_not_shift_following_table_cell_fields(self):
        t = self._make_table()
        t.rows[2].cells[2].text = 'total: {total}'
        fd = {'key': 'items', 'label': 'items', 'field_type': 'table',
              'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1},
              'columns': [{'key': 'product', 'label': 'product'},
                          {'key': 'qty', 'label': 'qty'},
                          {'key': 'price', 'label': 'price'}]}
        docx_builder.apply_table_field(self.doc, fd, [])
        docx_builder.apply_text_field(self.doc,
            {'type': 'table_cell', 'table_index': 0, 'row_index': 2,
             'col_index': 2, 'placeholder': '{total}'},
            '100', 'total', 'total')
        self.assertIn('total: 100', self.doc.tables[0].rows[2].cells[2].text)

    def test_table_index_oob_raises(self):
        with self.assertRaises(ValueError):
            docx_builder.apply_table_field(self.doc,
                {'key': 'x', 'field_type': 'table',
                 'location': {'type': 'table', 'table_index': 99, 'template_row_index': 0},
                 'columns': [{'key': 'a', 'label': 'A'}]}, [])

    def test_missing_columns_no_crash(self):
        self._make_table()
        docx_builder.apply_table_field(self.doc,
            {'key': 'x', 'field_type': 'table',
             'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1}},
            [{'col_0': 'x'}])


class GenerateFromScratchTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
    def tearDown(self):
        import shutil; shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_text_fields_rendered(self):
        tpl = {'template_name': 'test', 'fields': [
            {'key': 'a', 'label': 'A', 'field_type': 'text'},
            {'key': 'b', 'label': 'B', 'field_type': 'text'}]}
        out = os.path.join(self.tmpdir, 't.docx')
        docx_builder.generate_from_scratch(tpl, {'a': 'hello', 'b': 'world'}, out)
        doc = Document(out)
        text = '\n'.join(p.text or '' for p in doc.paragraphs)
        self.assertIn('hello', text)
        self.assertIn('world', text)

    def test_table_field_rendered(self):
        tpl = {'template_name': 'test', 'fields': [
            {'key': 'items', 'label': 'items', 'field_type': 'table',
             'columns': [{'key': 'n', 'label': 'name'}, {'key': 'q', 'label': 'qty'}]}]}
        out = os.path.join(self.tmpdir, 't.docx')
        docx_builder.generate_from_scratch(tpl, {'items': [
            {'n': 'A', 'q': '10'}, {'n': 'B', 'q': '20'}]}, out)
        doc = Document(out)
        self.assertEqual(len(doc.tables), 1)
        self.assertIn('A', doc.tables[0].rows[1].cells[0].text)

    def test_too_many_rows_raises(self):
        tpl = {'template_name': 't', 'fields': [
            {'key': 't', 'field_type': 'table',
             'columns': [{'key': 'a', 'label': 'A'}]}]}
        out = os.path.join(self.tmpdir, 't.docx')
        with self.assertRaises(ValueError):
            docx_builder.generate_from_scratch(tpl,
                {'t': [{'a': str(i)} for i in range(MAX_TABLE_ROWS + 1)]}, out)

    def test_normal_rows_accepted(self):
        tpl = {'template_name': 't', 'fields': [
            {'key': 't', 'field_type': 'table',
             'columns': [{'key': 'a', 'label': 'A'}]}]}
        out = os.path.join(self.tmpdir, 't.docx')
        docx_builder.generate_from_scratch(tpl,
            {'t': [{'a': str(i)} for i in range(5)]}, out)
        self.assertTrue(os.path.exists(out))


class EndToEndTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.src = os.path.join(self.tmpdir, 's.docx')
        d = Document()
        d.add_paragraph('A: {A}')
        d.add_paragraph('B: {B}')
        d.save(self.src)
        self.doc = Document(self.src)
    def tearDown(self):
        import shutil; shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_full_workflow(self):
        fields = [
            {'key': 'a', 'label': 'A', 'field_type': 'text',
             'location': {'type': 'paragraph', 'body_index': 0, 'placeholder': '{A}'}},
            {'key': 'b', 'label': 'B', 'field_type': 'text',
             'location': {'type': 'paragraph', 'body_index': 1, 'placeholder': '{B}'}}]
        vals = {'a': 'hello', 'b': 'world'}
        for f in field_eval.sort_fields_by_dependency(fields):
            docx_builder.apply_text_field(self.doc, f.get('location', {}),
                vals.get(f['key'], ''), f.get('label', ''), f['key'])
        out = os.path.join(self.tmpdir, 'o.docx')
        self.doc.save(out)
        doc = Document(out)
        text = '\n'.join(p.text or '' for p in doc.paragraphs)
        self.assertIn('hello', text)
        self.assertIn('world', text)

