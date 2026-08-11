import io
import os
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor
from contextlib import contextmanager

import pytest
from docx import Document

import ledger_store
from services.contract_import_service import (
    ContractImportRequest,
    ContractImportService,
)
from services.generation_recovery_service import GenerationRecoveryService
from utils import helpers


DOCX_MIME = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
)


class SimulatedProcessTermination(BaseException):
    pass


def _contract_docx_bytes(*, contract_no='HT-2026-001', suffix=''):
    document = Document()
    table = document.add_table(rows=6, cols=2)
    values = [
        ('合同编号', contract_no),
        ('合同名称', f'设备维保合同{suffix}'),
        ('乙方', '示例供应商有限公司'),
        ('合同金额', '人民币 1,234,567.89 元'),
        ('签订日期', '2026年7月1日'),
        ('有效期至', '2027年6月30日'),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
    document.add_paragraph('合同签订后10日内支付合同金额的30%作为预付款。')
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _csrf(client, value='contract-import-token'):
    with client.session_transaction() as flask_session:
        flask_session['_csrf_token'] = value
    return value


def _preview(client, payload, filename='外部合同.docx'):
    token = _csrf(client)
    return client.post(
        '/contracts/import/preview',
        data={'csrf_token': token, 'file': (io.BytesIO(payload), filename)},
        content_type='multipart/form-data',
        follow_redirects=False,
    )


def _confirm_form(preview, token, **overrides):
    summary = preview['summary']
    form = {
        'csrf_token': token,
        'contract_no': summary.get('contract_no') or '',
        'title': summary.get('title') or '',
        'counterparty': summary.get('counterparty') or '',
        'amount': summary.get('amount') if summary.get('amount') is not None else '',
        'sign_date': summary.get('sign_date') or '',
        'expiry_date': summary.get('expiry_date') or '',
        'owner': summary.get('owner') or '',
        'status': summary.get('status') or 'draft',
        'project_name': summary.get('project_name') or '',
        'subsystem_name': summary.get('subsystem_name') or '',
        'coverage_mode': summary.get('coverage_mode') or 'not_applicable',
        'coverage_start': summary.get('coverage_start') or '',
        'coverage_end': summary.get('coverage_end') or '',
        'plan_count': len(preview.get('plans') or []),
        'rule_count': len(preview.get('rules') or []),
    }
    for index, plan in enumerate(preview.get('plans') or []):
        prefix = f'plan_{index}_'
        form.update({
            prefix + 'include': '1',
            prefix + 'phase_name': plan.get('phase_name') or '',
            prefix + 'payment_type': plan.get('payment_type') or 'conditional',
            prefix + 'trigger_event': plan.get('trigger_event') or '',
            prefix + 'trigger_days': (
                plan.get('trigger_days')
                if plan.get('trigger_days') is not None else ''
            ),
            prefix + 'expected_trigger_date': (
                plan.get('expected_trigger_date') or ''
            ),
            prefix + 'due_date': plan.get('due_date') or '',
            prefix + 'ratio': (
                plan.get('ratio') if plan.get('ratio') is not None else ''
            ),
            prefix + 'due_amount': (
                plan.get('due_amount')
                if plan.get('due_amount') is not None else ''
            ),
            prefix + 'paid_amount': '0',
            prefix + 'paid_date': '',
            prefix + 'condition_text': plan.get('condition_text') or '',
            prefix + 'source_text': plan.get('source_text') or '',
            prefix + 'confidence': plan.get('confidence') or 'low',
            prefix + 'confirm_status': 'pending',
            prefix + 'payment_status': 'unpaid',
            prefix + 'remark': plan.get('remark') or '',
        })
    for index, rule in enumerate(preview.get('rules') or []):
        prefix = f'rule_{index}_'
        form[prefix + 'include'] = '1'
        for key in (
            'group_key', 'phase_name', 'rule_type', 'scope',
            'trigger_event_type', 'trigger_event', 'trigger_days', 'due_date',
            'conditions_json', 'condition_logic', 'amount_basis',
            'amount_basis_text', 'ratio', 'explicit_amount',
            'calculated_amount', 'repeat_mode', 'source_text', 'source_block',
            'rule_fingerprint', 'source_fingerprint', 'extractor_version',
            'rule_version', 'parse_status', 'reason_codes_json',
            'confirm_status',
        ):
            value = rule.get(key)
            form[prefix + key] = '' if value is None else value
    form.update(overrides)
    return form


def _active_import(client):
    with client.session_transaction() as flask_session:
        sid = flask_session['contract_import_sid']
    return sid, helpers.load_session_data(sid)


def test_preview_extracts_table_fields_diagnostics_and_pending_plans(
    tmp_db, tmp_path
):
    uploads = tmp_path / 'uploads'
    output = tmp_path / 'output'
    uploads.mkdir()
    source = uploads / 'external.docx'
    source.write_bytes(_contract_docx_bytes())
    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=output,
    )

    preview = service.preview_file(source, '外部设备合同.docx')

    assert preview.summary == {
        'contract_no': 'HT-2026-001',
        'title': '设备维保合同',
        'counterparty': '示例供应商有限公司',
        'amount': 1234567.89,
        'sign_date': '2026-07-01',
        'expiry_date': '2027-06-30',
        'owner': '',
        'status': 'draft',
        'project_name': '',
        'subsystem_name': '',
        'coverage_mode': '',
        'coverage_not_applicable': False,
        'coverage_start': None,
        'coverage_end': None,
    }
    diagnostics = {item['field']: item for item in preview.diagnostics}
    assert diagnostics['contract_no']['confidence'] == 'high'
    assert '合同编号' in diagnostics['contract_no']['evidence']
    assert preview.plans
    assert all(plan['confirm_status'] == 'pending' for plan in preview.plans)
    assert all(plan['payment_status'] == 'unpaid' for plan in preview.plans)


def test_preview_uses_filename_and_never_invents_contract_number(tmp_db, tmp_path):
    uploads = tmp_path / 'uploads'
    uploads.mkdir()
    source = uploads / 'fallback.docx'
    document = Document()
    document.add_paragraph('这是没有结构化字段的正文。')
    document.save(source)
    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=tmp_path / 'output',
    )

    preview = service.preview_file(source, '人工编写采购合同.docx')

    assert preview.summary['title'] == '人工编写采购合同'
    assert preview.summary['contract_no'] == ''


def test_preview_ignores_unfilled_template_placeholders(tmp_db, tmp_path):
    uploads = tmp_path / 'uploads'
    uploads.mkdir()
    source = uploads / 'template.docx'
    document = Document()
    table = document.add_table(rows=3, cols=2)
    for row, values in zip(table.rows, (
        ('合同编号', '{合同编号}'),
        ('合同名称', '{合同名称}'),
        ('乙方', '{乙方单位名称}'),
    )):
        row.cells[0].text, row.cells[1].text = values
    document.save(source)
    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=tmp_path / 'output',
    )

    preview = service.preview_file(source, '待填写采购合同.docx')

    assert preview.summary['contract_no'] == ''
    assert preview.summary['title'] == '待填写采购合同'
    assert preview.summary['counterparty'] == ''


@pytest.mark.parametrize('archive_kind', ['forged', 'traversal', 'bomb'])
def test_preview_rejects_unsafe_docx_archives(tmp_db, tmp_path, archive_kind):
    uploads = tmp_path / 'uploads'
    uploads.mkdir()
    source = uploads / 'unsafe.docx'
    if archive_kind == 'forged':
        source.write_bytes(b'not-a-zip')
    else:
        with zipfile.ZipFile(source, 'w', zipfile.ZIP_DEFLATED) as archive:
            archive.writestr('[Content_Types].xml', b'<Types/>')
            archive.writestr('word/document.xml', b'<document/>')
            if archive_kind == 'traversal':
                archive.writestr('../escape.xml', b'x')
            else:
                archive.writestr('word/expanded.xml', b'A' * (2 * 1024 * 1024))
    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=tmp_path / 'output',
    )

    with pytest.raises(ValueError):
        service.preview_file(source, 'unsafe.docx')


def test_preview_enforces_total_upload_size(tmp_db, tmp_path):
    uploads = tmp_path / 'uploads'
    uploads.mkdir()
    source = uploads / 'oversized.docx'
    source.write_bytes(_contract_docx_bytes())
    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=tmp_path / 'output',
        max_upload_bytes=source.stat().st_size - 1,
    )

    with pytest.raises(ValueError, match='上传大小'):
        service.preview_file(source, 'oversized.docx')


def test_import_route_full_flow_preserves_bytes_and_is_idempotent(app, client):
    payload = _contract_docx_bytes()
    response = _preview(client, payload, '供应商原稿.docx')
    assert response.status_code == 302
    sid, data = _active_import(client)

    review = client.get(response.headers['Location'])
    assert review.status_code == 200
    assert '供应商原稿.docx' in review.get_data(as_text=True)
    form = _confirm_form(
        data['preview'], 'contract-import-token',
        title='人工复核后的设备维保合同', owner='张三', project_name='园区维保',
    )
    confirmed = client.post(
        f'/contracts/import/{sid}/confirm', data=form, follow_redirects=False
    )
    assert confirmed.status_code == 302
    contract_id = int(confirmed.headers['Location'].rstrip('/').rsplit('/', 1)[-1])

    contract = ledger_store.get_contract(contract_id)
    assert contract['record_origin'] == 'imported'
    assert contract['original_filename'] == '供应商原稿.docx'
    assert contract['title'] == '人工复核后的设备维保合同'
    assert contract['status'] == 'draft'
    plans = ledger_store.list_payment_plans(contract_id=contract_id)
    assert plans
    assert all(plan['confirm_status'] == 'pending' for plan in plans)
    assert all(plan['payment_status'] == 'unpaid' for plan in plans)

    detail = client.get(f'/contracts/{contract_id}')
    assert detail.status_code == 200
    detail_html = detail.get_data(as_text=True)
    assert '外部导入' in detail_html
    assert '供应商原稿.docx' in detail_html
    listing = client.get('/contracts?q=人工复核后的设备维保合同')
    assert '外部导入' in listing.get_data(as_text=True)
    download = client.get(f'/contracts/{contract_id}/download')
    assert download.status_code == 200
    assert download.headers['Content-Type'].startswith(DOCX_MIME)
    assert download.get_data() == payload
    download.close()

    repeated = client.post(
        f'/contracts/import/{sid}/confirm', data=form, follow_redirects=False
    )
    assert repeated.status_code == 302
    assert repeated.headers['Location'].endswith(f'/contracts/{contract_id}')
    assert ledger_store.list_contracts()['total'] == 1


def test_review_uses_string_status_values_and_preserves_invalid_submission(app, client):
    assert _preview(client, _contract_docx_bytes()).status_code == 302
    sid, data = _active_import(client)

    review = client.get(f'/contracts/import/{sid}/review')
    html = review.get_data(as_text=True)
    assert 'value="draft" selected' in html
    assert 'value="ContractStatus.DRAFT"' not in html

    form = _confirm_form(
        data['preview'], 'contract-import-token',
        title='需要保留的人工复核标题', status='invalid-status',
    )
    rejected = client.post(f'/contracts/import/{sid}/confirm', data=form)
    rejected_html = rejected.get_data(as_text=True)

    assert rejected.status_code == 409
    assert '合同状态无效' in rejected_html
    assert 'value="需要保留的人工复核标题"' in rejected_html
    cancelled = client.post(
        f'/contracts/import/{sid}/cancel',
        data={'csrf_token': 'contract-import-token'},
    )
    assert cancelled.status_code == 302


def test_duplicate_sha_and_deleted_contract_number_link_existing(app, client):
    first_payload = _contract_docx_bytes()
    assert _preview(client, first_payload).status_code == 302
    first_sid, first_data = _active_import(client)
    first = client.post(
        f'/contracts/import/{first_sid}/confirm',
        data=_confirm_form(first_data['preview'], 'contract-import-token'),
        follow_redirects=False,
    )
    contract_id = int(first.headers['Location'].rstrip('/').rsplit('/', 1)[-1])
    ledger_store.soft_delete_contract(contract_id)

    duplicate_file = _preview(client, first_payload)
    assert duplicate_file.status_code == 409
    assert f'/contracts/{contract_id}' in duplicate_file.get_data(as_text=True)

    second_payload = _contract_docx_bytes(suffix='补充版')
    assert _preview(client, second_payload, '另一合同.docx').status_code == 302
    second_sid, second_data = _active_import(client)
    duplicate_number = client.post(
        f'/contracts/import/{second_sid}/confirm',
        data=_confirm_form(second_data['preview'], 'contract-import-token'),
        follow_redirects=False,
    )
    assert duplicate_number.status_code == 409
    assert f'/contracts/{contract_id}' in duplicate_number.get_data(as_text=True)
    assert ledger_store.list_contracts(include_deleted=True)['total'] == 1


def test_move_failure_rolls_back_database_and_keeps_staged_upload(app, client):
    payload = _contract_docx_bytes()
    assert _preview(client, payload).status_code == 302
    sid, data = _active_import(client)
    staging_path = helpers.safe_uploaded_docx_path(data['staging_name'])
    service = app.extensions['contract_tool'].contract_import
    original_replace = service.replace_file
    service.replace_file = lambda *_args: (_ for _ in ()).throw(OSError('move failed'))
    try:
        response = client.post(
            f'/contracts/import/{sid}/confirm',
            data=_confirm_form(data['preview'], 'contract-import-token'),
        )
    finally:
        service.replace_file = original_replace

    assert response.status_code == 500
    assert ledger_store.list_contracts()['total'] == 0
    assert os.path.isfile(staging_path)
    assert not list(app.extensions['runtime_paths'].output_dir.glob('imported_*.docx'))


def test_commit_failure_moves_file_back_without_orphan(tmp_path):
    uploads = tmp_path / 'uploads'
    output = tmp_path / 'output'
    uploads.mkdir()
    source = uploads / 'staged.docx'
    payload = _contract_docx_bytes()
    source.write_bytes(payload)

    class FailingCommitLedger:
        CONTRACT_STATUSES = {'draft'}

        @staticmethod
        def create_generation_job(*_args, **_kwargs):
            return None

        @staticmethod
        def update_generation_job(*_args, **_kwargs):
            return None

        @staticmethod
        def get_contract_by_source_sha256(_digest):
            return None

        @staticmethod
        @contextmanager
        def get_conn():
            yield object()
            raise RuntimeError('commit failed')

        @staticmethod
        def create_contract_with_plans(*_args, **_kwargs):
            return 1, 0

    service = ContractImportService(
        ledger_store=FailingCommitLedger,
        uploads_dir=uploads,
        output_dir=output,
    )
    digest = service.sha256_file(source)
    request = ContractImportRequest(
        staging_path=str(source),
        original_filename='原稿.docx',
        source_sha256=digest,
        summary={
            'title': '回滚测试合同', 'status': 'draft',
            'coverage_mode': 'not_applicable',
        },
        plans=[],
    )

    with pytest.raises(RuntimeError, match='commit failed'):
        service.finalize(request)

    assert source.read_bytes() == payload
    assert not list(output.glob('imported_*.docx'))


def test_import_recovery_isolates_file_moved_before_commit(tmp_db, tmp_path):
    uploads = tmp_path / 'uploads'
    output = tmp_path / 'output'
    uploads.mkdir()
    source = uploads / 'staged.docx'
    source.write_bytes(_contract_docx_bytes(contract_no=''))

    def terminate_after_move(source_path, target_path):
        os.replace(source_path, target_path)
        raise SimulatedProcessTermination('power loss after import move')

    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=output,
        replace_file=terminate_after_move,
    )
    digest = service.sha256_file(source)

    with pytest.raises(SimulatedProcessTermination):
        service.finalize(ContractImportRequest(
            staging_path=str(source),
            original_filename='原稿.docx',
            source_sha256=digest,
            summary={
                'title': '断电回滚合同', 'status': 'draft',
                'contract_no': '', 'coverage_mode': 'not_applicable',
            },
            plans=[],
        ))

    assert ledger_store.list_contracts()['total'] == 0
    jobs = ledger_store.list_unfinished_generation_jobs()
    assert len(jobs) == 1 and jobs[0]['state'] == 'staged'
    moved = list(output.glob('imported_*.docx'))
    assert len(moved) == 1

    report = GenerationRecoveryService(
        ledger_store=ledger_store,
        output_dir=output,
    ).reconcile()

    assert report['recovered'] == 1
    assert not moved[0].exists()
    assert len(list((output / '.recovery').iterdir())) == 1


def test_import_recovery_finalizes_commit_before_terminal_marker(tmp_db, tmp_path):
    uploads = tmp_path / 'uploads'
    output = tmp_path / 'output'
    uploads.mkdir()
    source = uploads / 'staged.docx'
    source.write_bytes(_contract_docx_bytes(contract_no=''))

    def terminate_after_commit(_result):
        raise SimulatedProcessTermination('power loss after import commit')

    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=output,
        after_commit=terminate_after_commit,
    )
    digest = service.sha256_file(source)

    with pytest.raises(SimulatedProcessTermination):
        service.finalize(ContractImportRequest(
            staging_path=str(source),
            original_filename='原稿.docx',
            source_sha256=digest,
            summary={
                'title': '断电提交合同', 'status': 'draft',
                'contract_no': '', 'coverage_mode': 'not_applicable',
            },
            plans=[],
        ))

    jobs = ledger_store.list_unfinished_generation_jobs()
    assert len(jobs) == 1 and jobs[0]['state'] == 'file_moved'
    assert ledger_store.get_contract(jobs[0]['contract_id']) is not None

    report = GenerationRecoveryService(
        ledger_store=ledger_store,
        output_dir=output,
    ).reconcile()

    assert report['completed'] == 1
    assert not ledger_store.list_unfinished_generation_jobs()
    assert list(output.glob('imported_*.docx'))


def test_concurrent_confirmation_allows_only_one_source_sha(tmp_db, tmp_path, monkeypatch):
    uploads = tmp_path / 'uploads'
    output = tmp_path / 'output'
    uploads.mkdir()
    payload = _contract_docx_bytes(contract_no='')
    sources = [uploads / f'staged-{index}.docx' for index in range(2)]
    for source in sources:
        source.write_bytes(payload)
    service = ContractImportService(
        ledger_store=ledger_store,
        uploads_dir=uploads,
        output_dir=output,
    )
    digest = service.sha256_file(sources[0])
    monkeypatch.setattr(ledger_store, 'get_contract_by_source_sha256', lambda _sha: None)

    def finalize(source):
        return service.finalize(ContractImportRequest(
            staging_path=str(source),
            original_filename='并发合同.docx',
            source_sha256=digest,
            summary={
                'title': '并发导入合同', 'status': 'draft',
                'contract_no': '', 'coverage_mode': 'not_applicable',
            },
            plans=[],
        ))

    with ThreadPoolExecutor(max_workers=2) as executor:
        outcomes = []
        for future in [executor.submit(finalize, source) for source in sources]:
            try:
                outcomes.append(future.result())
            except ValueError as exc:
                outcomes.append(exc)

    assert sum(not isinstance(result, Exception) for result in outcomes) == 1
    assert sum(isinstance(result, ValueError) for result in outcomes) == 1
    assert ledger_store.list_contracts()['total'] == 1
    assert len(list(output.glob('imported_*.docx'))) == 1


def test_preview_parse_failure_removes_staged_upload(app, client):
    uploads = app.extensions['runtime_paths'].uploads_dir
    response = _preview(client, b'not a docx', '伪造.docx')

    assert response.status_code == 400
    assert not list(uploads.glob('contract_import_*.docx'))


def test_cancel_and_expiry_clean_import_session_files(app, client):
    payload = _contract_docx_bytes()
    assert _preview(client, payload).status_code == 302
    sid, data = _active_import(client)
    staging = app.extensions['runtime_paths'].uploads_dir / data['staging_name']
    session_file = app.extensions['runtime_paths'].sessions_dir / f'{sid}.json'
    assert staging.is_file() and session_file.is_file()

    cancelled = client.post(
        f'/contracts/import/{sid}/cancel',
        data={'csrf_token': 'contract-import-token'},
    )
    assert cancelled.status_code == 302
    assert not staging.exists() and not session_file.exists()

    assert _preview(client, payload, '过期合同.docx').status_code == 302
    expired_sid, expired_data = _active_import(client)
    expired_staging = (
        app.extensions['runtime_paths'].uploads_dir / expired_data['staging_name']
    )
    expired_session = (
        app.extensions['runtime_paths'].sessions_dir / f'{expired_sid}.json'
    )
    old = time.time() - 8 * 24 * 3600
    os.utime(expired_session, (old, old))

    review = client.get(f'/contracts/import/{expired_sid}/review')
    assert review.status_code == 302
    assert not expired_staging.exists() and not expired_session.exists()
