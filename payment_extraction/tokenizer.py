"""Lexical layer for deterministic payment-clause extraction."""

from __future__ import annotations

import math
import re
from datetime import datetime

from docx import Document

from utils.field_utils import normalize_date


PAYMENT_KEYWORDS = (
    '付款', '支付', '付清', '结算', '预付款', '预付', '定金', '尾款', '余款',
    '质保金', '保证金', '货款', '价款', '款项', '验收', '发票', '开票',
    '货到', '到货', '投产通知', '生产通知', '排产通知',
)
STRONG_PAYMENT_KEYWORDS = (
    '付款', '支付', '付清', '结算', '预付款', '定金', '尾款', '余款',
    '质保金', '保证金', '货款', '价款', '款项',
)
PAYMENT_ACTION_KEYWORDS = (
    '付款', '支付', '付清', '结算', '预付款', '预付', '定金', '尾款',
    '余款', '质保金', '保证金', '一次性总付', '分期支付', '扣留',
)
TRIGGER_ONLY_KEYWORDS = (
    '验收', '发票', '开票', '货到', '到货', '投产通知', '生产通知',
    '排产通知', '对账',
)
PLANISH_NO_MONEY_KEYWORDS = (
    '一次性总付', '分期支付', '预付款', '尾款', '余款', '质保金',
    '付款期限', '结算方式', '投产通知', '生产通知', '排产通知',
)
EXCLUDE_PAYMENT_KEYWORDS = (
    '不得', '报销', '赔偿', '违约金', '罚金', '滞纳金', '责任方',
)
PRICE_SUMMARY_KEYWORDS = (
    '合同总价款', '合同款项', '订购', '不含税金额', '税额', '增值税',
    '税率', '含税', '大写', '小写',
)

AMOUNT_BASIS_PATTERNS = (
    ('production_notice_total', r'(?:该|本|当)?(?:次|份)?(?:投产|生产|排产)通知(?:内|中|所列|项下)?(?:的)?(?:产品|货物)?总?价(?:款|金额)?'),
    ('batch_delivery_total', r'(?:该|本|当)?(?:批|批次)(?:产品|货物)?(?:的)?总?(?:价款|金额|货款)'),
    ('accepted_product_total', r'(?:已|实际)?验收(?:合格)?(?:产品|货物)?(?:的)?总?(?:价款|金额)'),
    ('settlement_amount', r'(?:当期|本期|该期)?(?:结算款|结算金额)'),
    ('invoice_amount', r'(?:发票|开票)(?:含税)?金额'),
    ('remaining_contract_amount', r'(?:剩余|余下|未付)(?:合同)?(?:价款|款项|金额)|尾款|余款'),
    ('contract_total_tax_exclusive', r'(?:不含税|税前)(?:合同)?(?:总价|价款|金额)'),
    ('contract_total_tax_inclusive', r'(?:含税)?(?:合同总价款|合同总价|合同金额|合同价款|总合同额)'),
)

PAYMENT_SECTION_HEADING = re.compile(
    r'^(?:第[一二三四五六七八九十\d]+条\s*)?'
    r'(?:合同价款及支付|价款及支付|付款方式|支付方式|结算方式|付款条件|发票与付款|投产及付款)[:：]?$'
)
GENERIC_SECTION_HEADING = re.compile(
    r'^(?:第[一二三四五六七八九十\d]+条|[一二三四五六七八九十]+、)\s*[^，。；]{2,30}$'
)
ACTION_PATTERN = re.compile(r'一次性总付|分期支付|支付|付款|付清|预付|结算|扣留')
NUMBERED_BOUNDARY = re.compile(
    r'(?=(?:（\s*[一二三四五六七八九十\d]+\s*）|'
    r'\(\s*\d+\s*\)|[①②③④⑤⑥⑦⑧⑨⑩]|'
    r'第[一二三四五六七八九十\d]+(?:期|笔|阶段)))'
)


def extract_docx_blocks(path):
    """Return paragraphs and table rows without flattening their boundaries."""
    doc = Document(path)
    try:
        blocks = []
        for index, para in enumerate(doc.paragraphs, start=1):
            text = (para.text or '').strip()
            if text:
                blocks.append({'kind': 'paragraph', 'index': index, 'text': text})
        for table_index, table in enumerate(doc.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                cells = []
                for cell in row.cells:
                    value = ''.join(p.text or '' for p in cell.paragraphs).strip()
                    cells.append(value)
                if any(cells):
                    blocks.append({
                        'kind': 'table_row',
                        'table_index': table_index,
                        'row_index': row_index,
                        'cells': cells,
                        'text': ' | '.join(cells),
                    })
        return blocks
    finally:
        del doc


def extract_docx_text(path):
    return '\n'.join(block['text'] for block in extract_docx_blocks(path))


def coerce_blocks(source):
    if isinstance(source, str):
        return [
            {'kind': 'paragraph', 'index': index, 'text': line.strip()}
            for index, line in enumerate(source.splitlines() or [source], start=1)
            if line.strip()
        ]
    blocks = []
    for index, item in enumerate(source or [], start=1):
        if isinstance(item, dict):
            text = str(item.get('text') or '').strip()
            if text:
                blocks.append({**item, 'text': text})
        else:
            text = str(item or '').strip()
            if text:
                blocks.append({'kind': 'paragraph', 'index': index, 'text': text})
    return blocks


def selected_payment_option(text):
    match = re.search(r'采用以下第\s*([一二三四五六七八九十\d]+)\s*种方式', text)
    if not match:
        return None
    parsed = parse_cn_number(match.group(1))
    return int(parsed) if parsed else None


def payment_snippets(text, payment_context=False):
    text = re.sub(r'[\t\u3000]+', ' ', str(text or ''))
    rough = []
    for numbered in NUMBERED_BOUNDARY.split(text):
        rough.extend(re.split(r'[。；;\n\r]+', numbered))
    snippets = []
    for chunk in rough:
        chunk = re.sub(r'\s+', ' ', chunk).strip(' ，,|')
        if PAYMENT_SECTION_HEADING.fullmatch(chunk):
            continue
        if len(chunk) >= 2 and is_candidate_clause(
            chunk, allow_bare_money=payment_context
        ):
            snippets.append(chunk)
    return snippets


def action_matches(text):
    matches = []
    for match in ACTION_PATTERN.finditer(text):
        suffix = text[match.end():match.end() + 3]
        if suffix.startswith(('条件', '方式', '条款', '计划')):
            continue
        matches.append(match)
    return matches


def split_segments(snippet):
    """Split around payment actions while keeping condition lists intact."""
    actions = action_matches(snippet)
    if len(actions) >= 2:
        boundaries = []
        for current, following in zip(actions, actions[1:]):
            between = snippet[current.end():following.start()]
            delimiter = re.search(r'[，,]', between)
            if delimiter:
                boundaries.append(current.end() + delimiter.start())
        if boundaries:
            pieces = []
            start = 0
            for boundary in boundaries:
                pieces.append(snippet[start:boundary].strip(' ，,'))
                start = boundary + 1
            pieces.append(snippet[start:].strip(' ，,'))
            useful = [piece for piece in pieces if is_candidate_clause(piece)]
            if len(useful) >= 2:
                return useful

    ratios = list(re.finditer(r'\d+(?:\.\d+)?\s*%', snippet))
    if len(ratios) >= 2:
        pieces = [part.strip() for part in re.split(r'[，,]', snippet) if part.strip()]
        ratio_pieces = [part for part in pieces if extract_ratios(part)]
        if len(ratio_pieces) >= 2:
            prefix = next((
                part for part in pieces
                if '方式' in part and not extract_ratios(part)
            ), '')
            return [
                f'{prefix} {part}'.strip() if prefix else part
                for part in ratio_pieces
            ]
    return [snippet]


def is_candidate_clause(text, allow_bare_money=False):
    if any(k in text for k in EXCLUDE_PAYMENT_KEYWORDS):
        return False
    if looks_like_price_summary(text):
        return False
    has_money = bool(extract_ratios(text) or extract_amounts(text))
    has_action = bool(action_matches(text)) or any(
        k in text for k in PAYMENT_ACTION_KEYWORDS
    )
    has_strong = any(k in text for k in STRONG_PAYMENT_KEYWORDS)
    has_trigger = any(k in text for k in TRIGGER_ONLY_KEYWORDS)
    has_planish = any(k in text for k in PLANISH_NO_MONEY_KEYWORDS)
    has_dynamic_basis = detect_amount_basis(text)[0] != 'unknown'
    if allow_bare_money and has_money:
        return True
    if has_money and (has_action or has_trigger or has_dynamic_basis):
        return True
    if has_planish or (has_action and (has_strong or has_trigger)):
        return True
    return False


def mark_payment_context(blocks):
    """Mark blocks following a payment-section heading without flattening them."""
    marked = [dict(block) for block in blocks]
    remaining = 0
    for block in marked:
        text = str(block.get('text') or '').strip()
        if PAYMENT_SECTION_HEADING.fullmatch(text):
            block['payment_context'] = True
            remaining = 12
            continue
        if remaining > 0 and GENERIC_SECTION_HEADING.fullmatch(text):
            remaining = 0
        if remaining > 0:
            block['payment_context'] = True
            remaining -= 1
        elif is_candidate_clause(text):
            block['payment_context'] = False
    return marked


def looks_like_price_summary(text):
    if not any(k in text for k in PRICE_SUMMARY_KEYWORDS):
        return False
    if action_matches(text):
        return False
    if any(k in text for k in ('每次', '每批', '各批次')):
        return False
    return bool(extract_amounts(text) or extract_ratios(text))


def extract_ratios(text):
    ratios = []
    for match in re.finditer(r'(\d+(?:\.\d+)?)\s*%', text):
        context = text[max(0, match.start() - 8):min(len(text), match.end() + 8)]
        if any(k in context for k in ('增值税', '税率', '税额')):
            continue
        ratios.append(float(match.group(1)))
    for match in re.finditer(r'百分之([零一二两三四五六七八九十百千万亿\d.]+)', text):
        parsed = parse_cn_number(match.group(1))
        if parsed is not None:
            ratios.append(float(parsed))
    return ratios


def extract_amounts(text):
    amounts = []
    pattern = r'(?:人民币|RMB|[￥¥])?\s*([0-9][0-9,]*(?:\.\d+)?)\s*(亿元|万元|千元|百元|元)'
    for match in re.finditer(pattern, text, re.I):
        prefix = text[max(0, match.start() - 10):match.start()]
        if any(kw in prefix for kw in ('已付', '已支付', '已预付', '实付', '扣减')):
            continue
        raw = match.group(1).replace(',', '')
        if len(raw.replace('.', '')) > 15:
            continue
        try:
            value = float(raw)
        except (ValueError, OverflowError):
            continue
        if not math.isfinite(value):
            continue
        unit = match.group(2)
        if unit == '亿元':
            value *= 100000000
        elif unit == '万元':
            value *= 10000
        elif unit == '千元':
            value *= 1000
        elif unit == '百元':
            value *= 100
        amounts.append(round(value, 2))
    return amounts


def extract_days(text):
    match = re.search(r'(\d+)\s*(?:个)?(?:工作日|日|天)\s*内?', text)
    return int(match.group(1)) if match else None


def extract_date(text, sign_date=''):
    result = normalize_date(text)
    if result:
        return result
    match = re.search(r'(\d{1,2})月(\d{1,2})日', text)
    if match:
        year = (sign_date or '')[:4] or str(datetime.now().year)
        month, day = match.groups()
        try:
            return datetime(int(year), int(month), int(day)).strftime('%Y-%m-%d')
        except ValueError:
            return ''
    return ''


def detect_amount_basis(text):
    for basis, pattern in AMOUNT_BASIS_PATTERNS:
        match = re.search(pattern, text)
        if match:
            return basis, match.group(0)
    return 'unknown', ''


def parse_cn_number(text):
    """Parse Chinese integer text, including section units 万 and 亿."""
    if re.fullmatch(r'\d+(?:\.\d+)?', text):
        return float(text)
    digit_map = {
        '零': 0, '一': 1, '二': 2, '两': 2, '三': 3, '四': 4,
        '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
    }
    unit_map = {'十': 10, '百': 100, '千': 1000}
    section_units = {'万': 10000, '亿': 100000000}
    result = 0
    section = 0
    digit = 0
    for char in text:
        if char in digit_map:
            digit = digit_map[char]
        elif char in unit_map:
            section += (digit or 1) * unit_map[char]
            digit = 0
        elif char in section_units:
            section = (section + digit) * section_units[char]
            result += section
            section = 0
            digit = 0
        else:
            return None
    result += section + digit
    return float(result) if result > 0 else None
