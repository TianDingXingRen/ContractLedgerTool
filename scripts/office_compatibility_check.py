"""Run structural DOCX compatibility checks."""

from __future__ import annotations

import argparse
import json
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import docx_builder  # noqa: E402
from utils.security import validate_office_archive  # noqa: E402


def build_reference_document(path: Path) -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = '合同兼容性基准'
    paragraph = document.add_paragraph('合同编号：')
    paragraph.add_run('{合同')
    paragraph.add_run('编号}')
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].merge(table.rows[0].cells[1]).text = '物资信息'
    table.rows[0].cells[2].text = '金额'
    for cell, marker in zip(table.rows[1].cells, ('{name}', '{quantity}', '{amount}')):
        cell.text = marker
    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.footer.is_linked_to_previous = False
    second.footer.paragraphs[0].text = '兼容性测试附件'
    document.add_paragraph('中文、English、￥、①')
    document.save(path)


def run_check(work_dir: Path) -> dict:
    source = work_dir / 'compatibility-source.docx'
    output = work_dir / 'compatibility-generated.docx'
    build_reference_document(source)
    validate_office_archive(source)

    document = Document(source)
    docx_builder.apply_text_field(
        document,
        {'type': 'paragraph', 'body_index': 0, 'placeholder': '{合同编号}'},
        'COMPAT-2026-001', '合同编号', 'contract_no',
    )
    docx_builder.apply_table_field(document, {
        'key': 'items', 'label': '采购明细', 'field_type': 'table',
        'location': {'type': 'table', 'table_index': 0, 'template_row_index': 1},
        'columns': [
            {'key': 'name', 'label': '品名'},
            {'key': 'quantity', 'label': '数量'},
            {'key': 'amount', 'label': '金额'},
        ],
    }, [
        {'name': '轴承', 'quantity': '10', 'amount': '1200.00'},
        {'name': '联轴器', 'quantity': '2', 'amount': '680.00'},
    ])
    document.save(output)
    validate_office_archive(output)

    reopened = Document(output)
    body_text = '\n'.join(paragraph.text for paragraph in reopened.paragraphs)
    table_text = [
        [cell.text for cell in row.cells]
        for row in reopened.tables[0].rows
    ]
    checks = {
        'split_run_placeholder': 'COMPAT-2026-001' in body_text,
        'merged_table_header': table_text[0][:2] == ['物资信息', '物资信息'],
        'repeating_table_rows': table_text[1:] == [
            ['轴承', '10', '1200.00'], ['联轴器', '2', '680.00'],
        ],
        'sections_preserved': len(reopened.sections) == 2,
        'header_preserved': reopened.sections[0].header.paragraphs[0].text == '合同兼容性基准',
        'unicode_preserved': '中文、English、￥、①' in body_text,
    }
    if not all(checks.values()):
        failed = [name for name, passed in checks.items() if not passed]
        raise RuntimeError(f'DOCX structural compatibility failed: {failed}')

    return {
        'schema_version': 1,
        'checked_at': datetime.now(timezone.utc).isoformat(),
        'structural_checks': checks,
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument('--output', default=str(ROOT / 'build' / 'office-compatibility.json'))
    args = parser.parse_args()

    with tempfile.TemporaryDirectory(prefix='contract-tool-office-') as temp_dir:
        report = run_check(Path(temp_dir))
    output = Path(args.output).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f'Office compatibility check passed: {output}')
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
