"""文档写入模块

精确地在 .docx 文档的指定位置写入字段值，
支持段落文本替换、表格行克隆等操作。
也支持从模板定义直接生成新文档。
"""

import logging
from copy import deepcopy
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


_log = logging.getLogger('contract_tool')


class DocxBuildError(ValueError):
    """DOCX 文档写入错误，携带字段上下文信息"""

    def __init__(self, message, field_label='', field_key=''):
        self.field_label = field_label
        self.field_key = field_key
        super().__init__(message)


def apply_text_field(doc, location, value, field_label='', field_key=''):
    """在文档的指定位置写入文本值

    参数:
        doc: python-docx Document 对象
        location: {'type': 'paragraph', 'body_index': int, 'placeholder': str}
        value: 要写入的值
        field_label: 字段标签，用于无占位符时的回退匹配
        field_key: 字段 key，用于无占位符时的回退匹配
    """
    try:
        loc_type = location.get('type')
        if loc_type == 'paragraph':
            _apply_to_paragraph(doc, location, str(value), field_label, field_key)
        elif loc_type == 'table_cell':
            _apply_to_table_cell(doc, location, str(value))
        else:
            raise DocxBuildError(
                f'不支持的位置类型: {loc_type}',
                field_label, field_key,
            )
    except DocxBuildError:
        raise
    except Exception as e:
        _log.error(
            '文本字段写入失败: %s(%s), error=%s',
            field_label, field_key, e, exc_info=True,
        )
        raise DocxBuildError(
            f'"{field_label or field_key}" 写入失败: {e}',
            field_label, field_key,
        ) from e


def _apply_to_paragraph(doc, location, value, field_label='', field_key=''):
    """按 body_index 定位段落并替换占位符"""
    body = doc.element.body
    body_index = location.get('body_index', -1)
    placeholder = location.get('placeholder', '')

    para = _find_para_by_index(body, body_index)
    if para is None:
        raise DocxBuildError(
            f'未找到 body_index {body_index} 对应的段落',
            field_label, field_key,
        )

    if placeholder:
        _replace_in_para(para, placeholder, value)
        return

    # 无占位符时，尝试用 {标签} {key} 标签 key 匹配段落文本
    full_text = _get_para_text(para)
    # 先尝试带花括号的占位符格式
    for target in ['{' + field_label + '}', '{' + field_key + '}']:
        if target and target in full_text:
            _replace_in_para(para, target, value)
            return
    # 再尝试纯文本匹配（仅当 label/key 较长时，避免误匹配短词如"乙方"）
    for target in [field_label, field_key]:
        if not target or len(target) < 3:
            continue
        if target and target in full_text:
            _replace_in_para(para, target, value)
            return

    # 正则回退：搜索段落中的所有 {xxx} 标记，匹配标签或 key
    for m in re.finditer(r'\{([^}]+)\}', full_text):
        content = m.group(1).strip()
        if content == field_label or content == field_key:
            _replace_in_para(para, m.group(0), value)
            return

    # 所有匹配方式都失败，无法定位写入位置
    raise DocxBuildError(
        f'字段 "{field_label or field_key}" 的占位符在段落中未找到，'
        f'请检查模板是否正确包含该字段的占位符标记',
        field_label, field_key,
    )


def _apply_to_table_cell(doc, location, value):
    """按 table_index + row_index + col_index 定位单元格并写入"""
    table_index = location.get('table_index', -1)
    row_index = location.get('row_index', -1)
    col_index = location.get('col_index', -1)
    placeholder = location.get('placeholder', '')

    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise DocxBuildError(f'未找到 table_index {table_index}',
                             location.get('placeholder', ''))
    table = tables[table_index]
    if row_index < 0 or row_index >= len(table.rows):
        raise DocxBuildError(f'未找到 row_index {row_index}',
                             location.get('placeholder', ''))
    row = table.rows[row_index]
    if col_index < 0 or col_index >= len(row.cells):
        raise DocxBuildError(f'未找到 col_index {col_index}',
                             location.get('placeholder', ''))

    cell = row.cells[col_index]

    if placeholder:
        # 替换单元格中第一个段落的占位符
        for para in cell.paragraphs:
            replaced = _try_replace_in_para_xml(para, placeholder, value)
            if replaced:
                return
        # 没找到占位符，直接设置单元格文本
        _set_cell_text(cell, value)
    else:
        _set_cell_text(cell, value)


def _set_cell_text(cell, value):
    """设置单元格文本，尽量保留第一个 run 的样式。"""
    t_elems = [
        t
        for paragraph in cell.paragraphs
        for t in paragraph._p.iter(qn('w:t'))
    ]
    if t_elems:
        t_elems[0].text = str(value)
        for t in t_elems[1:]:
            t.text = ''
        return

    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(value))


def _cell_has_marker(cell):
    """检查单元格文本中是否包含 {xxx} 标记"""
    for p in cell.paragraphs:
        if re.search(r'\{[^}]+\}', p.text or ''):
            return True
    return False


def apply_table_field(doc, field_def, rows_data):
    """应用表格字段：克隆模板行并填入数据

    参数:
        doc: python-docx Document 对象
        field_def: 表格字段定义（含 location 和 columns）
        rows_data: [{col_key: value, ...}, ...]

    说明:
        - 只覆盖包含 {xxx} 标记的单元格，无标记的单元格保留原始内容
        - 对 malformed 文档抛出 DocxBuildError 而非原始异常
    """
    field_label = field_def.get('label', field_def.get('key', 'table'))
    try:
        _apply_table_field_impl(doc, field_def, rows_data)
    except DocxBuildError:
        raise
    except Exception as e:
        _log.error(
            '表格字段写入失败: %s, error=%s', field_label, e, exc_info=True,
        )
        raise DocxBuildError(
            f'"{field_label}" 表格写入失败: {e}',
            field_label, field_def.get('key', ''),
        ) from e


def _apply_table_field_impl(doc, field_def, rows_data):
    """apply_table_field 的内部实现，异常由外层包装"""
    location = field_def.get('location', {})
    table_index = location.get('table_index', -1)
    template_row_index = location.get('template_row_index', 0)

    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise DocxBuildError(f'未找到 table_index {table_index}',
                             field_def.get('label', ''))
    table = tables[table_index]
    if template_row_index < 0 or template_row_index >= len(table.rows):
        raise DocxBuildError(f'未找到 template_row_index {template_row_index}',
                             field_def.get('label', ''))

    columns = field_def.get('columns', [])
    template_row = table.rows[template_row_index]
    # 获取模板行的 XML 元素引用
    tr_element = template_row._tr

    # 计算列索引映射
    col_index_map = {}
    for ci, col in enumerate(columns):
        col_index_map[col['key']] = ci

    # 扫描模板行，标记哪些单元格包含 {xxx} 标记
    # 无标记的单元格在填入数据时不会被覆盖，保留原始内容
    template_cell_has_marker = [
        _cell_has_marker(template_row.cells[ci])
        for ci in range(len(template_row.cells))
    ]
    template_row_has_marker = any(template_cell_has_marker)

    # 包含占位符的模板行作为第一条数据行复用，后续数据行克隆该行。
    use_template_as_data = template_row_has_marker

    if not rows_data:
        # 无数据时保留模板行结构，避免后续固定 row_index 的字段定位漂移。
        if template_row_has_marker:
            _clear_markers_in_row(template_row)
        return

    # 克隆行数据（不包括模板行本身）
    clone_count = len(rows_data)
    if use_template_as_data:
        clone_count -= 1  # 模板行本身作为第一行数据

    if clone_count > 0:
        # 确定插入位置（在模板行之后）
        insert_after = tr_element
        for _ in range(clone_count):
            new_tr = deepcopy(tr_element)
            insert_after.addnext(new_tr)
            insert_after = new_tr

    # 填入数据（只替换有 {xxx} 标记的单元格）
    for i, row_data in enumerate(rows_data):
        row = table.rows[template_row_index + i]
        for col in columns:
            col_key = col.get('key')
            if col_key not in col_index_map:
                continue
            ci = col_index_map[col_key]
            if ci >= len(row.cells) or ci >= len(template_cell_has_marker):
                continue
            if template_cell_has_marker[ci]:
                cell = row.cells[ci]
                if not _replace_first_marker_in_cell(cell, str(row_data.get(col_key, ''))):
                    _set_cell_text(cell, str(row_data.get(col_key, '')))
            elif _is_index_column(col, ci):
                _set_cell_text(row.cells[ci], str(i + 1))
            elif str(row_data.get(col_key, '')).strip() and not _get_cell_text(row.cells[ci]).strip():
                _set_cell_text(row.cells[ci], str(row_data.get(col_key, '')))


def _is_index_column(col, col_index):
    label = str(col.get('label', '')).strip()
    key = str(col.get('key', '')).strip().lower()
    return col_index == 0 and (label in {'序号', '序', '编号'} or key in {'序号', 'index', 'no', 'number'})


def _get_cell_text(cell):
    return ''.join(paragraph.text or '' for paragraph in cell.paragraphs)


def _row_is_empty(row):
    """检查行是否为空（无文本内容）"""
    for cell in row.cells:
        for para in cell.paragraphs:
            if para.text.strip():
                return False
    return True


def _clear_markers_in_row(row):
    """清空模板行中的 {xxx} 占位符，同时保留表格行和非占位文本。"""
    for cell in row.cells:
        while _replace_first_marker_in_cell(cell, ''):
            pass


def _get_para_text(para_elem):
    """获取段落的所有文本"""
    return ''.join(t.text or '' for t in para_elem.iter(qn('w:t')))


def _find_para_by_index(body, index):
    """在 body 中查找第 index 个 <w:p> 元素"""
    count = 0
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            if count == index:
                return child
            count += 1
    return None


def _replace_in_para(para_elem, old_text, new_text):
    """在段落中替换文本，保留未命中 run 的原始样式。"""
    return _replace_in_text_elements(list(para_elem.iter(qn('w:t'))), old_text, new_text)


def _try_replace_in_para_xml(para_elem, old_text, new_text):
    """尝试在段落中替换，返回是否成功"""
    elem = getattr(para_elem, '_p', para_elem)
    return _replace_in_text_elements(list(elem.iter(qn('w:t'))), old_text, new_text)


def _replace_in_text_elements(t_elems, old_text, new_text):
    """Replace across w:t nodes while preserving surrounding run formatting."""
    if not t_elems:
        return False

    full = ''.join(t.text or '' for t in t_elems)
    start = full.find(old_text or '')
    if start < 0 or not old_text:
        return False
    end = start + len(old_text)

    cursor = 0
    inserted = False
    for t in t_elems:
        text = t.text or ''
        elem_start = cursor
        elem_end = cursor + len(text)
        cursor = elem_end

        if elem_end <= start or elem_start >= end:
            continue

        before = text[:max(0, start - elem_start)] if elem_start <= start < elem_end else ''
        after = text[max(0, end - elem_start):] if elem_start < end <= elem_end else ''

        if not inserted:
            t.text = before + str(new_text) + after
            inserted = True
        else:
            t.text = after
    return True


def _replace_first_marker_in_cell(cell, value):
    for para in cell.paragraphs:
        elem = para._p
        text = _get_para_text(elem)
        marker = re.search(r'\{([^}]+)\}', text)  # 只匹配完整的 {xxx} 占位符
        if marker and _replace_in_para(elem, marker.group(0), value):
            return True
    return False


def get_para_count(doc):
    """获取文档 body 中的段落总数（调试用）"""
    count = 0
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            count += 1
    return count


def get_table_count(doc):
    """获取文档中的表格总数"""
    return len(doc.tables)


def dump_body_structure(docx_path):
    """调试用：打印文档 body 结构"""
    from docx import Document as _Document
    doc = _Document(docx_path)
    body = doc.element.body

    from docx.oxml.ns import qn as _qn
    para_idx = 0
    table_idx = 0
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = ''
            for t in child.iter(_qn('w:t')):
                if t.text:
                    text += t.text
            logging.debug('  P[%d]: %s', para_idx, text[:80])
            para_idx += 1
        elif tag == 'tbl':
            rows = child.findall(_qn('w:tr'))
            logging.debug('  TABLE[%d]: %d 行', table_idx, len(rows))
            for ri, row in enumerate(rows):
                cells = row.findall(_qn('w:tc'))
                cell_texts = []
                for cell in cells:
                    ct = ''
                    for p in cell.findall(_qn('w:p')):
                        for t in p.iter(_qn('w:t')):
                            if t.text:
                                ct += t.text
                    cell_texts.append(ct.strip())
                logging.debug('    行%d: %s', ri, cell_texts)
            table_idx += 1


def generate_from_scratch(template_def, values, output_path):
    """从模板定义直接生成合同文档（无需源文档）

    参数:
        template_def: 模板定义字典，含 fields 列表
        values: 字段值字典 {field_key: value}
        output_path: 输出路径

    返回:
        output_path
    """
    from utils.security import MAX_TABLE_ROWS, MAX_TABLE_COLUMNS

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # 标题
    title = doc.add_heading(template_def.get('template_name', '合同'), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fields = template_def.get('fields', [])

    for field in fields:
        if field.get('field_type') == 'table':
            # 表格字段
            rows_data = values.get(field['key'], [])
            columns = field.get('columns', [])
            if len(rows_data) > MAX_TABLE_ROWS:
                raise ValueError(f'表格行数超过限制 ({MAX_TABLE_ROWS})')
            if len(columns) > MAX_TABLE_COLUMNS:
                raise ValueError(f'表格列数超过限制 ({MAX_TABLE_COLUMNS})')
            if columns:
                table = doc.add_table(rows=1 + len(rows_data), cols=len(columns))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # 表头
                for ci, col in enumerate(columns):
                    cell = table.rows[0].cells[ci]
                    cell.text = col['label']
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # 数据行
                for ri, row_data in enumerate(rows_data):
                    for ci, col in enumerate(columns):
                        val = str(row_data.get(col['key'], ''))
                        table.rows[ri + 1].cells[ci].text = val

                # 表后空行
                doc.add_paragraph('')
        else:
            # 文本类字段
            label = field.get('label', field.get('key', ''))
            value = str(values.get(field['key'], ''))
            p = doc.add_paragraph()
            run_label = p.add_run(f'{label}：')
            run_label.bold = True
            p.add_run(value)

    try:
        doc.save(output_path)
    except Exception as e:
        _log.error('从模板直接生成合同保存失败: %s', e, exc_info=True)
        raise DocxBuildError(f'保存合同文档失败: {e}') from e
    return output_path
