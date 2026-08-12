"""字段工具：键生成、数值/日期解析、标记检测、表格规范化 —— 从 helpers.py 拆分"""

import re
import json
import math
from typing import Any, Optional
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn

import field_eval
from utils.security import (
    limit_text, bounded_decimal_places,
    MAX_TABLE_COLUMNS, MAX_TABLE_ROWS,
    MAX_TEXT_VALUE_LENGTH,
)
from utils.constants import FieldType


# ═══════════════════════════════════════════════════════
#  Field key helpers
# ═══════════════════════════════════════════════════════

def field_key_from_label(label: str, fallback: str) -> str:
    key = re.sub(r'[^\w\u4e00-\u9fff]', '_', label or '')
    key = re.sub(r'_+', '_', key).strip('_')
    return key or fallback


def unique_key(base_key, fields):
    key = base_key
    counter = 1
    while any(f.get('key') == key for f in fields):
        key = f'{base_key}_{counter}'
        counter += 1
    return key


def safe_col_key(label: str, index: int, existing: set[str]) -> str:
    base = field_eval.make_col_key(label, index)
    base = re.sub(r'[^\w\u4e00-\u9fff]', '_', base).strip('_')[:40] or f'col_{index}'
    key = base
    suffix = 1
    while key in existing:
        key = f'{base}_{suffix}'
        suffix += 1
    existing.add(key)
    return key


# ═══════════════════════════════════════════════════════
#  String / number / date helpers
# ═══════════════════════════════════════════════════════

def safe_filename_part(value, fallback: str = '合同') -> str:
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', str(value or '')).strip(' ._')
    result = (text[:80] or fallback)
    # 检查 Windows 保留文件名（不区分大小写）
    _WIN_RESERVED = {
        'CON', 'PRN', 'AUX', 'NUL',
        'COM1', 'COM2', 'COM3', 'COM4', 'COM5', 'COM6', 'COM7', 'COM8', 'COM9',
        'LPT1', 'LPT2', 'LPT3', 'LPT4', 'LPT5', 'LPT6', 'LPT7', 'LPT8', 'LPT9',
    }
    name_noext = result.rsplit('.', 1)[0].upper()
    if name_noext in _WIN_RESERVED:
        result = '_' + result
    return result


def parse_number(value) -> Optional[float]:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    text = text.replace(',', '').replace('，', '')
    multiplier = 1
    if '亿' in text:
        multiplier *= 100000000
    if '万' in text:
        multiplier *= 10000
    matches = re.findall(r'-?\d+(?:\.\d+)?', text)
    if not matches:
        return None
    if multiplier > 1 and len(matches) > 1:
        numbers = [float(m) for m in matches]
        return round(max(numbers) * multiplier, 2)
    try:
        return round(float(matches[0]) * multiplier, 2)
    except ValueError:
        return None


def normalize_date(value) -> str:
    text = str(value or '').strip()
    if not text:
        return ''
    patterns = [
        r'(20\d{2})[年/-](\d{1,2})[月/-](\d{1,2})日?',
        r'(20\d{2})(\d{2})(\d{2})',
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if not m:
            continue
        y, mo, d = m.groups()
        try:
            return datetime(int(y), int(mo), int(d)).strftime('%Y-%m-%d')
        except ValueError:
            return ''
    iso_match = re.search(r'(20\d{2})-(\d{1,2})-(\d{1,2})', text)
    if iso_match:
        y, mo, d = iso_match.groups()
        try:
            datetime(int(y), int(mo), int(d))
            return iso_match.group(0)
        except ValueError:
            return ''
    return ''


def to_calc_number(value):
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, str):
        cleaned = value.strip().replace(',', '')
        is_pct = cleaned.endswith('%')
        if is_pct:
            cleaned = cleaned[:-1]
        try:
            num = float(cleaned)
            return num / 100 if is_pct else num
        except ValueError:
            return 0
    return 0


def float_or_none(value) -> Optional[float]:
    text = str(value or '').strip().replace(',', '').replace('，', '')
    if not text:
        return None
    try:
        parsed = float(text)
    except (TypeError, ValueError):
        return None
    return parsed if math.isfinite(parsed) else None


def normalize_number_field_value(value, field=None) -> str:
    """严格解析数字字段，并按字段精度和范围返回规范字符串。"""
    text = str(value or '').strip().replace(',', '').replace('，', '')
    if not text:
        return ''
    try:
        number = float(text)
    except (TypeError, ValueError):
        raise ValueError('必须是有效数字')
    if not math.isfinite(number):
        raise ValueError('必须是有限数字')

    definition = field or {}
    min_value = definition.get('min_value')
    max_value = definition.get('max_value')
    if min_value is not None and number < float(min_value):
        raise ValueError(f'不能小于 {min_value}')
    if max_value is not None and number > float(max_value):
        raise ValueError(f'不能大于 {max_value}')

    decimal_places = definition.get('decimal_places')
    if decimal_places is None:
        return str(int(number)) if number.is_integer() else format(number, '.15g')
    decimals = bounded_decimal_places(decimal_places)
    return f'{number:.{decimals}f}'


def int_or_none(value) -> Optional[int]:
    try:
        text = str(value or '').strip()
        return int(text) if text else None
    except ValueError:
        return None


# ═══════════════════════════════════════════════════════
#  Marker detection
# ═══════════════════════════════════════════════════════

def _looks_like_repeat_row(row_data, row_index):
    if row_index <= 0:
        return False
    marker_cells = sum(1 for cell in row_data if cell['markers'])
    has_formula = any(
        marker.group(1).strip().startswith('=')
        for cell in row_data
        for marker in cell['markers']
    )
    return marker_cells >= 2 or has_formula


def _add_table_cell_field(fields, table_idx, row_idx, col_idx, marker):
    name = marker.group(1).strip()
    key = unique_key(field_key_from_label(name, f'table_{table_idx}_r{row_idx}_c{col_idx}'), fields)
    fields.append({
        'key': key,
        'label': name,
        'field_type': FieldType.TEXT,
        'required': False,
        'location': {
            'type': 'table_cell',
            'table_index': table_idx,
            'row_index': row_idx,
            'col_index': col_idx,
            'placeholder': marker.group(0),
        },
    })


def detect_markers(docx_path: str) -> list[dict[str, Any]]:
    """Scan .docx for {field_name} markers and return field definitions."""
    doc = Document(docx_path)
    body = doc.element.body
    marker_re = re.compile(r'\{\{?(.+?)\}\}?')
    fields = []

    para_idx = 0
    table_idx = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
            for m in marker_re.finditer(text):
                name = m.group(1).strip()
                key = re.sub(r'[^\w\u4e00-\u9fff]', '_', name)
                key = re.sub(r'_+', '_', key).strip('_') or f'field_{para_idx}'
                key = unique_key(key, fields)
                fields.append({
                    'key': key, 'label': name, 'field_type': FieldType.TEXT,
                    'required': False,
                    'location': {'type': 'paragraph', 'body_index': para_idx, 'placeholder': m.group(0)},
                })
            para_idx += 1

        elif tag == 'tbl':
            rows = child.findall(qn('w:tr'))

            table_rows_data = []
            has_markers = False
            for ri, row in enumerate(rows):
                cells = row.findall(qn('w:tc'))
                row_data = []
                for ci, cell in enumerate(cells):
                    ctext = ''.join(
                        t.text or ''
                        for p in cell.findall(qn('w:p'))
                        for t in p.iter(qn('w:t'))
                    )
                    markers = list(marker_re.finditer(ctext))
                    if markers:
                        has_markers = True
                    row_data.append({'text': ctext, 'markers': markers})
                table_rows_data.append(row_data)

            if has_markers:
                repeat_row_index = None
                first_marker_row = None
                for ri, row_data in enumerate(table_rows_data):
                    if any(cell['markers'] for cell in row_data):
                        first_marker_row = ri
                        break
                if first_marker_row is not None and _looks_like_repeat_row(table_rows_data[first_marker_row], first_marker_row):
                    repeat_row_index = first_marker_row

                if repeat_row_index is not None:
                    header_row = table_rows_data[0] if table_rows_data else []
                    repeat_row = table_rows_data[repeat_row_index]
                    columns = []
                    col_keys_used = set()

                    for ci, cell_data in enumerate(header_row):
                        col_label = cell_data['text'].strip() or f'列{ci+1}'
                        col_key = field_eval.make_col_key(col_label, ci)
                        base = col_key
                        suffix = 1
                        while col_key in col_keys_used:
                            col_key = f'{base}_{suffix}'
                            suffix += 1
                        col_keys_used.add(col_key)
                        col_type = FieldType.TEXT
                        col_formula = ''

                        if ci < len(repeat_row):
                            for m in repeat_row[ci]['markers']:
                                mn = m.group(1).strip()
                                if mn.startswith('='):
                                    col_type = FieldType.CALCULATED
                                    col_formula = mn[1:].strip()

                        columns.append({
                            'key': col_key,
                            'label': col_label,
                            'field_type': col_type,
                            'formula': col_formula,
                        })

                    if not columns:
                        columns = [{'key': 'col_0', 'label': '内容', 'field_type': FieldType.TEXT, 'formula': ''}]

                    fields.append({
                        'key': unique_key(f'table_{table_idx}', fields),
                        'label': f'表格{table_idx + 1}',
                        'field_type': FieldType.TABLE,
                        'required': False,
                        'location': {
                            'type': 'table',
                            'table_index': table_idx,
                            'template_row_index': repeat_row_index,
                        },
                        'columns': columns,
                    })

                for ri, row_data in enumerate(table_rows_data):
                    if ri == repeat_row_index:
                        continue
                    for ci, cell_data in enumerate(row_data):
                        for marker in cell_data['markers']:
                            _add_table_cell_field(fields, table_idx, ri, ci, marker)

            table_idx += 1

    return fields


# ═══════════════════════════════════════════════════════
#  Table normalization
# ═══════════════════════════════════════════════════════

def filter_table_rows(field, rows_data):
    columns = field.get('columns', [])
    calc_value = FieldType.CALCULATED.value
    editable_keys = {
        col.get('key')
        for col in columns
        if col.get('field_type') != calc_value and col.get('key')
    }
    filtered = []
    for row in rows_data:
        if not isinstance(row, dict):
            raise ValueError(f'{field.get("label", field.get("key"))} 的表格行数据必须是对象')
        keys_to_check = editable_keys or set(row.keys())
        if any(str(row.get(key, '')).strip() for key in keys_to_check):
            normalized_row = {
                str(key)[:80]: limit_text(value)
                for key, value in row.items()
            }
            for col in columns:
                col_key = col.get('key')
                if not col_key or col.get('field_type') == FieldType.CALCULATED:
                    continue
                raw_value = normalized_row.get(col_key, '')
                if col.get('required') and not str(raw_value).strip():
                    raise ValueError(f'{field.get("label", field.get("key"))} 的 {col.get("label", col_key)} 不能为空')
                if not str(raw_value).strip():
                    continue
                if col.get('field_type') == FieldType.NUMBER:
                    try:
                        normalized_row[col_key] = normalize_number_field_value(raw_value, col)
                    except ValueError as e:
                        raise ValueError(f'{field.get("label", field.get("key"))} 的 {col.get("label", col_key)}{e}')
                elif col.get('field_type') == FieldType.SELECT:
                    options = [str(option) for option in col.get('options', [])]
                    if options and str(raw_value) not in options:
                        raise ValueError(f'{field.get("label", field.get("key"))} 的 {col.get("label", col_key)}选项无效')
            filtered.append(normalized_row)
    return filtered


def normalize_table_columns(field, submitted_cols):
    if not isinstance(submitted_cols, list):
        raise ValueError('列定义必须是数组')
    if len(submitted_cols) > MAX_TABLE_COLUMNS:
        raise ValueError(f'列数不能超过 {MAX_TABLE_COLUMNS}')

    normalized = []
    existing = set()
    for idx, col in enumerate(submitted_cols):
        if not isinstance(col, dict):
            raise ValueError('列定义格式错误')
        label = limit_text(col.get('label') or f'列{idx + 1}', 80).strip() or f'列{idx + 1}'
        key = str(col.get('key') or '').strip()
        if not re.fullmatch(r'[\w\u4e00-\u9fff]{1,40}', key) or key in existing:
            key = safe_col_key(label, idx, existing)
        else:
            existing.add(key)
        field_type = col.get('field_type') if col.get('field_type') in {'text', 'number', 'textarea', 'select', 'calculated'} else 'text'
        formula = ''
        decimal_places = 2
        if field_type == 'calculated':
            formula = str(col.get('formula') or '').strip()
            if formula:
                field_eval.validate_formula(formula)
            else:
                field_type = 'text'
            decimal_places = bounded_decimal_places(col.get('decimal_places', 2))
        normalized_col = {
            'key': key,
            'label': label,
            'field_type': field_type,
            'formula': formula if field_type == 'calculated' else '',
            'default_value': '' if field_type == 'calculated' else limit_text(col.get('default_value'), 2000),
        }
        if field_type == 'select':
            raw_options = col.get('options', [])
            if isinstance(raw_options, str):
                raw_options = raw_options.splitlines()
            normalized_col['options'] = [
                limit_text(option, 200) for option in raw_options
                if str(option).strip()
            ][:100]
        if field_type == 'number':
            normalized_col['decimal_places'] = bounded_decimal_places(col.get('decimal_places', 2))
            normalized_col['min_value'] = col.get('min_value')
            normalized_col['max_value'] = col.get('max_value')
        if field_type == 'calculated':
            normalized_col['decimal_places'] = decimal_places
        normalized.append(normalized_col)

    if not normalized:
        raise ValueError('列定义不能为空')
    field_eval.sort_table_columns_by_dependency(normalized)
    return normalized


def apply_submitted_table_columns(fields, form):
    errors = []
    for i, field in enumerate(fields):
        if field.get('field_type') != FieldType.TABLE:
            continue
        fid = field.get('id', i)
        cols_raw = form.get(f'table_cols_{fid}')
        if not cols_raw and fid != i:
            cols_raw = form.get(f'table_cols_{i}')
        if not cols_raw:
            continue
        try:
            submitted_cols = json.loads(cols_raw)
        except (json.JSONDecodeError, TypeError):
            errors.append(f'{field.get("label", field.get("key"))} 的列定义格式错误')
            continue
        try:
            field['columns'] = normalize_table_columns(field, submitted_cols)
        except (ValueError, field_eval.FormulaError) as e:
            errors.append(f'{field.get("label", field.get("key"))} 的列定义无效：{e}')
    return errors


def parse_submitted_field_values(fields, form, allow_empty_keys=None):
    field_values = {}
    input_errors = []
    allow_empty_keys = set(allow_empty_keys or [])
    for i, field in enumerate(fields):
        key = field['key']
        fid = field.get('id', i)
        raw_val = form.get(f'field_{fid}')
        if raw_val is None and fid != i:
            raw_val = form.get(f'field_{i}')
        raw_val = (raw_val or '').strip()

        if field.get('field_type') == FieldType.TABLE:
            try:
                rows_data = json.loads(raw_val) if raw_val else []
            except (json.JSONDecodeError, TypeError):
                input_errors.append(f'{field.get("label", key)} 的表格数据格式错误')
                rows_data = []
            if not isinstance(rows_data, list):
                input_errors.append(f'{field.get("label", key)} 的表格数据必须是数组')
                rows_data = []
            if len(rows_data) > MAX_TABLE_ROWS:
                input_errors.append(f'{field.get("label", key)} 的表格行数不能超过 {MAX_TABLE_ROWS}')
                rows_data = []
            try:
                rows_data = filter_table_rows(field, rows_data)
            except ValueError as e:
                input_errors.append(str(e))
                rows_data = []
            if field.get('required') and not rows_data:
                input_errors.append(f'{field.get("label", key)}不能为空')
            field_values[key] = rows_data
        else:
            if len(raw_val) > MAX_TEXT_VALUE_LENGTH:
                from utils.logger import get_logger
                get_logger().warning('字段 %s(%s) 的值超过 %d 字符，已截断', field.get('label', ''), key, MAX_TEXT_VALUE_LENGTH)
            if (field.get('required') and key not in allow_empty_keys
                    and field.get('field_type') != FieldType.CALCULATED and not raw_val):
                input_errors.append(f'{field.get("label", key)}不能为空')
            if field.get('field_type') == FieldType.NUMBER and raw_val:
                try:
                    raw_val = normalize_number_field_value(raw_val, field)
                except ValueError as e:
                    input_errors.append(f'{field.get("label", key)}{e}')
            elif field.get('field_type') == FieldType.SELECT and raw_val:
                options = [str(option) for option in field.get('options', [])]
                if options and raw_val not in options:
                    input_errors.append(f'{field.get("label", key)}选项无效')
            field_values[key] = limit_text(raw_val)
    return field_values, input_errors
