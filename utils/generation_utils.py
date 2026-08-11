"""合同生成工具：计算、台账入账、批量、付款辅助 —— 从 helpers.py 拆分"""

import os
import secrets
from datetime import date, datetime, timedelta

from docx import Document

import docx_builder
import field_eval
import ledger_store
import payment_extractor
from utils.field_utils import (
    to_calc_number, parse_number, normalize_date,
    apply_submitted_table_columns, parse_submitted_field_values,
)
from utils.keyword_maps import (
    find_scalar_semantic,
)
from utils.logger import get_logger
from utils.security import (
    MAX_COUNTERPARTY_LENGTH, MAX_PROJECT_NAME_LENGTH, MAX_SUBSYSTEM_NAME_LENGTH,
    bounded_int, limit_text,
)
from utils.constants import FieldType


# ═══════════════════════════════════════════════════════
#  Template source binding validation
# ═══════════════════════════════════════════════════════

def validate_template_source_bindings(fields, source_docx_path: str) -> list[str]:
    """Check that stored field locations still point to the source document."""
    if not source_docx_path:
        return []
    errors = []
    try:
        doc = Document(source_docx_path)
    except Exception as e:
        return [f'无法读取模板源文件：{e}']

    paragraphs = list(doc.paragraphs)
    body_text = '\n'.join(p.text or '' for p in paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                body_text += '\n' + '\n'.join(p.text or '' for p in cell.paragraphs)

    for field in fields or []:
        label = field.get('label') or field.get('key') or '未命名字段'
        location = field.get('location') or {}
        loc_type = location.get('type')
        placeholder = location.get('placeholder') or ''
        if placeholder and placeholder not in body_text:
            errors.append(f'{label} 的占位符 {placeholder} 在源文档中不存在')
        if loc_type == 'paragraph':
            body_index = location.get('body_index', -1)
            if not isinstance(body_index, int) or body_index < 0 or body_index >= len(paragraphs):
                errors.append(f'{label} 的段落位置无效')
        elif loc_type == 'table':
            table_index = location.get('table_index', -1)
            row_index = location.get('template_row_index', -1)
            if not isinstance(table_index, int) or table_index < 0 or table_index >= len(doc.tables):
                errors.append(f'{label} 的表格位置无效')
            elif not isinstance(row_index, int) or row_index < 0 or row_index >= len(doc.tables[table_index].rows):
                errors.append(f'{label} 的表格模板行无效')
        elif loc_type == 'table_cell':
            table_index = location.get('table_index', -1)
            row_index = location.get('row_index', -1)
            col_index = location.get('col_index', -1)
            if not isinstance(table_index, int) or table_index < 0 or table_index >= len(doc.tables):
                errors.append(f'{label} 的表格位置无效')
                continue
            table = doc.tables[table_index]
            if not isinstance(row_index, int) or row_index < 0 or row_index >= len(table.rows):
                errors.append(f'{label} 的表格行位置无效')
                continue
            if not isinstance(col_index, int) or col_index < 0 or col_index >= len(table.rows[row_index].cells):
                errors.append(f'{label} 的表格列位置无效')
    return errors


# ═══════════════════════════════════════════════════════
#  Calculation helpers
# ═══════════════════════════════════════════════════════

def calc_context(fields, field_values):
    context = {}
    for field in fields:
        key = field.get('key')
        if not key:
            continue
        if field.get('field_type') == FieldType.TABLE:
            context[key] = {
                '__table_rows__': field_values.get(key, []) or [],
                '__table_columns__': {
                    col.get('key') for col in field.get('columns', []) if col.get('key')
                },
            }
        else:
            context[key] = to_calc_number(field_values.get(key, ''))
    return context


def recalculate_scalar_fields(fields, field_values):
    errors = []
    try:
        ordered_fields = field_eval.sort_fields_by_dependency(fields)
    except field_eval.FormulaError as e:
        return [str(e)]
    for field in ordered_fields:
        if field.get('field_type') != FieldType.CALCULATED:
            continue
        key = field.get('key')
        formula = field.get('formula', '')
        try:
            value = field_eval.safe_eval_decimal(formula, calc_context(fields, field_values))
            decimals = int(field.get('decimal_places', 2))
            field_values[key] = field_eval.format_number_text(value, decimals)
        except (field_eval.FormulaError, ValueError, TypeError) as e:
            errors.append(f'{field.get("label", key)} 公式计算失败：{e}')
    return errors


def recalculate_table_fields(fields, field_values):
    errors = []
    for field in fields:
        if field.get('field_type') != FieldType.TABLE:
            continue
        rows_data = field_values.get(field['key'], [])
        if not isinstance(rows_data, list):
            rows_data = []
        columns = field.get('columns', [])
        for row_index, row in enumerate(rows_data, start=1):
            if not isinstance(row, dict):
                continue
            for col in columns:
                if col.get('field_type') != 'calculated' or not col.get('formula'):
                    continue
                col_key = col.get('key')
                if not col_key:
                    continue
                try:
                    ctx = {}
                    for c in columns:
                        ck = c.get('key')
                        if not ck:
                            continue
                        ctx[ck] = to_calc_number(row.get(ck, '0'))
                    result = field_eval.safe_eval_decimal(col['formula'], ctx)
                    decimals = int(col.get('decimal_places', 2))
                    row[col_key] = field_eval.format_number_text(result, decimals)
                except (field_eval.FormulaError, ValueError, TypeError) as exc:
                    row[col_key] = ''
                    errors.append(
                        f'{field.get("label", field.get("key", "表格"))}第 {row_index} 行'
                        f'“{col.get("label", col_key)}”公式计算失败：{exc}'
                    )
    return errors


def prepare_generation_values(fields, form, allow_empty_keys=None):
    """统一执行生成/预览前的表单解析、校验和公式计算。"""
    errors = apply_submitted_table_columns(fields, form)
    field_values, parse_errors = parse_submitted_field_values(
        fields, form, allow_empty_keys=allow_empty_keys
    )
    errors.extend(parse_errors)
    if errors:
        return field_values, errors
    errors.extend(recalculate_table_fields(fields, field_values))
    if errors:
        return field_values, errors
    errors.extend(recalculate_scalar_fields(fields, field_values))
    return field_values, errors


# ═══════════════════════════════════════════════════════
#  Contract generation helpers
# ═══════════════════════════════════════════════════════

def _find_field_by_semantic(fields, semantic_name):
    for field in fields:
        if field.get('field_type') == 'table':
            continue
        label = str(field.get('label') or '')
        key = str(field.get('key') or '')
        if find_scalar_semantic(label, key) == semantic_name:
            return field
    return None


def _value_by_semantic(fields, field_values, semantic_name, numeric=False, date_value=False):
    field = _find_field_by_semantic(fields, semantic_name)
    if not field:
        return None
    raw = field_values.get(field['key'], '')
    if numeric:
        parsed = parse_number(raw)
        return parsed if parsed is not None else None
    if date_value:
        return normalize_date(raw) or None
    return str(raw).strip() or None


def infer_contract_summary(tpl, fields, field_values):
    amount = _value_by_semantic(fields, field_values, 'amount', numeric=True)
    sign_date = _value_by_semantic(fields, field_values, 'sign_date', date_value=True)
    contract_no = _value_by_semantic(fields, field_values, 'contract_no')
    title = _value_by_semantic(fields, field_values, 'title')
    counterparty = _value_by_semantic(fields, field_values, 'counterparty')
    owner = _value_by_semantic(fields, field_values, 'owner')
    if not contract_no:
        contract_no = 'HT' + datetime.now().strftime('%Y%m%d%H%M%S') + secrets.token_hex(4)  # 8位随机后缀防碰撞
    return {
        'contract_no': contract_no,
        'title': title or tpl.name or '未命名合同',
        'counterparty': limit_text(counterparty or '', MAX_COUNTERPARTY_LENGTH),
        'amount': amount,
        'sign_date': sign_date or '',
        'owner': owner or '',
        'status': 'draft',
        'template_name': tpl.name,
    }


def parse_contract_classification(form):
    """Parse the required coverage mode and its classification fields."""
    project_name = limit_text(
        str(form.get('project_name', '') or '').strip(),
        MAX_PROJECT_NAME_LENGTH,
    )
    subsystem_name = limit_text(
        str(form.get('subsystem_name', '') or '').strip(),
        MAX_SUBSYSTEM_NAME_LENGTH,
    )
    start_raw = str(form.get('coverage_start', '') or '').strip()
    end_raw = str(form.get('coverage_end', '') or '').strip()
    coverage_mode = str(form.get('coverage_mode', '') or '').strip()

    if coverage_mode not in {'range', 'not_applicable'}:
        raise ValueError('请选择“填写数字范围”或“不适用”')

    if coverage_mode == 'not_applicable':
        if start_raw or end_raw:
            raise ValueError('选择“不适用”时不能填写起始发次或结束发次')
        return {
            'project_name': project_name,
            'subsystem_name': subsystem_name,
            'coverage_mode': coverage_mode,
            'coverage_not_applicable': True,
            'coverage_start': None,
            'coverage_end': None,
        }

    if not start_raw or not end_raw:
        raise ValueError('填写数字范围时，起始发次和结束发次必须同时填写')
    if not project_name:
        raise ValueError('填写发次范围前，请先填写项目名称')

    coverage_start = bounded_int(
        start_raw, min_value=1, max_value=1_000_000_000, label='起始发次'
    )
    coverage_end = bounded_int(
        end_raw, min_value=1, max_value=1_000_000_000, label='结束发次'
    )
    if coverage_start > coverage_end:
        raise ValueError('起始发次不能大于结束发次')

    return {
        'project_name': project_name,
        'subsystem_name': subsystem_name,
        'coverage_mode': coverage_mode,
        'coverage_not_applicable': False,
        'coverage_start': coverage_start,
        'coverage_end': coverage_end,
    }


def prepare_ledger_record(tpl, fields, field_values, document_path, classification=None):
    """Build the ledger summary and extracted plans without opening a transaction."""
    summary = infer_contract_summary(tpl, fields, field_values)
    if classification:
        summary.update({
            'project_name': classification.get('project_name') or '',
            'subsystem_name': classification.get('subsystem_name') or '',
            'coverage_mode': classification.get('coverage_mode') or '',
            'coverage_not_applicable': bool(
                classification.get('coverage_not_applicable')
            ),
            'coverage_start': classification.get('coverage_start'),
            'coverage_end': classification.get('coverage_end'),
        })
    # 付款计划提取在事务外完成（需要读取 DOCX 文件）
    try:
        doc_text = payment_extractor.extract_docx_text(document_path)
        plans = payment_extractor.extract_payment_plans(
            doc_text,
            contract_amount=summary.get('amount'),
            sign_date=summary.get('sign_date') or '',
        )
    except Exception:
        get_logger().error('Payment text extraction failed for %s', document_path, exc_info=True)
        plans = []
    return summary, plans


def prepare_ledger_payments(
    tpl, fields, field_values, document_path, classification=None
):
    """Build a ledger summary plus separated payment plans and rules."""
    summary = infer_contract_summary(tpl, fields, field_values)
    if classification:
        summary.update({
            'project_name': classification.get('project_name') or '',
            'subsystem_name': classification.get('subsystem_name') or '',
            'coverage_mode': classification.get('coverage_mode') or '',
            'coverage_not_applicable': bool(
                classification.get('coverage_not_applicable')
            ),
            'coverage_start': classification.get('coverage_start'),
            'coverage_end': classification.get('coverage_end'),
        })
    try:
        blocks = payment_extractor.extract_docx_blocks(document_path)
        extraction = payment_extractor.extract_payment_items(
            blocks,
            contract_amount=summary.get('amount'),
            sign_date=summary.get('sign_date') or '',
        )
        return summary, extraction.plans, extraction.rules
    except Exception:
        get_logger().error(
            'Payment rule extraction failed for %s', document_path, exc_info=True
        )
        return summary, [], []


def create_ledger_record(
    tpl, fields, field_values, output_path, classification=None, *, conn=None,
    document_path=None,
):
    """Create one contract and its payment plans in the supplied transaction."""
    summary, plans, rules = prepare_ledger_payments(
        tpl, fields, field_values, document_path or output_path, classification
    )
    # 合同创建与付款计划插入在同一事务中完成
    contract_id, plan_count = ledger_store.create_contract_with_plans(
        summary, field_values, output_path, plans, conn=conn, rules=rules,
    )
    if plan_count:
        get_logger().info('Created contract %d with %d payment plans', contract_id, plan_count)
    return contract_id


def docx_write_order(fields):
    """返回写入顺序：表格字段优先（表头+数据行），计算字段最后，输入/文本按依赖排序"""
    ordered = field_eval.sort_fields_by_dependency(fields)
    tables = [f for f in ordered if f.get('field_type') == FieldType.TABLE]
    calcs = [f for f in ordered if f.get('field_type') == FieldType.CALCULATED]
    # 非表格、非计算的输入型字段
    inputs = [f for f in ordered if f.get('field_type') not in (FieldType.TABLE, FieldType.CALCULATED)]
    return tables + inputs + calcs


def generate_docx_document(tpl_data, fields, field_values, source_docx, output_path):
    """统一的合同文档生成逻辑 —— generate() 与 generate_batch() 共用。

    返回 (errors, output_path)，errors 为空列表表示成功。
    """
    errors = []
    if source_docx:
        if not os.path.exists(source_docx):
            errors.append('找不到模板源文件，文件可能已被清理')
            return errors, output_path
        doc = Document(source_docx)
        ordered_fields = docx_write_order(fields)
        for field in ordered_fields:
            ftype = field['field_type']
            key = field['key']
            location = field.get('location', {})
            if ftype == 'table':
                try:
                    docx_builder.apply_table_field(doc, field, field_values.get(key, []))
                except Exception as e:
                    errors.append(f'{field.get("label", key)} 写入失败：{e}')
                    get_logger().error('Table field write failed: %s', e, exc_info=True)
            else:
                try:
                    docx_builder.apply_text_field(doc, location, field_values.get(key, ''), field.get('label', ''), key)
                except Exception as e:
                    errors.append(f'{field.get("label", key)} 写入失败：{e}')
                    get_logger().error('Text field write failed: %s', e, exc_info=True)
        if not errors:
            try:
                tmp_path = output_path + '.tmp'
                doc.save(tmp_path)
                os.replace(tmp_path, output_path)
            except Exception as e:
                errors.append(f'文档保存失败：{e}')
                get_logger().error('doc.save failed: %s', e, exc_info=True)
                # 清理可能残留的临时文件
                try:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)
                except OSError:
                    get_logger().warning('无法删除合同生成临时文件: %s', tmp_path, exc_info=True)
    else:
        docx_builder.generate_from_scratch(tpl_data, field_values, output_path)
    return errors, output_path


# ═══════════════════════════════════════════════════════
#  Batch helpers
# ═══════════════════════════════════════════════════════

def counterparty_batch_keys(fields, submitted_key=''):
    keys = []
    if submitted_key:
        target = next((f for f in fields if f.get('key') == submitted_key), None)
        keys.append(submitted_key)
        if target:
            label = target.get('label')
            if label is not None:
                for field in fields:
                    key = field.get('key')
                    if key and key != submitted_key and field.get('label') is not None and field.get('label') == label:
                        keys.append(key)
        return list(dict.fromkeys(keys))

    for field in fields:
        if field.get('field_type') == 'table':
            continue
        label = str(field.get('label') or '')
        key = str(field.get('key') or '')
        if find_scalar_semantic(label, key) == 'counterparty':
            if key:
                keys.append(key)
    return list(dict.fromkeys(keys))


def contract_number_keys(fields):
    keys = []
    for field in fields:
        if field.get('field_type') == 'table':
            continue
        label = str(field.get('label') or '')
        key = str(field.get('key') or '')
        if find_scalar_semantic(label, key) == 'contract_no':
            if key:
                keys.append(key)
    return keys


# ═══════════════════════════════════════════════════════
#  Date helpers
# ═══════════════════════════════════════════════════════

def next_month_ym(today=None):
    """返回下个月的年、月（整数元组），处理跨年。"""
    today = today or date.today()
    m = today.month + 1
    y = today.year
    if m > 12:
        m = 1
        y += 1
    return y, m


def next_month_range(today=None):
    today = today or date.today()
    if today.month == 12:
        start = date(today.year + 1, 1, 1)
    else:
        start = date(today.year, today.month + 1, 1)
    if start.month == 12:
        following = date(start.year + 1, 1, 1)
    else:
        following = date(start.year, start.month + 1, 1)
    end = following - timedelta(days=1)
    return start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')


# ═══════════════════════════════════════════════════════
#  Payment plan helpers
# ═══════════════════════════════════════════════════════

def has_payment_content(row):
    return any(str(row.get(key) or '').strip() for key in [
        'phase_name', 'trigger_event', 'due_date', 'ratio', 'due_amount',
        'condition_text', 'source_text', 'remark'
    ])


def can_bulk_confirm_payment(plan):
    parse_status = str(plan.get('parse_status') or '').strip()
    if parse_status in {'conflict', 'unsupported'}:
        return False
    if not parse_status and (plan.get('confidence') or 'low') == 'low':
        return False
    if plan.get('due_amount') is None and plan.get('ratio') is None:
        return False
    trigger_event = str(plan.get('trigger_event') or '').strip()
    if not plan.get('due_date') and trigger_event in ('', '其他'):
        return False
    return True
